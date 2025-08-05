import azure.functions as func
import logging
import json
import traceback
import base64
import os

from io import BytesIO, StringIO

import fitz  # PyMuPDF

import csv

from concurrent.futures import ThreadPoolExecutor, as_completed

from datetime import datetime

from typing import Dict, List, Any, Optional

import numpy as np

from uuid import uuid4

from openpyxl.worksheet.table import Table as ExcelTable, TableStyleInfo
from openpyxl.utils import get_column_letter, column_index_from_string

from pandas import DataFrame

from openpyxl import load_workbook, Workbook

from xlrd import open_workbook

import tempfile

from PIL import Image

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4

from pdf2docx import Converter

import mammoth

import xml.etree.ElementTree as ET

from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table as WordTable
from docx.text.paragraph import Paragraph

from docx import Document

from docx.enum.text import WD_BREAK

import glob
from lxml import etree
from copy import deepcopy

import shutil

from collections import defaultdict

import re

import html

from extract_msg import Message as ExtractMessage

from docx.shared import Inches

from docx.opc.constants import RELATIONSHIP_TYPE

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
import requests

from zipfile import ZipFile, ZIP_DEFLATED, BadZipFile

app = func.FunctionApp(http_auth_level=func.AuthLevel.FUNCTION)



@app.route(route="upsert_json_array")
def upsert_json_array(req: func.HttpRequest) -> func.HttpResponse:
    try:
        req_body = req.get_json()

        primary_key = req_body.get("primary_key")
        initial_json = req_body.get("initial_json", [])
        update_json = req_body.get("update_json", [])

        if not all([primary_key, initial_json, update_json]):
            return func.HttpResponse(
                json.dumps({"error": "Missing required parameters."}),
                status_code=400,
                mimetype="application/json"
            )

        # Extract all possible columns from initial_json
        all_columns = {key for item in initial_json for key in item.keys()}

        # Organize initial_json by primary key (one-to-many relationship)
        initial_dict = {}
        for item in initial_json:
            key = item[primary_key]
            if key not in initial_dict:
                initial_dict[key] = []
            initial_dict[key].append(item)

        # Convert update_json to a dictionary
        update_dict = {item[primary_key]: item for item in update_json}

        new_keys = set(update_dict.keys()) - set(initial_dict.keys())

        # Update existing records
        for key, records in initial_dict.items():
            if key in update_dict:
                update_values = update_dict[key]
                for record in records:
                    for update_field, update_value in update_values.items():
                        if update_field in record:
                            record[update_field] = update_value

        # Append new records with missing columns set to None
        for key in new_keys:
            new_record = {col: update_dict[key].get(col, None) for col in all_columns}
            initial_json.append(new_record)

        return func.HttpResponse(
            json.dumps(initial_json, indent=2),
            status_code=200,
            mimetype="application/json"
        )

    except Exception as e:
        error_details = traceback.format_exc()
        logging.error("Exception occurred: %s", error_details)
        return func.HttpResponse(
            json.dumps({"error": "Internal server error", "details": str(e), "traceback": error_details}),
            status_code=500,
            mimetype="application/json"
        )






@app.route(route="csv_to_json")
def csv_to_json(req: func.HttpRequest) -> func.HttpResponse:
    try:
        # Get parameters from request body as JSON
        req_body = req.get_body().decode('utf-8')
        if not req_body:
            return func.HttpResponse("Request body is empty", status_code=400)

        try:
            req_json = json.loads(req_body)
        except json.JSONDecodeError:
            return func.HttpResponse("Invalid JSON in request body", status_code=400)

        # Extract parameters from JSON with additional flexibility options
        file_content = req_json.get('file_content')
        delimiter = req_json.get('delimiter', ',')
        linebreak = req_json.get('linebreak', '\r\n')
        type_detect = req_json.get('type_detect', 'YES').upper() == 'YES'
        first_row_headers = req_json.get('first_row_headers', 'YES').upper() == 'YES'
        skip_top_rows = max(0, int(req_json.get('skip_top_rows', 0)))
        skip_bottom_rows = max(0, int(req_json.get('skip_bottom_rows', 0)))

        if not file_content:
            return func.HttpResponse("Missing required parameter: file_content", status_code=400)

        # Handle custom linebreak characters
        file_content = file_content.replace('\\r\\n', '\r\n').replace('\\n', '\n')
        if linebreak != '\r\n' and linebreak != '\n':
            # Convert string representation to actual line break characters
            linebreak = linebreak.replace('\\r\\n', '\r\n').replace('\\n', '\n')
            file_content = file_content.replace(linebreak, '\r\n')

        # Create CSV reader with custom delimiter
        csv_reader = csv.reader(StringIO(file_content), delimiter=delimiter)
        rows = list(csv_reader)

        if not rows:
            return func.HttpResponse(json.dumps([]), mimetype="application/json")

        # Apply row skipping logic
        if skip_top_rows > 0:
            if skip_top_rows >= len(rows):
                return func.HttpResponse("Skip top rows value exceeds number of rows in CSV", status_code=400)
            rows = rows[skip_top_rows:]

        if skip_bottom_rows > 0:
            if skip_bottom_rows >= len(rows):
                return func.HttpResponse("Skip bottom rows value exceeds number of rows in CSV", status_code=400)
            rows = rows[:-skip_bottom_rows]

        if not rows:
            return func.HttpResponse(json.dumps([]), mimetype="application/json")

        # Process headers after row skipping
        if first_row_headers:
            headers = rows.pop(0)
            # Ensure header names are valid and make them unique
            headers = handle_duplicate_headers([h if h else f"Column_{i+1}" for i, h in enumerate(headers)])
        else:
            headers = [f"Column_{i+1}" for i in range(len(rows[0]))]

        if type_detect and rows:
            column_types = detect_column_types(rows)
            json_data = []
            for row in rows:
                row_dict = {}
                for i, value in enumerate(row):
                    if i < len(headers):  # Protect against malformed rows
                        try:
                            row_dict[headers[i]] = column_types[i](value) if value else None
                        except:
                            row_dict[headers[i]] = value
                json_data.append(row_dict)
        else:
            json_data = []
            for row in rows:
                # Handle cases where rows might have different lengths
                row_dict = {}
                for i, value in enumerate(row):
                    if i < len(headers):
                        row_dict[headers[i]] = value
                json_data.append(row_dict)

        return func.HttpResponse(
            json.dumps(json_data, ensure_ascii=False), 
            mimetype="application/json"
        )

    except csv.Error as e:
        logging.error(f"CSV parsing error: {str(e)}")
        return func.HttpResponse(f"Error parsing CSV: {str(e)}", status_code=400)
    except json.JSONDecodeError as e:
        logging.error(f"JSON encoding error: {str(e)}")
        return func.HttpResponse(f"Error creating JSON: {str(e)}", status_code=500)
    except ValueError as e:
        logging.error(f"Value error: {str(e)}")
        return func.HttpResponse(f"Invalid parameter value: {str(e)}", status_code=400)
    except Exception as e:
        logging.error(traceback.format_exc())
        return func.HttpResponse(f"Error processing CSV: {str(e)}", status_code=500)

def handle_duplicate_headers(headers):
    """
    Makes duplicate headers unique by appending an index.
    Example: ["Name", "Name", "Name"] becomes ["Name", "Name2", "Name3"]
    """
    seen = {}
    unique_headers = []

    for header in headers:
        if header in seen:
            seen[header] += 1
            unique_headers.append(f"{header}_{seen[header]}")
        else:
            seen[header] = 1
            unique_headers.append(header)

    return unique_headers

def detect_column_types(rows):
    """Detects column data types based on all rows."""
    if not rows or not rows[0]:
        return []

    num_cols = max(len(row) for row in rows)
    column_types = [str] * num_cols

    # Sample up to 100 rows for type detection
    sample_rows = rows[:min(100, len(rows))]

    for col_idx in range(num_cols):
        col_values = [row[col_idx] for row in sample_rows if col_idx < len(row) and row[col_idx]]
        if not col_values:
            continue

        # Check for boolean values
        if all(val.lower() in ('true', 'false', 'yes', 'no', '1', '0') for val in col_values):
            column_types[col_idx] = lambda x: x.lower() in ('true', 'yes', '1')
            continue

        # Check for integers
        if all(re.match(r'^-?\d+$', val) for val in col_values):
            column_types[col_idx] = int
            continue

        # Check for floats
        if all(is_float(val) for val in col_values):
            column_types[col_idx] = float
            continue

        # Check for dates (simple ISO format check)
        if all(re.match(r'^\d{4}-\d{2}-\d{2}', val) for val in col_values):
            column_types[col_idx] = lambda x: x  # Keep as string for dates

    return column_types

def is_float(value):
    try:
        float(value)
        return True
    except (ValueError, TypeError):
        return False






@app.route(route="dummy")
async def dummy(req: func.HttpRequest) -> func.HttpResponse:
    # Dummy function that only returns a 200 response. Can call this function to keep the function app active & warm for other functions.
    return func.HttpResponse(
        status_code=200,
        mimetype="application/json"
    )




@app.route(route="aggregations")
def aggregations(req: func.HttpRequest) -> func.HttpResponse:
    try:
        # Get request body
        req_body = req.get_json()

        # Extract data and aggregation parameters
        data = req_body.get('data', [])
        min_cols = req_body.get('min', [])
        max_cols = req_body.get('max', [])
        sum_cols = req_body.get('sum', [])
        count_cols = req_body.get('count', [])
        avg_cols = req_body.get('avg', [])
        med_cols = req_body.get('med', [])
        stdev_cols = req_body.get('stdev', [])

        # Handle empty data
        if not data:
            return func.HttpResponse(json.dumps({}), mimetype="application/json")

        # Detect if data is an array of arrays or a single array
        is_array_of_arrays = all(isinstance(item, list) for item in data)

        # Process data accordingly
        if is_array_of_arrays:
            # Handle array of arrays case
            if all(len(arr) == 0 for arr in data):
                return func.HttpResponse(json.dumps([]), mimetype="application/json")

            results = []
            for data_array in data:
                result = process_aggregations(
                    data_array, min_cols, max_cols, sum_cols, 
                    count_cols, avg_cols, med_cols, stdev_cols
                )
                results.append(result)

            return func.HttpResponse(json.dumps(results), mimetype="application/json")
        else:
            # Handle single array case
            result = process_aggregations(
                data, min_cols, max_cols, sum_cols, 
                count_cols, avg_cols, med_cols, stdev_cols
            )

            return func.HttpResponse(json.dumps(result), mimetype="application/json")

    except Exception as e:
        error_details = traceback.format_exc()
        logging.error(f"Error processing request: {str(e)}\nTraceback: {error_details}")
        return func.HttpResponse(
            json.dumps({"error": str(e), "Traceback": error_details}),
            status_code=400,
            mimetype="application/json"
        )

def process_aggregations(
    data: List[Dict], 
    min_cols: List[str], 
    max_cols: List[str], 
    sum_cols: List[str], 
    count_cols: List[str], 
    avg_cols: List[str], 
    med_cols: List[str], 
    stdev_cols: List[str]
) -> Dict:
    """
    Process all requested aggregations on the data.
    """
    result = {}

    # Minimum aggregations
    for col in min_cols:
        result[f"min_{col}"] = calculate_min(data, col)

    # Maximum aggregations
    for col in max_cols:
        result[f"max_{col}"] = calculate_max(data, col)

    # Sum aggregations
    for col in sum_cols:
        result[f"sum_{col}"] = calculate_sum(data, col)

    # Count aggregations
    for col in count_cols:
        result[f"count_{col}"] = calculate_count(data, col)

    # Average aggregations
    for col in avg_cols:
        result[f"avg_{col}"] = calculate_avg(data, col)

    # Median aggregations
    for col in med_cols:
        result[f"med_{col}"] = calculate_median(data, col)

    # Standard deviation aggregations
    for col in stdev_cols:
        result[f"stdev_{col}"] = calculate_stdev(data, col)

    return result

def try_numeric_conversion(value):
    """Attempt to convert a value to a numeric type if it's a string number."""
    if not isinstance(value, str):
        return value

    try:
        # Try integer first
        return int(value)
    except ValueError:
        try:
            # Then try float
            return float(value)
        except ValueError:
            # Not a numeric string
            return value

def calculate_min(data: List[Dict], column: str) -> Any:
    """Calculate minimum value for a column across all data points."""
    # Extract values and attempt numeric conversion for string numbers
    values = []
    for item in data:
        if item.get(column) is not None:
            values.append(try_numeric_conversion(item.get(column)))

    if not values:
        return None

    # Try to determine the type of values
    sample_value = values[0]

    if isinstance(sample_value, (int, float)):
        # Get only the numeric values (including converted string numbers)
        numeric_values = [v for v in values if isinstance(v, (int, float))]
        if not numeric_values:
            return None
        return min(numeric_values)
    elif isinstance(sample_value, bool):
        # Boolean min (False < True)
        boolean_values = [v for v in values if isinstance(v, bool)]
        if not boolean_values:
            return None
        return False if False in boolean_values else True
    elif isinstance(sample_value, str):
        # Try to parse as date first
        try:
            # Attempt to parse dates
            date_values = []
            original_values = []
            for v in values:
                if isinstance(v, str):
                    try:
                        date_values.append(datetime.fromisoformat(v.replace('Z', '+00:00')))
                        original_values.append(v)
                    except (ValueError, AttributeError):
                        pass

            if date_values:
                min_date = min(date_values)
                # Find the original string that corresponds to this date
                for i, date in enumerate(date_values):
                    if date == min_date:
                        return original_values[i]

            # If not dates or empty date_values, treat as strings
            string_values = [v for v in values if isinstance(v, str)]
            if not string_values:
                return None
            return min(string_values)
        except (ValueError, AttributeError):
            # If error in date parsing, treat as strings
            string_values = [v for v in values if isinstance(v, str)]
            if not string_values:
                return None
            return min(string_values)

    # Fallback
    return min(values, default=None)

def calculate_max(data: List[Dict], column: str) -> Any:
    """Calculate maximum value for a column across all data points."""
    # Extract values and attempt numeric conversion for string numbers
    values = []
    for item in data:
        if item.get(column) is not None:
            values.append(try_numeric_conversion(item.get(column)))

    if not values:
        return None

    # Try to determine the type of values
    sample_value = values[0]

    if isinstance(sample_value, (int, float)):
        # Get only the numeric values (including converted string numbers)
        numeric_values = [v for v in values if isinstance(v, (int, float))]
        if not numeric_values:
            return None
        return max(numeric_values)
    elif isinstance(sample_value, bool):
        # Boolean max (True > False)
        boolean_values = [v for v in values if isinstance(v, bool)]
        if not boolean_values:
            return None
        return True if True in boolean_values else False
    elif isinstance(sample_value, str):
        # Try to parse as date first
        try:
            # Attempt to parse dates
            date_values = []
            original_values = []
            for v in values:
                if isinstance(v, str):
                    try:
                        date_values.append(datetime.fromisoformat(v.replace('Z', '+00:00')))
                        original_values.append(v)
                    except (ValueError, AttributeError):
                        pass

            if date_values:
                max_date = max(date_values)
                # Find the original string that corresponds to this date
                for i, date in enumerate(date_values):
                    if date == max_date:
                        return original_values[i]

            # If not dates or empty date_values, treat as strings
            string_values = [v for v in values if isinstance(v, str)]
            if not string_values:
                return None
            return max(string_values)
        except (ValueError, AttributeError):
            # If error in date parsing, treat as strings
            string_values = [v for v in values if isinstance(v, str)]
            if not string_values:
                return None
            return max(string_values)

    # Fallback
    return max(values, default=None)

def calculate_sum(data: List[Dict], column: str) -> Optional[float]:
    """Calculate sum of values for a column across all data points."""
    numeric_values = []

    for item in data:
        if item.get(column) is not None:
            value = try_numeric_conversion(item.get(column))
            if isinstance(value, (int, float)):
                numeric_values.append(value)

    if not numeric_values:
        return None

    return sum(numeric_values)

def calculate_count(data: List[Dict], column: str) -> int:
    """Calculate count of non-null values for a column across all data points."""
    return sum(1 for item in data if item.get(column) is not None)

def calculate_avg(data: List[Dict], column: str) -> Optional[float]:
    """Calculate average of values for a column across all data points."""
    numeric_values = []

    for item in data:
        if item.get(column) is not None:
            value = try_numeric_conversion(item.get(column))
            if isinstance(value, (int, float)):
                numeric_values.append(value)

    if not numeric_values:
        return None

    return sum(numeric_values) / len(numeric_values)

def calculate_median(data: List[Dict], column: str) -> Optional[float]:
    """Calculate median of values for a column across all data points."""
    numeric_values = []

    for item in data:
        if item.get(column) is not None:
            value = try_numeric_conversion(item.get(column))
            if isinstance(value, (int, float)):
                numeric_values.append(value)

    if not numeric_values:
        return None

    return float(np.median(numeric_values))

def calculate_stdev(data: List[Dict], column: str) -> Optional[float]:
    """Calculate standard deviation of values for a column across all data points."""
    numeric_values = []

    for item in data:
        if item.get(column) is not None:
            value = try_numeric_conversion(item.get(column))
            if isinstance(value, (int, float)):
                numeric_values.append(value)

    if len(numeric_values) < 2:  # Need at least 2 values for standard deviation
        return None

    return float(np.std(numeric_values, ddof=1))  # Using sample standard deviation






@app.route(route="regex_find_and_replace")
def regex_find_and_replace(req: func.HttpRequest) -> func.HttpResponse:
    try:
        req_body = req.get_json()
        input_string = req_body.get("string")
        regex_pattern = req_body.get("regex")
        replace_text = req_body.get("replace_text", None)

        if not input_string or not regex_pattern:
            return func.HttpResponse(
                json.dumps({"error": "Both 'string' and 'regex' parameters are required."}),
                status_code=400,
                mimetype="application/json"
            )

        matches = []
        new_string = input_string

        # Find matches first
        for match in re.finditer(regex_pattern, input_string):
            match_data = {
                "value": match.group(),
                "start_index": match.start(),
                "end_index": match.end()
            }
            matches.append(match_data)

        # Perform replacement if replace_text is provided
        if replace_text is not None:
            updated_matches = []
            new_string = re.sub(regex_pattern, replace_text, input_string)
            current_index = 0

            for match in re.finditer(regex_pattern, input_string):
                new_match_start = new_string.find(replace_text, current_index)
                new_match_end = new_match_start + len(replace_text)
                current_index = new_match_end

                updated_match_data = {
                    "value": match.group(),
                    "start_index": match.start(),
                    "end_index": match.end(),
                    "new_value": replace_text,
                    "new_start_index": new_match_start,
                    "new_end_index": new_match_end
                }
                updated_matches.append(updated_match_data)

            matches = updated_matches

            response = {"matches": matches, "new_string": new_string}

        else:
            response = {"matches": matches}

        return func.HttpResponse(
            json.dumps(response),
            status_code=200,
            mimetype="application/json"
        )
    except Exception as e:
        return func.HttpResponse(
            json.dumps({
                "error": "An error occurred.",
                "details": str(e),
                "traceback": traceback.format_exc()
            }),
            status_code=500,
            mimetype="application/json"
        )

### Python function for transforming array items ###





@app.route(route="py_transform_array")
def py_transform_array(req: func.HttpRequest) -> func.HttpResponse:
    try:
        # Parse request body
        req_body = req.get_json()
        transform_from = req_body.get('transform_from')
        transform_expr = req_body.get('transformation')

        # Basic validation for required fields
        if not transform_expr or not transform_from:
            return func.HttpResponse("Error: 'transform_expr' and 'transform_from' are required fields", status_code=400)

        # Attempt to create the transform function using eval()
        try:
            transform_item = eval('lambda item: ' + transform_expr)
        except Exception as e:
            return func.HttpResponse(f"Error: Invalid transformation expression - {str(e)}", status_code=400)

        # Process each item applying the transformation
        results = [transform_item(item) for item in transform_from]

        # Return the transformed array directly
        return func.HttpResponse(json.dumps(results), mimetype="application/json")

    except Exception as e:
        return func.HttpResponse(f"Error: {str(e)}", status_code=500)






### Simple python filter array py_filter
### Return a JSON array of all records matching the filter_query

@app.route(route="py_filter")
def py_filter(req: func.HttpRequest) -> func.HttpResponse:
    try:
        # Parse request body
        req_body = req.get_json()
        filter_from = req_body.get('filter_from')
        filter_query = req_body.get('filter_query')
        # filter_query example 1 #
        '''item['key'] != "String1"'''
        # filter_query example 2 #
        '''(item['key'] not in "String1 String2" and item['numeric_key'] >= 0) or regex(r'^A.*', item['key'])'''

        # Basic validation for required fields 
        if not filter_query or not filter_from: 
            return func.HttpResponse("Error: 'filter_query' and 'filter_from' are required fields", status_code=400)

        # Attempt to create the filter function using eval()
        try:
            filter_item = eval('lambda item: ' + filter_query)
        except Exception as e:
            return func.HttpResponse(f"Error: Invalid filter query - {str(e)}", status_code=400)

        # Process each item applying the filter
        results = [item for item in filter_from if filter_item(item)]

        # Return the response
        response = {
            "values": results,
        }
        return func.HttpResponse(json.dumps(response), mimetype="application/json")

    except Exception as e:
        return func.HttpResponse(f"Error: {str(e)}", status_code=500)



### Lookup the 1st record in lookup_from matching the lookup_query for each value in loop_array
### Return a single JSON array with the 1st record found for each loop_array value

@app.route(route="for_each_lookup")
def for_each_lookup(req: func.HttpRequest) -> func.HttpResponse:
    try:
        # Parse request body
        req_body = req.get_json()
        loop_array = req_body.get('loop_array', [0])
        lookup_from = req_body.get('lookup_from')
        lookup_query = req_body.get('lookup_query')
        # lookup_query example 1 #
        '''item['key'] != LoopValue'''
        # lookup_query example 2 #
        '''LoopValue['key'].upper() in item['key'].upper())'''
        # lookup_query example 3 #
        '''(item['key'] not in 'String1 String2' and item['numeric_key'] >= 0) or (item['key'] != LoopValue['key'] and regex(r'^A.*', item['key']))'''

        # Basic validation for required fields 
        if not lookup_query or not lookup_from: 
            return func.HttpResponse("Error: 'lookup_query' and 'lookup_from' are required fields", status_code=400)

        # Attempt to create the filter function using eval()
        try:
            filter_item = eval('lambda item, LoopValue: ' + lookup_query)
        except Exception as e:
            return func.HttpResponse(f"Error: Invalid lookup query - {str(e)}", status_code=400)

        # Function to process each item
        def process_item(LoopValue):
            for item in lookup_from:
                if filter_item(item, LoopValue):
                    item['LoopValue'] = LoopValue
                    return item, None
            return None, LoopValue

        # Use ThreadPoolExecutor to manage concurrent tasks
        values = []
        no_value_loop_array = []
        with ThreadPoolExecutor() as executor:
            future_to_iterator = {executor.submit(process_item, LoopValue): LoopValue for LoopValue in loop_array}
            for future in as_completed(future_to_iterator):
                try:
                    value, no_value_loop = future.result()
                    if value:
                        values.append(value)
                    elif no_value_loop:
                        no_value_loop_array.append(no_value_loop)
                except Exception as e:
                    return func.HttpResponse(f"Error: {str(e)}", status_code=500)

        # Return the response
        response = {
            "values": values,
            "no_value_loop_array": no_value_loop_array
        }
        return func.HttpResponse(json.dumps(response), mimetype="application/json")

    except Exception as e:
        return func.HttpResponse(f"Error: {str(e)}", status_code=500)

def regex(pattern, value):
    return re.search(pattern, value) is not None



### Filter to all records in filter_from matching the filter_query for each value in loop_array
### Return an array of JSON arrays, one for each set of records found for each loop_array value

@app.route(route="for_each_filter")
def for_each_filter(req: func.HttpRequest) -> func.HttpResponse:
    try:
        # Parse request body
        req_body = req.get_json()
        loop_array = req_body.get('loop_array', [0])
        filter_query = req_body.get('filter_query')
        filter_from = req_body.get('filter_from')

        # Basic validation for required fields 
        if not filter_query or not filter_from: 
            return func.HttpResponse("Error: 'filter_query' and 'filter_from' are required fields", status_code=400)

        # Attempt to create the filter function using eval()
        try:
            filter_item = eval('lambda item, LoopValue: ' + filter_query)
        except Exception as e:
            return func.HttpResponse(f"Error: Invalid filter query - {str(e)}", status_code=400)

        # Function to process each item 
        def process_item(LoopValue): 
            results = [item for item in filter_from if filter_item(item, LoopValue)] 
            if results: 
                for result in results: 
                    result['LoopValue'] = LoopValue 
                return results, None 
            return None, LoopValue

        # Use ThreadPoolExecutor to manage concurrent tasks
        values = []
        no_value_loop_array = []
        with ThreadPoolExecutor() as executor:
            future_to_iterator = {executor.submit(process_item, LoopValue): LoopValue for LoopValue in loop_array}
            for future in as_completed(future_to_iterator):
                try:
                    value, no_value_loop = future.result()
                    if value:
                        values.append(value)
                    elif no_value_loop:
                        no_value_loop_array.append(no_value_loop)
                except Exception as e:
                    return func.HttpResponse(f"Error: {str(e)}", status_code=500)

        # Return the response
        response = {
            "values": values,
            "no_value_loop_array": no_value_loop_array
        }
        return func.HttpResponse(json.dumps(response), mimetype="application/json")

    except Exception as e:
        return func.HttpResponse(f"Error: {str(e)}", status_code=500)

def regex(pattern, value):
    return re.search(pattern, value) is not None







@app.route(route="index_duplicates")
def index_duplicates(req: func.HttpRequest) -> func.HttpResponse:
    try:
        req_body = req.get_json()
    except ValueError:
        err = traceback.format_exc()
        logging.error("Invalid JSON body:\n%s", err)
        return func.HttpResponse(
            f"Invalid JSON body. Details:\n{err}",
            status_code=400
        )

    try:
        # Determine the field to index
        index_field = req_body.get('index_field')
        if not index_field:
            raise ValueError("Missing 'index_field' parameter.")

        index_name = req_body.get('index_name', f"Index_{index_field}")

        # Extract items array
        if isinstance(req_body, dict) and 'json' in req_body:
            items = req_body['json']
        elif isinstance(req_body, list):
            items = req_body
        else:
            raise TypeError("JSON body must be an array or an object with 'json' property.")

        if not isinstance(items, list):
            raise TypeError("'json' property must be an array.")

        # Build indexed result
        counters = {}
        result = []
        for item in items:
            key = item.get(index_field)
            count = counters.get(key, 0) + 1
            counters[key] = count

            new_item = dict(item)
            new_item[index_name] = count
            result.append(new_item)

        # Return updated JSON
        return func.HttpResponse(
            json.dumps(result),
            mimetype="application/json"
        )

    except Exception:
        err = traceback.format_exc()
        logging.error("Error processing request:\n%s", err)
        return func.HttpResponse(
            f"Error processing request. Details:\n{err}",
            status_code=500
        )








@app.route(route="json_to_excel_table")
def json_to_excel_table(req: func.HttpRequest) -> func.HttpResponse:
    try:
        # Get request body
        req_body = req.get_json()

        # Validate JSON input
        if 'json' not in req_body:
            return func.HttpResponse(
                json.dumps({"error": "Please provide a 'json' array in the request body"}),
                status_code=400,
                mimetype="application/json"
            )

        # Define maximum length constraints for Excel
        MAX_SHEET_NAME_LENGTH = 31
        MAX_TABLE_NAME_LENGTH = 255

        # Get JSON data and optional parameters
        json_data = req_body.get('json')
        sheet_name = req_body.get('sheet_name', 'Sheet1')[:MAX_SHEET_NAME_LENGTH]
        table_name = req_body.get('table_name', 'Table1')[:MAX_TABLE_NAME_LENGTH]

        # Convert JSON to DataFrame
        df = DataFrame(json_data)

        # Check if an existing file was provided
        file_content = req_body.get('file_content')
        if file_content:
            # Decode the provided Base64 Excel content and load the workbook
            excel_base64_str = file_content.get('$content')
            excel_bytes = base64.b64decode(excel_base64_str)
            wb = load_workbook(BytesIO(excel_bytes))
        else:
            # Create a new workbook and rename the default sheet
            wb = Workbook()
            ws = wb.active
            ws.title = sheet_name  # Rename the default sheet

        # Search for the table in all worksheets
        table_found = False
        target_worksheet = None
        target_table = None

        for ws in wb.worksheets:
            # Check if the table name exists in this worksheet's tables
            if table_name in ws._tables:
                table_found = True
                target_worksheet = ws
                target_table = ws._tables[table_name]
                break

        if table_found:
            # Table exists, append data to it
            ws = target_worksheet

            # Parse table range to find the location and dimensions
            table_ref = target_table.ref
            range_parts = table_ref.split(':')
            start_cell = range_parts[0]
            end_cell = range_parts[1]

            # Get column coordinates
            start_col = ''.join(filter(str.isalpha, start_cell))
            start_row = int(''.join(filter(str.isdigit, start_cell)))
            end_col = ''.join(filter(str.isalpha, end_cell))
            end_row = int(''.join(filter(str.isdigit, end_cell)))

            # Add new rows below the existing table
            new_row_idx = end_row + 1

            # Get headers from the existing table (header row)
            headers = []
            for col in range(column_index_from_string(start_col), column_index_from_string(end_col) + 1):
                cell_value = ws.cell(row=start_row, column=col).value
                headers.append(cell_value)

            # Check if the table has a blank first row
            has_blank_first_row = True
            first_data_row = start_row + 1
            if first_data_row == end_row:  # Make sure table has exactly one initial empty data row
                for col in range(column_index_from_string(start_col), column_index_from_string(end_col) + 1):
                    if ws.cell(row=first_data_row, column=col).value is not None:
                        has_blank_first_row = False
                        break
            else:
                has_blank_first_row = False  # No data rows at all

            # Map JSON keys to column indexes based on headers
            column_mapping = {}
            for idx, header in enumerate(headers):
                if header in df.columns:
                    column_mapping[header] = column_index_from_string(start_col) + idx

            rows_to_add = len(df)
            if rows_to_add > 0:
                # Insert new rows in the worksheet
                ws.insert_rows(end_row + 1, amount=rows_to_add)

                # Update references for all tables that might be affected by the insertion
                update_table_references(ws, new_row_idx, rows_to_add)

                # Add all new rows with data
                for i, (_, row_data) in enumerate(df.iterrows()):
                    for header, col_idx in column_mapping.items():
                        if header in row_data:
                            # Ensure boolean values are properly handled
                            value = row_data[header]
                            ws.cell(row=new_row_idx + i, column=col_idx, value=value)

                # Update table range to include the new rows
                new_table_ref = f"{start_cell}:{end_col}{end_row + rows_to_add}"
                target_table.ref = new_table_ref

                # If there was an empty first row and we added data, delete the empty first row
                if has_blank_first_row:
                    # Delete the first row of the table (which is empty)
                    ws.delete_rows(first_data_row, amount=1)
                    
                    # Update the table reference to reflect the deletion
                    new_table_start_row = start_row  # Header row stays the same
                    new_table_end_row = end_row + rows_to_add - 1  # End row is now one less due to deletion
                    new_table_ref = f"{start_col}{new_table_start_row}:{end_col}{new_table_end_row}"
                    target_table.ref = new_table_ref
                    
                    # Update references for all other tables that might be affected by the deletion
                    update_table_references_after_deletion(ws, first_data_row, 1)

        else:
            # Table doesn't exist, create new sheet and table
            if file_content:
                # If workbook exists but sheet doesn't, create a new sheet
                if sheet_name in wb.sheetnames:
                    i = 1
                    new_sheet_name = f"{sheet_name}_{i}"
                    while new_sheet_name in wb.sheetnames:
                        i += 1
                        new_sheet_name = f"{sheet_name}_{i}"
                    sheet_name = new_sheet_name

                wb.create_sheet(title=sheet_name)

            # Get the appropriate worksheet
            ws = wb[sheet_name]

            # Write DataFrame to the sheet
            for idx, col_name in enumerate(df.columns, start=1):
                ws.cell(row=1, column=idx, value=col_name)  # Write headers
            for r_idx, row in enumerate(df.itertuples(index=False), start=2):
                for c_idx, value in enumerate(row, start=1):
                    ws.cell(row=r_idx, column=c_idx, value=value)

            # Define table range
            table_range = f"A1:{get_column_letter(len(df.columns))}{len(df) + 1}"

            # Create a formatted table
            tab = ExcelTable(displayName=table_name, ref=table_range)
            style = TableStyleInfo(
                name="TableStyleMedium9", 
                showFirstColumn=False,
                showLastColumn=False, 
                showRowStripes=True, 
                showColumnStripes=False
            )
            tab.tableStyleInfo = style
            ws.add_table(tab)

        # Save the updated workbook to an in-memory file
        output = BytesIO()
        wb.save(output)
        output.seek(0)
        excel_data = output.getvalue()
        excel_base64 = base64.b64encode(excel_data).decode('utf-8')

        # Build and return the response
        response = {
            "$content-type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "$content": excel_base64
        }
        return func.HttpResponse(
            json.dumps(response),
            mimetype="application/json"
        )

    except Exception as e:
        # Log detailed traceback for debugging
        error_traceback = traceback.format_exc()
        logging.error(f"Exception: {str(e)}\nTraceback: {error_traceback}")

        return func.HttpResponse(
            json.dumps({
                "error": str(e),
                "traceback": error_traceback
            }),
            status_code=500,
            mimetype="application/json"
        )

def update_table_references(worksheet, insertion_row, rows_added):
    for table_name, table_ref in worksheet._tables.items():
        range_parts = table_ref.split(':')
        start_cell = range_parts[0]
        end_cell = range_parts[1]

        start_col = ''.join(filter(str.isalpha, start_cell))
        start_row = int(''.join(filter(str.isdigit, start_cell)))
        end_col = ''.join(filter(str.isalpha, end_cell))
        end_row = int(''.join(filter(str.isdigit, end_cell)))

        # If the table starts at or below our insertion point, adjust its reference
        if start_row >= insertion_row:
            new_start_row = start_row + rows_added
            new_end_row = end_row + rows_added
            new_ref = f"{start_col}{new_start_row}:{end_col}{new_end_row}"
            # Update the table reference directly in the worksheet's table collection
            worksheet._tables[table_name].ref = new_ref

def update_table_references_after_deletion(worksheet, deletion_row, rows_deleted):
    for table_name, table_ref in worksheet._tables.items():
        range_parts = table_ref.split(':')
        start_cell = range_parts[0]
        end_cell = range_parts[1]

        start_col = ''.join(filter(str.isalpha, start_cell))
        start_row = int(''.join(filter(str.isdigit, start_cell)))
        end_col = ''.join(filter(str.isalpha, end_cell))
        end_row = int(''.join(filter(str.isdigit, end_cell)))

        # For tables that start after the deletion point
        if start_row > deletion_row:
            new_start_row = start_row - rows_deleted
            new_end_row = end_row - rows_deleted
            new_ref = f"{start_col}{new_start_row}:{end_col}{new_end_row}"
            worksheet._tables[table_name].ref = new_ref












@app.route(route="excel_to_csvs")
def excel_to_csvs(req: func.HttpRequest) -> func.HttpResponse:
    """
    Azure Function to convert Excel file to array of CSVs (one per worksheet)
    """    
    try:
        # Parse the request body
        req_body = req.get_json()
        
        if not req_body or "file_content" not in req_body:
            return func.HttpResponse(
                json.dumps({"error": "Request body with file_content parameter is required"}),
                status_code=400,
                mimetype="application/json"
            )
        
        content_block = req_body["file_content"]
        
        # Validate required fields
        if "$content" not in content_block:
            return func.HttpResponse(
                json.dumps({"error": "file_content parameter with $content field is required"}),
                status_code=400,
                mimetype="application/json"
            )
        
        # Decode base64 content
        try:
            excel_data = base64.b64decode(content_block["$content"])
        except Exception as e:
            logging.error(f"Failed to decode base64 content: {str(e)}")
            return func.HttpResponse(
                json.dumps({
                    "error": "Failed to decode base64 content",
                    "details": str(e),
                    "traceback": traceback.format_exc()
                }),
                status_code=400,
                mimetype="application/json"
            )
        
        # Load Excel workbook from bytes
        try:
            workbook = load_workbook(BytesIO(excel_data), data_only=True)
        except Exception as e:
            logging.error(f"Failed to load Excel workbook: {str(e)}")
            return func.HttpResponse(
                json.dumps({
                    "error": "Failed to load Excel workbook",
                    "details": str(e),
                    "traceback": traceback.format_exc()
                }),
                status_code=400,
                mimetype="application/json"
            )
        
        # Process each worksheet
        result = []
        
        for sheet_name in workbook.sheetnames:
            try:
                worksheet = workbook[sheet_name]
                csv_data = convert_sheet_to_csv(worksheet)
                
                result.append({
                    "sheet_name": sheet_name,
                    "csv": csv_data
                })
                
                logging.info(f"Successfully processed sheet: {sheet_name}")
                
            except Exception as e:
                logging.error(f"Failed to process sheet '{sheet_name}': {str(e)}")
                # Continue processing other sheets, but log the error
                result.append({
                    "sheet_name": sheet_name,
                    "csv": "",
                    "error": f"Failed to process sheet: {str(e)}"
                })
        
        # Close workbook to free resources
        workbook.close()
        
        return func.HttpResponse(
            json.dumps(result),
            status_code=200,
            mimetype="application/json"
        )
    
    except Exception as e:
        logging.error(f"Unexpected error in main function: {str(e)}")
        logging.error(traceback.format_exc())
        
        return func.HttpResponse(
            json.dumps({
                "error": "Unexpected error occurred",
                "details": str(e),
                "traceback": traceback.format_exc()
            }),
            status_code=500,
            mimetype="application/json"
        )


def convert_sheet_to_csv(worksheet) -> str:
    """
    Convert an Excel worksheet to CSV format with proper handling of commas and quotes
    
    Args:
        worksheet: openpyxl worksheet object
        
    Returns:
        str: CSV formatted string
    """
    
    # Get all rows with data (openpyxl automatically handles used range)
    rows_data = []
    
    # Find the actual used range by getting all rows and columns with data
    max_row = worksheet.max_row
    max_col = worksheet.max_column
    
    # If worksheet is empty, return empty string
    if max_row == 1 and max_col == 1:
        cell_value = worksheet.cell(1, 1).value
        if cell_value is None or cell_value == "":
            return ""
    
    # Extract data from each row
    for row_num in range(1, max_row + 1):
        row_data = []
        for col_num in range(1, max_col + 1):
            cell_value = worksheet.cell(row_num, col_num).value
            
            # Convert None to empty string, everything else to string
            if cell_value is None:
                cell_str = ""
            else:
                cell_str = str(cell_value)
            
            row_data.append(cell_str)
        
        rows_data.append(row_data)
    
    # Remove trailing empty rows
    while rows_data and all(cell == "" for cell in rows_data[-1]):
        rows_data.pop()
    
    # If no data after cleanup, return empty string
    if not rows_data:
        return ""
    
    # Convert to CSV using Python's csv module for proper escaping
    output = StringIO()
    csv_writer = csv.writer(output, quoting=csv.QUOTE_MINIMAL, lineterminator='\r\n')
    
    for row in rows_data:
        # Remove trailing empty cells from each row
        while row and row[-1] == "":
            row.pop()
        csv_writer.writerow(row)
    
    csv_content = output.getvalue()
    output.close()
    
    return csv_content










@app.route(route="get_or_edit_excel_sheet_names")
def get_or_edit_excel_sheet_names(req: func.HttpRequest) -> func.HttpResponse:
    try:
        MAX_SHEET_NAME_LENGTH = 31
        req_body = req.get_json()

        # Extract parameters from request
        file_content = req_body.get('file_content')
        index_edit = req_body.get('index_edit')
        name_edit = req_body.get('name_edit')
        regex_pattern = req_body.get('regex')
        replace_text = req_body.get('replace_text')

        if not file_content:
            return func.HttpResponse(
                json.dumps({"error": "file_content is required"}),
                status_code=400,
                mimetype="application/json"
            )

        # Validate and extract the base64 content
        if "$content-type" not in file_content or file_content["$content-type"] != "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            return func.HttpResponse(
                json.dumps({"error": "Invalid file content type"}),
                status_code=400,
                mimetype="application/json"
            )

        if "$content" not in file_content:
            return func.HttpResponse(
                json.dumps({"error": "Missing file content"}),
                status_code=400,
                mimetype="application/json"
            )

        # Process the Excel file
        excel_base64 = file_content["$content"]
        excel_data = base64.b64decode(excel_base64)
        wb = load_workbook(BytesIO(excel_data))

        # Get all sheet names in their index order
        original_sheet_names = wb.sheetnames
        sheet_names = original_sheet_names.copy()
        modified = False

        # Process index_edit if provided
        if index_edit:
            for index_str, new_name in index_edit.items():
                try:
                    index = int(index_str) - 1
                    if 0 <= index < len(sheet_names):
                        old_name = sheet_names[index]
                        wb[old_name].title = new_name[:MAX_SHEET_NAME_LENGTH]
                        sheet_names[index] = new_name[:MAX_SHEET_NAME_LENGTH]
                        modified = True
                    else:
                        logging.warning(f"Index {index} out of range. Skipping.")
                except ValueError:
                    logging.warning(f"Invalid index: {index_str}. Skipping.")

        # Process name_edit if provided
        if name_edit:
            for old_name, new_name in name_edit.items():
                if old_name in wb.sheetnames:
                    wb[old_name].title = new_name[:MAX_SHEET_NAME_LENGTH]
                    # Update the sheet_names list
                    for i, name in enumerate(sheet_names):
                        if name == old_name:
                            sheet_names[i] = new_name[:MAX_SHEET_NAME_LENGTH]
                    modified = True
                else:
                    logging.warning(f"Sheet name '{old_name}' not found. Skipping.")

        # Process regex replacement if both regex and replace_text are provided
        if regex_pattern and replace_text is not None:
            pattern = re.compile(regex_pattern)
            for i, sheet_name in enumerate(sheet_names):
                if pattern.search(sheet_name):
                    new_name = pattern.sub(replace_text, sheet_name)
                    old_sheet = wb[sheet_name]
                    old_sheet.title = new_name[:MAX_SHEET_NAME_LENGTH]
                    sheet_names[i] = new_name[:MAX_SHEET_NAME_LENGTH]
                    modified = True

        # If changes were made, save the file and prepare it for return
        if modified:
            output = BytesIO()
            wb.save(output)
            output.seek(0)

            # Get updated sheet names (may be different from our sheet_names list due to Excel's constraints)
            final_sheet_names = wb.sheetnames

            # Encode the modified file to base64
            modified_excel_base64 = base64.b64encode(output.getvalue()).decode('utf-8')

            return func.HttpResponse(
                json.dumps({
                    "sheet_names": final_sheet_names,
                    "file_content": {
                        "$content-type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        "$content": modified_excel_base64
                    }
                }),
                mimetype="application/json"
            )
        else:
            # If no changes, just return the original sheet names
            return func.HttpResponse(
                json.dumps({"sheet_names": original_sheet_names}),
                mimetype="application/json"
            )

    except Exception as e:
        # Catch any exceptions and return error details
        error_traceback = traceback.format_exc()
        logging.error(f"Error processing request: {str(e)}\n{error_traceback}")
        return func.HttpResponse(
            json.dumps({
                "error": str(e),
                "traceback": error_traceback
            }),
            status_code=500,
            mimetype="application/json"
        )








@app.route(route="append_or_remove_excel_sheets")
def append_or_remove_excel_sheets(req: func.HttpRequest) -> func.HttpResponse:
    try:
        req_body = req.get_json()
        file_content = req_body.get('file_content', [])
        select_sheets = req_body.get('select')
        remove_sheets = req_body.get('remove')
        select_sheets_list = select_sheets if select_sheets else None
        remove_sheets_list = remove_sheets if remove_sheets else None

        if not isinstance(file_content, list) or len(file_content) == 0:
            return func.HttpResponse(json.dumps({"error": "Request must contain an array of Excel file content objects"}), mimetype="application/json", status_code=400)

        # Check if there's only one input workbook
        if len(file_content) == 1 and (select_sheets_list or remove_sheets_list):
            # Use simplified logic for single workbook sheet selection/removal
            return process_single_workbook(file_content[0], select_sheets_list, remove_sheets_list)

        ET.register_namespace('', "http://schemas.openxmlformats.org/spreadsheetml/2006/main")
        ET.register_namespace('r', "http://schemas.openxmlformats.org/officeDocument/2006/relationships")
        ET.register_namespace('ct', "http://schemas.openxmlformats.org/package/2006/content-types")
        ET.register_namespace('mc', "http://schemas.openxmlformats.org/markup-compatibility/2006")
        ET.register_namespace('x14ac', "http://schemas.microsoft.com/office/spreadsheetml/2009/9/ac") 
        ET.register_namespace('xr', "http://schemas.microsoft.com/office/spreadsheetml/2014/revision")
        ET.register_namespace('xr2', "http://schemas.microsoft.com/office/spreadsheetml/2015/revision2")
        ET.register_namespace('xr3', "http://schemas.microsoft.com/office/spreadsheetml/2016/revision3")
        ET.register_namespace('x14', "http://schemas.microsoft.com/office/spreadsheetml/2009/9/main")
        ET.register_namespace('xpdl', "http://schemas.microsoft.com/office/spreadsheetml/2016/pivotdefaultlayout")
        ET.register_namespace('a', "http://schemas.openxmlformats.org/drawingml/2006/main")
        ET.register_namespace('c', "http://schemas.openxmlformats.org/drawingml/2006/chart")
        ET.register_namespace('xdr', "http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing")
        ET.register_namespace('a16', "http://schemas.microsoft.com/office/drawing/2014/main")
        ET.register_namespace('c14', "http://schemas.microsoft.com/office/drawing/2007/8/2/chart")
        ET.register_namespace('c16', "http://schemas.microsoft.com/office/drawing/2014/chart")
        ET.register_namespace('c16r2', "http://schemas.microsoft.com/office/drawing/2015/06/chart")

        sheet_name_counts, count_error = count_sheet_names(file_content)
        if count_error:
            return func.HttpResponse(json.dumps({"error": count_error}), mimetype="application/json", status_code=400)

        sheets_to_process = build_sheets_to_process(file_content, sheet_name_counts, select_sheets_list, remove_sheets_list)

        if len(sheets_to_process) == 0:
            return func.HttpResponse(json.dumps({"error": "No sheets were included in the output file based on select/remove criteria"}), mimetype="application/json", status_code=400)

        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = os.path.join(temp_dir, "output.xlsx")

            workbook_xml_root = create_empty_workbook_xml()
            ns = {'ns': 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'}
            sheets_element = workbook_xml_root.find('.//ns:sheets', ns)

            workbook_rels_root = create_empty_workbook_rels()

            with ZipFile(output_path, 'w') as output_zip:
                next_sheet_rel_id = 1000

                add_relationship(workbook_rels_root, "rId1", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/sharedStrings", "sharedStrings.xml")
                add_relationship(workbook_rels_root, "rId2", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles", "styles.xml")
                add_relationship(workbook_rels_root, "rId3", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme", "theme/theme1.xml")

                shared_strings_map, pivot_cache_map, pivot_table_map, table_map, final_pivot_names, final_table_names, pivot_table_cacheids, dxf_map, chart_map, drawing_map, cellXfs_map, name_map = copy_shared_resources(file_content, output_zip, temp_dir, len(sheets_to_process), sheets_to_process)

                pivot_rel_id_start = len(workbook_rels_root.findall('.//{*}Relationship')) + 1
                next_pivot_rel_id = pivot_rel_id_start

                for (wb_idx, old_id), new_id in pivot_cache_map.items():
                    pivot_rel_id = f"rId{next_pivot_rel_id}"
                    add_relationship(workbook_rels_root, pivot_rel_id, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/pivotCacheDefinition", f"pivotCache/pivotCacheDefinition{new_id}.xml")
                    next_pivot_rel_id += 1

                source_files = []
                for idx, file_content_obj in enumerate(file_content):
                    source_path = os.path.join(temp_dir, f"source_{idx}.xlsx")
                    if not os.path.exists(source_path):
                        with open(source_path, 'wb') as f:
                            f.write(base64.b64decode(file_content_obj.get("$content")))
                    source_files.append(source_path)

                for index, sheet_info in enumerate(sheets_to_process):
                    source_idx = sheet_info['source_idx']
                    source_path = source_files[source_idx]

                    with ZipFile(source_path, 'r') as source_zip:
                        sheet_rel_id, next_sheet_rel_id = process_sheet(sheet_info, source_zip, output_zip, index + 1, next_sheet_rel_id, workbook_rels_root, sheets_element, temp_dir, shared_strings_map, pivot_cache_map, pivot_table_map,
                                                                        table_map, final_pivot_names, final_table_names, pivot_table_cacheids, dxf_map, chart_map, drawing_map, cellXfs_map, name_map)

                # Only create pivotCaches element if there are pivot tables
                if pivot_table_cacheids:
                    main_ns = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
                    pivot_caches_element = ET.SubElement(workbook_xml_root, '{' + main_ns + '}pivotCaches')
                    current_pivot_rel_id = pivot_rel_id_start

                    for (wb_idx, pivot_id), original_cache_id in pivot_table_cacheids.items():
                        # Create a new pivotCache element with the original cacheId value.
                        pivot_cache = ET.SubElement(pivot_caches_element, '{' + main_ns + '}pivotCache')
                        pivot_cache.set('cacheId', str(original_cache_id))
                        pivot_cache.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id', f'rId{current_pivot_rel_id}')
                        current_pivot_rel_id += 1

                workbook_xml = ET.tostring(workbook_xml_root, encoding='UTF-8')
                output_zip.writestr('xl/workbook.xml', workbook_xml)

                workbook_rels_xml = ET.tostring(workbook_rels_root, encoding='UTF-8')
                output_zip.writestr('xl/_rels/workbook.xml.rels', workbook_rels_xml)

            with open(output_path, 'rb') as f:
                output_data = f.read()

        output_base64 = base64.b64encode(output_data).decode('utf-8')

        response = {"$content-type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "$content": output_base64}

        return func.HttpResponse(json.dumps(response), mimetype="application/json")

    except Exception as e:
        error_response = {"error": str(e), "traceback": traceback.format_exc()}
        return func.HttpResponse(json.dumps(error_response), mimetype="application/json", status_code=500)

def process_single_workbook(file_content_obj, select_sheets_list, remove_sheets_list):
    """
    Process a single workbook to select or remove specific sheets without 
    using the more complex merging logic.
    """
    try:
        content_type = file_content_obj.get("$content-type")
        content = file_content_obj.get("$content")

        if not content_type or not content or content_type != "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            return func.HttpResponse(json.dumps({"error": "Invalid content object format"}), mimetype="application/json", status_code=400)

        decoded_content = base64.b64decode(content)
        file_stream = BytesIO(decoded_content)

        # Load the workbook
        workbook = load_workbook(file_stream, data_only=False)

        # Get all sheet names
        all_sheets = workbook.sheetnames

        # Determine which sheets to keep
        sheets_to_keep = set(all_sheets)

        # Apply select filter if provided
        if select_sheets_list:
            sheets_to_keep = {sheet for sheet in all_sheets if sheet in select_sheets_list}

        # Apply remove filter if provided
        if remove_sheets_list:
            sheets_to_keep = {sheet for sheet in sheets_to_keep if sheet not in remove_sheets_list}

        if not sheets_to_keep:
            return func.HttpResponse(json.dumps({"error": "No sheets were included in the output file based on select/remove criteria"}), mimetype="application/json", status_code=400)

        # Create a list of sheets to remove
        sheets_to_remove = [sheet for sheet in all_sheets if sheet not in sheets_to_keep]

        # Remove sheets from the workbook
        for sheet_name in sheets_to_remove:
            sheet = workbook[sheet_name]
            workbook.remove(sheet)

        # Save the modified workbook to a BytesIO object
        output_stream = BytesIO()
        workbook.save(output_stream)
        output_stream.seek(0)

        # Convert to base64
        output_data = output_stream.getvalue()
        output_base64 = base64.b64encode(output_data).decode('utf-8')

        response = {"$content-type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "$content": output_base64}

        return func.HttpResponse(json.dumps(response), mimetype="application/json")

    except Exception as e:
        error_response = {"error": str(e), "traceback": traceback.format_exc()}
        return func.HttpResponse(json.dumps(error_response), mimetype="application/json", status_code=500)

def create_empty_workbook_xml():
    main_ns = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
    rel_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

    workbook = ET.Element('{' + main_ns + '}workbook')

    ET.register_namespace('', main_ns)
    ET.register_namespace('r', rel_ns)

    file_version = ET.SubElement(workbook, '{' + main_ns + '}fileVersion')
    file_version.set('appName', 'xl')
    file_version.set('lastEdited', '1')
    file_version.set('lowestEdited', '1')
    file_version.set('rupBuild', '14420')

    workbook_pr = ET.SubElement(workbook, '{' + main_ns + '}workbookPr')
    workbook_pr.set('defaultThemeVersion', '153222')

    book_views = ET.SubElement(workbook, '{' + main_ns + '}bookViews')
    workbook_view = ET.SubElement(book_views, '{' + main_ns + '}workbookView')
    workbook_view.set('xWindow', '0')
    workbook_view.set('yWindow', '0')
    workbook_view.set('windowWidth', '20490')
    workbook_view.set('windowHeight', '7755')

    sheets = ET.SubElement(workbook, '{' + main_ns + '}sheets')

    return workbook

def create_empty_workbook_rels():
    ns = 'http://schemas.openxmlformats.org/package/2006/relationships'
    rels = ET.Element('{'+ns+'}Relationships')
    rels.set('xmlns', ns)
    return rels

def add_relationship(root, rel_id, rel_type, target):
    ns = 'http://schemas.openxmlformats.org/package/2006/relationships'
    rel = ET.SubElement(root, '{'+ns+'}Relationship')
    rel.set('Id', rel_id)
    rel.set('Type', rel_type)
    rel.set('Target', target)
    return rel

def count_sheet_names(file_content):
    sheet_name_counts = {}

    for idx, file_content_obj in enumerate(file_content):
        content_type = file_content_obj.get("$content-type")
        content = file_content_obj.get("$content")

        if not content_type or not content or content_type != "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            return None, f"Invalid content object format at index {idx}"

        try:
            decoded_content = base64.b64decode(content)
            file_stream = BytesIO(decoded_content)

            workbook = load_workbook(file_stream, data_only=False)

            for sheet_name in workbook.sheetnames:
                if sheet_name in sheet_name_counts:
                    sheet_name_counts[sheet_name] += 1
                else:
                    sheet_name_counts[sheet_name] = 1

        except Exception as e:
            return None, f"Error processing Excel file at index {idx}: {str(e)}"

    return sheet_name_counts, None

def build_sheets_to_process(file_content, sheet_name_counts, select_sheets_list, remove_sheets_list):
    indexed_sheet_names = {}
    current_counts = {name: 0 for name in sheet_name_counts}
    sheets_to_process = []

    # Create a mapping of potential indexed names to their base sheet names
    # This will help us match indexed names to non-indexed sheets
    potential_indexed_names = {}
    for sheet_name in sheet_name_counts:
        for i in range(1, sheet_name_counts[sheet_name] + 1):
            indexed_name = f"{i}_{sheet_name}"
            potential_indexed_names[indexed_name] = sheet_name

    for idx, file_content_obj in enumerate(file_content):
        content = file_content_obj.get("$content")

        decoded_content = base64.b64decode(content)
        file_stream = BytesIO(decoded_content)

        workbook = load_workbook(file_stream, data_only=False)

        for sheet_name in workbook.sheetnames:
            current_counts[sheet_name] += 1

            if sheet_name_counts[sheet_name] > 1:
                new_sheet_name = f"{current_counts[sheet_name]}_{sheet_name}"
            else:
                new_sheet_name = sheet_name

            indexed_sheet_names[new_sheet_name] = sheet_name

            should_include = True

            if select_sheets_list:
                # Check if the sheet is in the select list, either by its original name,
                # new indexed name, or potential indexed format (like 1_SheetName even if only one exists)
                sheet_in_list = (sheet_name in select_sheets_list or new_sheet_name in select_sheets_list)

                # Check if any potential indexed version of this sheet is in the select list
                if not sheet_in_list and sheet_name_counts[sheet_name] == 1:
                    potential_indexed_name = f"1_{sheet_name}"
                    sheet_in_list = potential_indexed_name in select_sheets_list

                should_include = sheet_in_list

            if remove_sheets_list and should_include:
                # Similar logic for removal - check all possible ways this sheet could be referenced
                remove_sheet = (sheet_name in remove_sheets_list or new_sheet_name in remove_sheets_list)

                # Check if any potential indexed version of this sheet is in the remove list
                if not remove_sheet and sheet_name_counts[sheet_name] == 1:
                    potential_indexed_name = f"1_{sheet_name}"
                    remove_sheet = potential_indexed_name in remove_sheets_list

                should_include = not remove_sheet

            if should_include:
                sheets_to_process.append({'workbook': workbook, 'sheet_name': sheet_name, 'new_sheet_name': new_sheet_name, 'file_stream': file_stream, 'source_idx': idx})

    return sheets_to_process

def update_sheet_shared_strings(sheet_tree, source_wb_idx, shared_strings_map):
    ns = get_namespace(sheet_tree)
    if ns:
        ns_dict = {'ns': ns}

        for cell in sheet_tree.findall('.//ns:c[@t="s"]', ns_dict):
            v_elem = cell.find('./ns:v', ns_dict)
            if v_elem is not None and v_elem.text:
                old_idx = int(v_elem.text)

                if (source_wb_idx, old_idx) in shared_strings_map:
                    v_elem.text = str(shared_strings_map[(source_wb_idx, old_idx)])

def update_sheet_dxf_references(sheet_tree, source_wb_idx, dxf_map):
    """Update dataDxfId references in conditional formatting rules"""
    ns = get_namespace(sheet_tree)
    if ns:
        ns_dict = {'ns': ns}

        for cf_rule in sheet_tree.findall('.//ns:conditionalFormatting/ns:cfRule[@dxfId]', ns_dict):
            if 'dxfId' in cf_rule.attrib:
                old_dxf_id = int(cf_rule.get('dxfId'))
                if (source_wb_idx, old_dxf_id) in dxf_map:
                    cf_rule.set('dxfId', str(dxf_map[(source_wb_idx, old_dxf_id)]))

        for cf_rule in sheet_tree.findall('.//ns:conditionalFormatting/ns:cfRule[@dataDxfId]', ns_dict):
            if 'dataDxfId' in cf_rule.attrib:
                old_dxf_id = int(cf_rule.get('dataDxfId'))
                if (source_wb_idx, old_dxf_id) in dxf_map:
                    cf_rule.set('dataDxfId', str(dxf_map[(source_wb_idx, old_dxf_id)]))

        for pivot_table in sheet_tree.findall('.//ns:pivotTableStyleInfo', ns_dict):
            if 'dxfId' in pivot_table.attrib:
                old_dxf_id = int(pivot_table.get('dxfId'))
                if (source_wb_idx, old_dxf_id) in dxf_map:
                    pivot_table.set('dxfId', str(dxf_map[(source_wb_idx, old_dxf_id)]))

def copy_shared_resources(file_content, output_zip, temp_dir, num_sheets, sheets_to_process):
    pivot_cache_map = {}
    pivot_table_map = {}
    table_map = {}
    name_map = {}
    pivot_table_name_map = {}
    all_shared_strings = []
    merged_shared_strings_map = {}

    next_pivot_cache_id = 1
    next_pivot_table_id = 1
    next_table_id = 1

    # Add these new maps
    chart_map = {}
    drawing_map = {}
    next_chart_id = 1
    next_drawing_id = 1

    # Build mapping from (source_idx, original_sheet_name) to new_sheet_name
    sheet_name_mapping = {}
    for sheet_info in sheets_to_process:
        key = (sheet_info['source_idx'], sheet_info['sheet_name'])
        sheet_name_mapping[key] = sheet_info['new_sheet_name']

    source_files = []
    for idx, file_content_obj in enumerate(file_content):
        source_path = os.path.join(temp_dir, f"source_{idx}.xlsx")
        with open(source_path, 'wb') as f:
            f.write(base64.b64decode(file_content_obj.get("$content")))
        source_files.append(source_path)

    added_files = set()

    for idx, source_path in enumerate(source_files):
        with ZipFile(source_path, 'r') as source_zip:
            if 'xl/sharedStrings.xml' in source_zip.namelist():
                shared_strings_xml = source_zip.read('xl/sharedStrings.xml')
                shared_strings_root = ET.fromstring(shared_strings_xml)

                ns = get_namespace(shared_strings_root)
                ns_dict = {'ns': ns} if ns else {}

                for si_idx, si in enumerate(shared_strings_root.findall('.//ns:si', ns_dict) if ns else shared_strings_root.findall('.//si')):
                    si_str = ET.tostring(si, encoding='UTF-8').decode('UTF-8')

                    found = False
                    for existing_idx, existing_str in enumerate(all_shared_strings):
                        if si_str == existing_str:
                            merged_shared_strings_map[(idx, si_idx)] = existing_idx
                            found = True
                            break

                    if not found:
                        new_idx = len(all_shared_strings)
                        all_shared_strings.append(si_str)
                        merged_shared_strings_map[(idx, si_idx)] = new_idx

            for item in source_zip.namelist():
                if 'pivotCache/pivotCacheDefinition' in item:
                    cache_id = extract_number(item)
                    if cache_id is not None:
                        pivot_cache_map[(idx, cache_id)] = next_pivot_cache_id
                        next_pivot_cache_id += 1

            for item in source_zip.namelist():
                if 'pivotTables/pivotTable' in item:
                    table_id = extract_number(item)
                    if table_id is not None:
                        with source_zip.open(item) as pivot_file:
                            pivot_tree = ET.parse(pivot_file)
                            pivot_root = pivot_tree.getroot()
                            original_name = pivot_root.get('name', f'PivotTable{table_id}')
                            pivot_table_name_map[(idx, table_id)] = original_name

                        pivot_table_map[(idx, table_id)] = next_pivot_table_id
                        next_pivot_table_id += 1

            pivot_table_cacheids = {}
            for idx, source_path in enumerate(source_files):
                with ZipFile(source_path, 'r') as source_zip:
                    for item in source_zip.namelist():
                        if 'pivotTables/pivotTable' in item:
                            table_id = extract_number(item)
                            if table_id is not None:
                                with source_zip.open(item) as pivot_file:
                                    pivot_tree = ET.parse(pivot_file)
                                    pivot_root = pivot_tree.getroot()
                                    cache_id = pivot_root.get('cacheId')
                                    if cache_id:
                                        pivot_table_cacheids[(idx, table_id)] = int(cache_id)

            pivot_names = {}
            final_pivot_names = {}

            for (wb_idx, old_id), original_name in pivot_table_name_map.items():
                new_id = pivot_table_map.get((wb_idx, old_id))
                if new_id is None:
                    continue

                if original_name in pivot_names:
                    counter = 1
                    for existing_key in pivot_names[original_name]:
                        final_pivot_names[existing_key] = f"{counter}_{original_name}"
                        counter += 1
                    final_pivot_names[(wb_idx, old_id)] = f"{counter}_{original_name}"
                    pivot_names[original_name].append((wb_idx, old_id))
                else:
                    pivot_names[original_name] = [(wb_idx, old_id)]
                    final_pivot_names[(wb_idx, old_id)] = original_name

            table_name_map = {}
            for idx, source_path in enumerate(source_files):
                with ZipFile(source_path, 'r') as source_zip:
                    for item in source_zip.namelist():
                        if 'tables/table' in item:
                            table_id = extract_number(item)
                            if table_id is not None:
                                with source_zip.open(item) as table_file:
                                    table_tree = ET.parse(table_file)
                                    table_root = table_tree.getroot()
                                    original_name = table_root.get('name')
                                    display_name = table_root.get('displayName', original_name)
                                    table_name_map[(idx, table_id)] = original_name

            all_table_names = {}
            for (wb_idx, table_id), original_name in table_name_map.items():
                if original_name in all_table_names:
                    all_table_names[original_name].append((wb_idx, table_id))
                else:
                    all_table_names[original_name] = [(wb_idx, table_id)]

            final_table_names = {}

            next_table_id = 1
            for table_name, entries in all_table_names.items():
                if len(entries) == 1:
                    wb_idx, table_id = entries[0]
                    final_table_names[(wb_idx, table_id)] = table_name
                    table_map[(wb_idx, table_id)] = next_table_id
                    next_table_id += 1
                else:
                    for i, (wb_idx, table_id) in enumerate(entries, 1):
                        prefixed_name = f"{table_name}_{i}"
                        final_table_names[(wb_idx, table_id)] = prefixed_name
                        table_map[(wb_idx, table_id)] = next_table_id
                        next_table_id += 1

            name_map = {}
            for (wb_idx, table_id), table_name in final_table_names.items():
                original_name = table_name_map.get((wb_idx, table_id))
                if original_name:
                    name_map[(wb_idx, original_name)] = table_name

            ### Scan and map charts and drawings in each file
            for idx, source_path in enumerate(source_files):
                with ZipFile(source_path, 'r') as source_zip:
                    # Map chart files
                    for item in source_zip.namelist():
                        if 'xl/charts/chart' in item and '.xml' in item:
                            chart_id = extract_number(item)
                            if chart_id is not None:
                                chart_map[(idx, chart_id)] = next_chart_id
                                next_chart_id += 1

                    # Map drawing files
                    for item in source_zip.namelist():
                        if 'xl/drawings/drawing' in item and '.xml' in item:
                            drawing_id = extract_number(item)
                            if drawing_id is not None:
                                drawing_map[(idx, drawing_id)] = next_drawing_id
                                next_drawing_id += 1
            ###

    if all_shared_strings:
        shared_strings_xml = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        shared_strings_xml += f'<sst xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" count="{len(all_shared_strings)}" uniqueCount="{len(all_shared_strings)}">\n'

        for string_xml in all_shared_strings:
            shared_strings_xml += string_xml

        shared_strings_xml += '</sst>'
        output_zip.writestr('xl/sharedStrings.xml', shared_strings_xml)
        added_files.add('xl/sharedStrings.xml')
    '''
    with ZipFile(source_files[0], 'r') as first_zip:
        if 'xl/styles.xml' in first_zip.namelist() and 'xl/styles.xml' not in added_files:
            styles_xml = first_zip.read('xl/styles.xml')
            output_zip.writestr('xl/styles.xml', styles_xml)
            added_files.add('xl/styles.xml')
    '''
    # Merge styles from all workbooks
    all_dxfs = []
    all_fills = []
    all_fonts = []
    all_borders = []
    all_number_formats = []
    all_cell_styles = []
    all_cell_style_xfs = []
    all_cell_xfs = []
    dxf_map = {}  # Map workbook_idx,dxf_id -> new_dxf_id
    num_fmt_map = {}  # Map workbook_idx,numFmtId -> new_numFmtId
    fill_map = {}  # Map workbook_idx,fill_id -> new_fill_id
    font_map = {}  # Map workbook_idx,font_id -> new_font_id 
    border_map = {}  # Map workbook_idx,border_id -> new_border_id
    cellStyleXfs_map = {}  # Map workbook_idx,xf_id -> new_xf_id
    cellXfs_map = {}  # Map workbook_idx,xf_id -> new_xf_id

    # Start with numFmtId 164 for custom formats (0-163 are reserved)
    next_num_fmt_id = 164

    for idx, source_path in enumerate(source_files):
        with ZipFile(source_path, 'r') as source_zip:
            if 'xl/styles.xml' in source_zip.namelist():
                styles_xml = source_zip.read('xl/styles.xml')
                styles_root = ET.fromstring(styles_xml)

                # Process numFmts
                num_fmts_elem = styles_root.find('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}numFmts')
                if num_fmts_elem is not None:
                    for num_fmt in num_fmts_elem.findall('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}numFmt'):
                        fmt_id = int(num_fmt.get('numFmtId'))
                        fmt_code = num_fmt.get('formatCode')

                        # Skip built-in formats (0-163)
                        if fmt_id < 164:
                            continue

                        # Check if this format code already exists
                        existing_id = None
                        for i, existing_fmt in enumerate(all_number_formats):
                            if existing_fmt.get('formatCode') == fmt_code:
                                existing_id = int(existing_fmt.get('numFmtId'))
                                break

                        if existing_id is not None:
                            num_fmt_map[(idx, fmt_id)] = existing_id
                        else:
                            num_fmt_map[(idx, fmt_id)] = next_num_fmt_id
                            num_fmt.set('numFmtId', str(next_num_fmt_id))
                            all_number_formats.append(num_fmt)
                            next_num_fmt_id += 1

                # Process fills
                fills_elem = styles_root.find('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}fills')
                if fills_elem is not None:
                    for fill_idx, fill in enumerate(fills_elem.findall('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}fill')):
                        fill_str = ET.tostring(fill, encoding='UTF-8').decode('UTF-8')

                        # Check if this fill already exists
                        existing_idx = None
                        for i, existing_fill_str in enumerate(all_fills):
                            if fill_str == existing_fill_str:
                                existing_idx = i
                                break

                        if existing_idx is not None:
                            fill_map[(idx, fill_idx)] = existing_idx
                        else:
                            new_fill_idx = len(all_fills)
                            fill_map[(idx, fill_idx)] = new_fill_idx
                            all_fills.append(fill_str)

                # Process fonts
                fonts_elem = styles_root.find('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}fonts')
                if fonts_elem is not None:
                    for font_idx, font in enumerate(fonts_elem.findall('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}font')):
                        font_str = ET.tostring(font, encoding='UTF-8').decode('UTF-8')

                        # Check if this font already exists
                        existing_idx = None
                        for i, existing_font_str in enumerate(all_fonts):
                            if font_str == existing_font_str:
                                existing_idx = i
                                break

                        if existing_idx is not None:
                            font_map[(idx, font_idx)] = existing_idx
                        else:
                            new_font_idx = len(all_fonts)
                            font_map[(idx, font_idx)] = new_font_idx
                            all_fonts.append(font_str)

                # Process borders with better handling of border styles and colors
                borders_elem = styles_root.find('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}borders')
                if borders_elem is not None:
                    for border_idx, border in enumerate(borders_elem.findall('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}border')):
                        # Process each border side (left, right, top, bottom)
                        ns = 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'
                        for side in ['left', 'right', 'top', 'bottom', 'diagonal']:
                            side_elem = border.find(f'.//{{{ns}}}' + side)
                            if side_elem is not None and side_elem.get('style') is not None:
                                # This border has a style, ensure it's preserved
                                color_elem = side_elem.find(f'.//{{{ns}}}color')
                                if color_elem is not None and 'indexed' in color_elem.attrib:
                                    # Map indexed colors if necessary
                                    pass

                        border_str = ET.tostring(border, encoding='UTF-8').decode('UTF-8')

                        # Check if this border already exists
                        existing_idx = None
                        for i, existing_border_str in enumerate(all_borders):
                            if border_str == existing_border_str:
                                existing_idx = i
                                break

                        if existing_idx is not None:
                            border_map[(idx, border_idx)] = existing_idx
                        else:
                            new_border_idx = len(all_borders)
                            border_map[(idx, border_idx)] = new_border_idx
                            all_borders.append(border_str)

                # Process cellStyleXfs
                cellStyleXfs_elem = styles_root.find('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}cellStyleXfs')
                if cellStyleXfs_elem is not None:
                    for xf_idx, xf in enumerate(cellStyleXfs_elem.findall('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}xf')):
                        # Update references to other style elements
                        if 'numFmtId' in xf.attrib:
                            fmt_id = int(xf.get('numFmtId'))
                            if fmt_id >= 164 and (idx, fmt_id) in num_fmt_map:
                                xf.set('numFmtId', str(num_fmt_map[(idx, fmt_id)]))

                        if 'fontId' in xf.attrib:
                            font_id = int(xf.get('fontId'))
                            if (idx, font_id) in font_map:
                                xf.set('fontId', str(font_map[(idx, font_id)]))

                        if 'fillId' in xf.attrib:
                            fill_id = int(xf.get('fillId'))
                            if (idx, fill_id) in fill_map:
                                xf.set('fillId', str(fill_map[(idx, fill_id)]))

                        if 'borderId' in xf.attrib:
                            border_id = int(xf.get('borderId'))
                            if (idx, border_id) in border_map:
                                xf.set('borderId', str(border_map[(idx, border_id)]))

                        xf_str = ET.tostring(xf, encoding='UTF-8').decode('UTF-8')

                        # Check if this xf already exists
                        existing_idx = None
                        for i, existing_xf_str in enumerate(all_cell_style_xfs):
                            if xf_str == existing_xf_str:
                                existing_idx = i
                                break

                        if existing_idx is not None:
                            cellStyleXfs_map[(idx, xf_idx)] = existing_idx
                        else:
                            new_xf_idx = len(all_cell_style_xfs)
                            cellStyleXfs_map[(idx, xf_idx)] = new_xf_idx
                            all_cell_style_xfs.append(xf_str)

                # Process cellXfs
                cellXfs_elem = styles_root.find('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}cellXfs')
                if cellXfs_elem is not None:
                    for xf_idx, xf in enumerate(cellXfs_elem.findall('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}xf')):
                        # Update references to other style elements
                        if 'numFmtId' in xf.attrib:
                            fmt_id = int(xf.get('numFmtId'))
                            if fmt_id >= 164 and (idx, fmt_id) in num_fmt_map:
                                xf.set('numFmtId', str(num_fmt_map[(idx, fmt_id)]))

                        if 'fontId' in xf.attrib:
                            font_id = int(xf.get('fontId'))
                            if (idx, font_id) in font_map:
                                xf.set('fontId', str(font_map[(idx, font_id)]))

                        if 'fillId' in xf.attrib:
                            fill_id = int(xf.get('fillId'))
                            if (idx, fill_id) in fill_map:
                                xf.set('fillId', str(fill_map[(idx, fill_id)]))

                        if 'borderId' in xf.attrib:
                            border_id = int(xf.get('borderId'))
                            if (idx, border_id) in border_map:
                                xf.set('borderId', str(border_map[(idx, border_id)]))

                        if 'xfId' in xf.attrib:
                            style_xf_id = int(xf.get('xfId'))
                            if (idx, style_xf_id) in cellStyleXfs_map:
                                xf.set('xfId', str(cellStyleXfs_map[(idx, style_xf_id)]))

                        xf_str = ET.tostring(xf, encoding='UTF-8').decode('UTF-8')

                        # Check if this xf already exists
                        existing_idx = None
                        for i, existing_xf_str in enumerate(all_cell_xfs):
                            if xf_str == existing_xf_str:
                                existing_idx = i
                                break

                        if existing_idx is not None:
                            cellXfs_map[(idx, xf_idx)] = existing_idx
                        else:
                            new_xf_idx = len(all_cell_xfs)
                            cellXfs_map[(idx, xf_idx)] = new_xf_idx
                            all_cell_xfs.append(xf_str)

                # Process dxfs (differential formats used for conditional formatting)
                dxfs_elem = styles_root.find('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}dxfs')
                if dxfs_elem is not None:
                    for dxf_idx, dxf in enumerate(dxfs_elem.findall('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}dxf')):
                        # Process font elements within dxf
                        font_elem = dxf.find('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}font')
                        if font_elem is not None:
                            # Update font color references if needed
                            color_elem = font_elem.find('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}color')
                            if color_elem is not None and 'indexed' in color_elem.attrib:
                                # You may need to map indexed colors if necessary
                                pass

                        # Process fill elements within dxf
                        fill_elem = dxf.find('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}fill')
                        if fill_elem is not None:
                            # Process patternFill elements
                            pattern_fill = fill_elem.find('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}patternFill')
                            if pattern_fill is not None:
                                bg_color = pattern_fill.find('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}bgColor')
                                if bg_color is not None and 'indexed' in bg_color.attrib:
                                    # Map indexed colors if necessary
                                    pass

                        # Process number format elements
                        numFmt_elem = dxf.find('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}numFmt')
                        if numFmt_elem is not None and 'numFmtId' in numFmt_elem.attrib:
                            fmt_id = int(numFmt_elem.get('numFmtId'))
                            if fmt_id >= 164 and (idx, fmt_id) in num_fmt_map:
                                numFmt_elem.set('numFmtId', str(num_fmt_map[(idx, fmt_id)]))

                        # Convert to string for comparison/storage
                        dxf_str = ET.tostring(dxf, encoding='UTF-8').decode('UTF-8')

                        # Check if this dxf already exists
                        existing_idx = None
                        for i, existing_dxf_str in enumerate(all_dxfs):
                            if dxf_str == existing_dxf_str:
                                existing_idx = i
                                break

                        if existing_idx is not None:
                            dxf_map[(idx, dxf_idx)] = existing_idx
                        else:
                            new_dxf_idx = len(all_dxfs)
                            dxf_map[(idx, dxf_idx)] = new_dxf_idx
                            all_dxfs.append(dxf_str)

                # Process cellStyles
                cellStyles_elem = styles_root.find('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}cellStyles')
                if cellStyles_elem is not None:
                    for cellStyle in cellStyles_elem.findall('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}cellStyle'):
                        if 'xfId' in cellStyle.attrib:
                            xf_id = int(cellStyle.get('xfId'))
                            if (idx, xf_id) in cellStyleXfs_map:
                                cellStyle.set('xfId', str(cellStyleXfs_map[(idx, xf_id)]))

                        cellStyle_str = ET.tostring(cellStyle, encoding='UTF-8').decode('UTF-8')
                        if cellStyle_str not in all_cell_styles:
                            all_cell_styles.append(cellStyle_str)

                ### Process tableStyles if present
                tableStyles_elem = styles_root.find('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}tableStyles')
                if tableStyles_elem is not None:
                    defaultTableStyle = tableStyles_elem.get('defaultTableStyle', 'TableStyleMedium2')
                    defaultPivotStyle = tableStyles_elem.get('defaultPivotStyle', 'PivotStyleLight16')
                ###

    # Create new merged styles.xml
    styles_xml = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
    styles_xml += '<styleSheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" xmlns:x14ac="http://schemas.microsoft.com/office/spreadsheetml/2009/9/ac" mc:Ignorable="x14ac">\n'

    # Add numFmts
    if all_number_formats:
        styles_xml += f'<numFmts count="{len(all_number_formats)}">\n'
        for num_fmt in all_number_formats:
            styles_xml += ET.tostring(num_fmt, encoding='UTF-8').decode('UTF-8')
        styles_xml += '</numFmts>\n'

    # Add fonts
    if all_fonts:
        styles_xml += f'<fonts count="{len(all_fonts)}" x14ac:knownFonts="1">\n'
        for font_str in all_fonts:
            styles_xml += font_str
        styles_xml += '</fonts>\n'

    # Add fills
    if all_fills:
        styles_xml += f'<fills count="{len(all_fills)}">\n'
        for fill_str in all_fills:
            styles_xml += fill_str
        styles_xml += '</fills>\n'

    # Add borders
    if all_borders:
        styles_xml += f'<borders count="{len(all_borders)}">\n'
        for border_str in all_borders:
            styles_xml += border_str
        styles_xml += '</borders>\n'

    # Add cellStyleXfs
    if all_cell_style_xfs:
        styles_xml += f'<cellStyleXfs count="{len(all_cell_style_xfs)}">\n'
        for xf_str in all_cell_style_xfs:
            styles_xml += xf_str
        styles_xml += '</cellStyleXfs>\n'

    # Add cellXfs
    if all_cell_xfs:
        styles_xml += f'<cellXfs count="{len(all_cell_xfs)}">\n'
        for xf_str in all_cell_xfs:
            styles_xml += xf_str
        styles_xml += '</cellXfs>\n'

    # Add cellStyles
    if all_cell_styles:
        styles_xml += f'<cellStyles count="{len(all_cell_styles)}">\n'
        for cellStyle_str in all_cell_styles:
            styles_xml += cellStyle_str
        styles_xml += '</cellStyles>\n'

    # Add dxfs
    if all_dxfs:
        styles_xml += f'<dxfs count="{len(all_dxfs)}">\n'
        for dxf_str in all_dxfs:
            styles_xml += dxf_str
        styles_xml += '</dxfs>\n'

    # Add default elements if needed
    styles_xml += f'<tableStyles count="0" defaultTableStyle="{defaultTableStyle}" defaultPivotStyle="{defaultPivotStyle}"/>\n'
    styles_xml += '</styleSheet>'

    output_zip.writestr('xl/styles.xml', styles_xml)
    added_files.add('xl/styles.xml')
    ###

    with ZipFile(source_files[0], 'r') as first_zip:
        for item in first_zip.namelist():
            if item.startswith('xl/theme/') and item not in added_files:
                theme_content = first_zip.read(item)
                output_zip.writestr(item, theme_content)
                added_files.add(item)

        essential_files = ['docProps/app.xml', 'docProps/core.xml', '_rels/.rels']

        for file_path in essential_files:
            if file_path in first_zip.namelist() and file_path not in added_files:
                content = first_zip.read(file_path)
                output_zip.writestr(file_path, content)
                added_files.add(file_path)

    for (wb_idx, old_id), new_id in pivot_cache_map.items():
        source_path = source_files[wb_idx]
        with ZipFile(source_path, 'r') as source_zip:
            old_def_path = f'xl/pivotCache/pivotCacheDefinition{old_id}.xml'
            new_def_path = f'xl/pivotCache/pivotCacheDefinition{new_id}.xml'

            if old_def_path in source_zip.namelist() and new_def_path not in added_files:
                content = source_zip.read(old_def_path)
                content_tree = ET.fromstring(content)

                xr_ns = "http://schemas.microsoft.com/office/spreadsheetml/2014/revision"
                uid_attr = f"{{{xr_ns}}}uid"
                if uid_attr in content_tree.attrib:
                    del content_tree.attrib[uid_attr]
                new_uid = "{" + str(uuid4()).upper() + "}"
                content_tree.set(uid_attr, new_uid)

                for elem in content_tree.iter():
                    if 'xmlns:x14' in elem.attrib:
                        del elem.attrib['xmlns:x14']

                ns = {'main': 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'}
                worksheet_source = content_tree.find('.//main:worksheetSource', ns)
                if worksheet_source is not None:
                    table_ref = worksheet_source.get('table')
                    if table_ref:
                        has_eq = table_ref.startswith('=')
                        original_name = table_ref.lstrip('=')

                        for (source_wb_idx, orig_name), mapped_name in name_map.items():
                            if source_wb_idx == wb_idx and orig_name == original_name:
                                worksheet_source.set('table', f'={mapped_name}' if has_eq else mapped_name)
                                break

                cache_source = content_tree.find('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}cacheSource')
                if cache_source is not None:
                    worksheet_source = cache_source.find('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}worksheetSource')
                    if worksheet_source is not None:
                        original_table_name = worksheet_source.get('name')
                        if original_table_name:
                            # Check if this table name has been renamed
                            if (wb_idx, original_table_name) in name_map:
                                new_table_name = name_map[(wb_idx, original_table_name)]
                                worksheet_source.set('name', new_table_name)

                output_zip.writestr(new_def_path, ET.tostring(content_tree, encoding='UTF-8'))
                added_files.add(new_def_path)

            old_rec_path = f'xl/pivotCache/pivotCacheRecords{old_id}.xml'
            new_rec_path = f'xl/pivotCache/pivotCacheRecords{new_id}.xml'

            if old_rec_path in source_zip.namelist() and new_rec_path not in added_files:
                content = source_zip.read(old_rec_path)
                output_zip.writestr(new_rec_path, content)
                added_files.add(new_rec_path)

            old_rels_path = f'xl/pivotCache/_rels/pivotCacheDefinition{old_id}.xml.rels'
            new_rels_path = f'xl/pivotCache/_rels/pivotCacheDefinition{new_id}.xml.rels'

            if old_rels_path in source_zip.namelist() and new_rels_path not in added_files:
                content = source_zip.read(old_rels_path)
                content_tree = ET.fromstring(content)

                ns = {'r': 'http://schemas.openxmlformats.org/package/2006/relationships'}
                for rel in content_tree.findall('.//r:Relationship', ns):
                    target = rel.get('Target')
                    if f'pivotCacheRecords{old_id}.xml' in target:
                        rel.set('Target', target.replace(f'pivotCacheRecords{old_id}.xml', f'pivotCacheRecords{new_id}.xml'))

                output_zip.writestr(new_rels_path, ET.tostring(content_tree, encoding='UTF-8'))
                added_files.add(new_rels_path)

    for (wb_idx, old_id), new_id in pivot_table_map.items():
        source_path = source_files[wb_idx]
        with ZipFile(source_path, 'r') as source_zip:
            old_path = f'xl/pivotTables/pivotTable{old_id}.xml'
            new_path = f'xl/pivotTables/pivotTable{new_id}.xml'

            if old_path in source_zip.namelist() and new_path not in added_files:
                content = source_zip.read(old_path)
                content_tree = ET.fromstring(content)

                ns = 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'

                pivot_name = final_pivot_names.get((wb_idx, old_id), f'PivotTable{new_id}')

                name_attr = f'{{{ns}}}name'
                if name_attr in content_tree.attrib:
                    pivot_name = content_tree.attrib[name_attr]
                    for attr_name in list(content_tree.attrib.keys()):
                        if attr_name.endswith('name') and attr_name != name_attr:
                            del content_tree.attrib[attr_name]
                    content_tree.attrib[name_attr] = pivot_name

                cache_id_attr = f'{{{ns}}}cacheId'
                if cache_id_attr in content_tree.attrib:
                    old_cache_id = int(content_tree.attrib[cache_id_attr])
                    if (wb_idx, old_cache_id) in pivot_cache_map:
                        content_tree.attrib[cache_id_attr] = str(pivot_cache_map[(wb_idx, old_cache_id)])

                output_zip.writestr(new_path, ET.tostring(content_tree, encoding='UTF-8'))
                added_files.add(new_path)

            old_rels_path = f'xl/pivotTables/_rels/pivotTable{old_id}.xml.rels'
            new_rels_path = f'xl/pivotTables/_rels/pivotTable{new_id}.xml.rels'

            if old_rels_path in source_zip.namelist() and new_rels_path not in added_files:
                content = source_zip.read(old_rels_path)
                content_tree = ET.fromstring(content)

                ns = 'http://schemas.openxmlformats.org/package/2006/relationships'
                for rel in content_tree.findall('.//{'+ns+'}Relationship'):
                    target = rel.get('Target')
                    if 'pivotCacheDefinition' in target:
                        old_cache_id = extract_number(target)
                        if old_cache_id is not None and (wb_idx, old_cache_id) in pivot_cache_map:
                            new_cache_id = pivot_cache_map[(wb_idx, old_cache_id)]
                            new_target = target.replace(f'pivotCacheDefinition{old_cache_id}', f'pivotCacheDefinition{new_cache_id}')
                            rel.set('Target', new_target)

                output_zip.writestr(new_rels_path, ET.tostring(content_tree, encoding='UTF-8'))
                added_files.add(new_rels_path)

    for (wb_idx, old_id), new_id in table_map.items():
        source_path = source_files[wb_idx]
        with ZipFile(source_path, 'r') as source_zip:
            old_path = f'xl/tables/table{old_id}.xml'
            new_path = f'xl/tables/table{new_id}.xml'

            if old_path in source_zip.namelist() and new_path not in added_files:
                content = source_zip.read(old_path)
                content_tree = ET.fromstring(content)

                content_tree.attrib['id'] = str(new_id)

                final_name = final_table_names.get((wb_idx, old_id), f'Table{new_id}')
                content_tree.set('name', final_name)
                content_tree.set('displayName', final_name)

                xr_ns = "http://schemas.microsoft.com/office/spreadsheetml/2014/revision"
                uid_attr = f"{{{xr_ns}}}uid"
                if uid_attr in content_tree.attrib:
                    del content_tree.attrib[uid_attr]
                new_table_uid = "{" + str(uuid4()).upper() + "}"
                content_tree.set(uid_attr, new_table_uid)

                auto_filter = content_tree.find('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}autoFilter')
                if auto_filter is not None:
                    auto_filter_uid_attr = f"{{{xr_ns}}}uid"
                    if auto_filter_uid_attr in auto_filter.attrib:
                        del auto_filter.attrib[auto_filter_uid_attr]
                    new_auto_filter_uid = "{" + str(uuid4()).upper() + "}"
                    auto_filter.set(auto_filter_uid_attr, new_auto_filter_uid)

                xr3_ns = "http://schemas.microsoft.com/office/spreadsheetml/2016/revision3"
                for table_column in content_tree.findall('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}tableColumn'):
                    xr3_uid_attr = f"{{{xr3_ns}}}uid"
                    if xr3_uid_attr in table_column.attrib:
                        del table_column.attrib[xr3_uid_attr]
                    new_column_uid = "{" + str(uuid4()).upper() + "}"
                    table_column.set(xr3_uid_attr, new_column_uid)

                output_zip.writestr(new_path, ET.tostring(content_tree, encoding='UTF-8'))
                added_files.add(new_path)

    ### Copy chart files
    for (wb_idx, old_id), new_id in chart_map.items():
        source_path = source_files[wb_idx]
        with ZipFile(source_path, 'r') as source_zip:
            old_chart_path = f'xl/charts/chart{old_id}.xml'
            new_chart_path = f'xl/charts/chart{new_id}.xml'

            if old_chart_path in source_zip.namelist() and new_chart_path not in added_files:
                # Read and parse the chart XML
                chart_content = source_zip.read(old_chart_path)
                chart_tree = ET.fromstring(chart_content)

                # Update c:f elements with new sheet names
                ns = {'c': 'http://schemas.openxmlformats.org/drawingml/2006/chart'}
                for f_element in chart_tree.findall('.//c:f', ns):
                    formula = f_element.text
                    if formula:
                        # Regex to parse sheet name from formula (handles quoted and unquoted names)
                        match = re.match(r"^(?:'([^']+)'!|([^!]+)!)(.*)$", formula)
                        if match:
                            quoted_sheet, unquoted_sheet, cell_ref = match.groups()
                            original_sheet = quoted_sheet if quoted_sheet is not None else unquoted_sheet
                            if original_sheet:
                                # Look up new sheet name in the mapping
                                new_sheet = sheet_name_mapping.get((wb_idx, original_sheet))
                                if new_sheet:
                                    # Rebuild formula with new sheet name, handling quotes if needed
                                    if quoted_sheet is not None:
                                        # Original was quoted; check if new name needs quotes
                                        if ' ' in new_sheet or "'" in new_sheet:
                                            new_sheet_temp = new_sheet.replace("'", "''")
                                            new_sheet_quoted = f"'{new_sheet_temp}'"
                                        else:
                                            new_sheet_quoted = new_sheet
                                        new_formula = f"{new_sheet_quoted}!{cell_ref}"
                                    else:
                                        # Original was unquoted; check if new name needs quotes
                                        if ' ' in new_sheet or "'" in new_sheet:
                                            new_sheet_temp2 = new_sheet.replace("'", "''")
                                            new_sheet_quoted = f"'{new_sheet_temp2}'"
                                            new_formula = f"{new_sheet_quoted}!{cell_ref}"
                                        else:
                                            new_formula = f"{new_sheet}!{cell_ref}"
                                    f_element.text = new_formula

                # Save modified chart XML
                modified_chart_content = ET.tostring(chart_tree, encoding='UTF-8')
                output_zip.writestr(new_chart_path, modified_chart_content)
                added_files.add(new_chart_path)

            # Process chart relationships file
            old_chart_rels_path = f'xl/charts/_rels/chart{old_id}.xml.rels'
            new_chart_rels_path = f'xl/charts/_rels/chart{new_id}.xml.rels'

            if old_chart_rels_path in source_zip.namelist() and new_chart_rels_path not in added_files:
                chart_rels_content = source_zip.read(old_chart_rels_path)
                chart_rels_tree = ET.fromstring(chart_rels_content)

                # Update relationship targets if needed
                ns = 'http://schemas.openxmlformats.org/package/2006/relationships'
                for rel in chart_rels_tree.findall('.//{'+ns+'}Relationship'):
                    target = rel.get('Target')

                    # Handle style files
                    if 'style' in target:
                        style_id = extract_number(target)
                        if style_id is not None:
                            new_style_path = f'xl/charts/style{new_id}.xml'
                            rel.set('Target', f'style{new_id}.xml')

                            # Copy the style file
                            old_style_path = f'xl/charts/style{style_id}.xml'
                            if old_style_path in source_zip.namelist() and new_style_path not in added_files:
                                style_content = source_zip.read(old_style_path)
                                output_zip.writestr(new_style_path, style_content)
                                added_files.add(new_style_path)

                    # Handle colors files
                    if 'colors' in target:
                        colors_id = extract_number(target)
                        if colors_id is not None:
                            new_colors_path = f'xl/charts/colors{new_id}.xml'
                            rel.set('Target', f'colors{new_id}.xml')

                            # Copy the colors file
                            old_colors_path = f'xl/charts/colors{colors_id}.xml'
                            if old_colors_path in source_zip.namelist() and new_colors_path not in added_files:
                                colors_content = source_zip.read(old_colors_path)
                                output_zip.writestr(new_colors_path, colors_content)
                                added_files.add(new_colors_path)

                output_zip.writestr(new_chart_rels_path, ET.tostring(chart_rels_tree, encoding='UTF-8'))
                added_files.add(new_chart_rels_path)

    # Copy drawing files
    for (wb_idx, old_id), new_id in drawing_map.items():
        source_path = source_files[wb_idx]
        with ZipFile(source_path, 'r') as source_zip:
            # Process drawing file
            old_drawing_path = f'xl/drawings/drawing{old_id}.xml'
            new_drawing_path = f'xl/drawings/drawing{new_id}.xml'

            if old_drawing_path in source_zip.namelist() and new_drawing_path not in added_files:
                drawing_content = source_zip.read(old_drawing_path)
                drawing_tree = ET.fromstring(drawing_content)

                # Update chart references in the drawing
                xdr_ns = "http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing"
                r_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

                # Update chart references
                for chart_ref in drawing_tree.findall(f'.//{{{xdr_ns}}}graphicFrame//{{{r_ns}}}id'):
                    r_id = chart_ref.get('id')
                    # We'll need to update this ID using the drawing relationships file
                    # This will be handled when processing the drawing relationships

                output_zip.writestr(new_drawing_path, ET.tostring(drawing_tree, encoding='UTF-8'))
                added_files.add(new_drawing_path)

            # Process drawing relationships file
            old_drawing_rels_path = f'xl/drawings/_rels/drawing{old_id}.xml.rels'
            new_drawing_rels_path = f'xl/drawings/_rels/drawing{new_id}.xml.rels'

            if old_drawing_rels_path in source_zip.namelist() and new_drawing_rels_path not in added_files:
                drawing_rels_content = source_zip.read(old_drawing_rels_path)
                drawing_rels_tree = ET.fromstring(drawing_rels_content)

                # Update relationship targets for charts
                ns = 'http://schemas.openxmlformats.org/package/2006/relationships'
                for rel in drawing_rels_tree.findall('.//{'+ns+'}Relationship'):
                    rel_id = rel.get('Id')
                    target = rel.get('Target')

                    if '../charts/chart' in target:
                        old_chart_id = extract_number(target)
                        if old_chart_id is not None and (wb_idx, old_chart_id) in chart_map:
                            new_chart_id = chart_map[(wb_idx, old_chart_id)]
                            rel.set('Target', target.replace(f'chart{old_chart_id}', f'chart{new_chart_id}'))

                output_zip.writestr(new_drawing_rels_path, ET.tostring(drawing_rels_tree, encoding='UTF-8'))
                added_files.add(new_drawing_rels_path)
    ###

    update_content_types(source_files, output_zip, num_sheets, pivot_cache_map, pivot_table_map, table_map, chart_map, drawing_map)

    return merged_shared_strings_map, pivot_cache_map, pivot_table_map, table_map, final_pivot_names, final_table_names, pivot_table_cacheids, dxf_map, chart_map, drawing_map, cellXfs_map, name_map

def update_content_types(source_files, output_zip, num_sheets, pivot_cache_map=None, pivot_table_map=None, table_map=None, chart_map=None, drawing_map=None):
    ct_ns = '{http://schemas.openxmlformats.org/package/2006/content-types}'

    existing_types_root = None
    try:
        with ZipFile(source_files[0], 'r') as source_zip:
            if '[Content_Types].xml' in source_zip.namelist():
                content = source_zip.read('[Content_Types].xml')
                existing_types_root = ET.fromstring(content)
    except Exception:
        pass

    if existing_types_root is None:
        types_root = ET.Element(ct_ns + 'Types')
        types_root.set('xmlns', 'http://schemas.openxmlformats.org/package/2006/content-types')

        default_types = {'xml': 'application/xml', 'rels': 'application/vnd.openxmlformats-package.relationships+xml', 'png': 'image/png', 'jpeg': 'image/jpeg', 'jpg': 'image/jpeg', 'vml': 'application/vnd.openxmlformats-officedocument.vmlDrawing'
        }

        for extension, content_type in default_types.items():
            default = ET.SubElement(types_root, ct_ns + 'Default')
            default.set('Extension', extension)
            default.set('ContentType', content_type)

        required_overrides = {
            '/xl/workbook.xml': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml',
            '/xl/styles.xml': 'application/vnd.openxmlformats-officedocument.spreadsheetml.styles+xml',
            '/xl/sharedStrings.xml': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sharedStrings+xml',
            '/xl/theme/theme1.xml': 'application/vnd.openxmlformats-officedocument.theme+xml',
            '/docProps/core.xml': 'application/vnd.openxmlformats-package.core-properties+xml',
            '/docProps/app.xml': 'application/vnd.openxmlformats-officedocument.extended-properties+xml'
        }

        for part_name, content_type in required_overrides.items():
            override = ET.SubElement(types_root, ct_ns + 'Override')
            override.set('PartName', part_name)
            override.set('ContentType', content_type)
    else:
        types_root = existing_types_root

    existing_overrides = set()
    for override in types_root.findall(f'{ct_ns}Override'):
        existing_overrides.add(override.get('PartName'))

    for i in range(1, num_sheets + 1):
        part_name = f'/xl/worksheets/sheet{i}.xml'
        if part_name not in existing_overrides:
            override = ET.SubElement(types_root, ct_ns + 'Override')
            override.set('PartName', part_name)
            override.set('ContentType', 'application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml')
            existing_overrides.add(part_name)

    if pivot_table_map:
        for new_id in pivot_table_map.values():
            part_name = f'/xl/pivotTables/pivotTable{new_id}.xml'
            if part_name not in existing_overrides:
                override = ET.SubElement(types_root, ct_ns + 'Override')
                override.set('PartName', part_name)
                override.set('ContentType', 'application/vnd.openxmlformats-officedocument.spreadsheetml.pivotTable+xml')
                existing_overrides.add(part_name)

    if pivot_cache_map:
        for new_id in pivot_cache_map.values():
            part_name = f'/xl/pivotCache/pivotCacheDefinition{new_id}.xml'
            if part_name not in existing_overrides:
                override = ET.SubElement(types_root, ct_ns + 'Override')
                override.set('PartName', part_name)
                override.set('ContentType', 'application/vnd.openxmlformats-officedocument.spreadsheetml.pivotCacheDefinition+xml')
                existing_overrides.add(part_name)

            part_name = f'/xl/pivotCache/pivotCacheRecords{new_id}.xml'
            if part_name not in existing_overrides:
                override = ET.SubElement(types_root, ct_ns + 'Override')
                override.set('PartName', part_name)
                override.set('ContentType', 'application/vnd.openxmlformats-officedocument.spreadsheetml.pivotCacheRecords+xml')
                existing_overrides.add(part_name)

    if table_map:
        for new_id in table_map.values():
            part_name = f'/xl/tables/table{new_id}.xml'
            if part_name not in existing_overrides:
                override = ET.SubElement(types_root, ct_ns + 'Override')
                override.set('PartName', part_name)
                override.set('ContentType', 'application/vnd.openxmlformats-officedocument.spreadsheetml.table+xml')
                existing_overrides.add(part_name)

    ### Add content types for charts and chart resources
    if 'chart_map' in locals() or 'chart_map' in globals():
        for new_id in chart_map.values():
            # Chart file
            part_name = f'/xl/charts/chart{new_id}.xml'
            if part_name not in existing_overrides:
                override = ET.SubElement(types_root, ct_ns + 'Override')
                override.set('PartName', part_name)
                override.set('ContentType', 'application/vnd.openxmlformats-officedocument.drawingml.chart+xml')
                existing_overrides.add(part_name)

            # Style file
            part_name = f'/xl/charts/style{new_id}.xml'
            if part_name not in existing_overrides and any(f'xl/charts/style{x}.xml' in source_zip.namelist() for source_zip in [ZipFile(f, 'r') for f in source_files] for x in range(1, 20)):
                override = ET.SubElement(types_root, ct_ns + 'Override')
                override.set('PartName', part_name)
                override.set('ContentType', 'application/vnd.ms-office.chartstyle+xml')
                existing_overrides.add(part_name)

            # Colors file
            part_name = f'/xl/charts/colors{new_id}.xml'
            if part_name not in existing_overrides and any(f'xl/charts/colors{x}.xml' in source_zip.namelist() for source_zip in [ZipFile(f, 'r') for f in source_files] for x in range(1, 20)):
                override = ET.SubElement(types_root, ct_ns + 'Override')
                override.set('PartName', part_name)
                override.set('ContentType', 'application/vnd.ms-office.chartcolorstyle+xml')
                existing_overrides.add(part_name)

    # Add content types for drawings
    if 'drawing_map' in locals() or 'drawing_map' in globals():
        for new_id in drawing_map.values():
            part_name = f'/xl/drawings/drawing{new_id}.xml'
            if part_name not in existing_overrides:
                override = ET.SubElement(types_root, ct_ns + 'Override')
                override.set('PartName', part_name)
                override.set('ContentType', 'application/vnd.openxmlformats-officedocument.drawing+xml')
                existing_overrides.add(part_name)
    ###

    for source_file in source_files:
        with ZipFile(source_file, 'r') as source_zip:
            for item in source_zip.namelist():
                part_name = '/' + item

                if item.startswith('xl/drawings/'):
                    if part_name not in existing_overrides:
                        override = ET.SubElement(types_root, ct_ns + 'Override')
                        override.set('PartName', part_name)
                        override.set('ContentType', 'application/vnd.openxmlformats-officedocument.drawing+xml')
                        existing_overrides.add(part_name)

                elif item.startswith('xl/charts/'):
                    if part_name not in existing_overrides:
                        override = ET.SubElement(types_root, ct_ns + 'Override')
                        override.set('PartName', part_name)
                        override.set('ContentType', 'application/vnd.openxmlformats-officedocument.drawingml.chart+xml')
                        existing_overrides.add(part_name)

    content_types_xml = ET.tostring(types_root, encoding='UTF-8', xml_declaration=True)
    output_zip.writestr('[Content_Types].xml', content_types_xml)

def update_cell_style_references(sheet_tree, source_wb_idx, cellXfs_map):
    """Update cell style indices (s attribute) to reference merged styles.xml cellXfs"""
    ns = get_namespace(sheet_tree)
    ns_dict = {'ns': ns} if ns else {}

    for cell in sheet_tree.findall('.//ns:c', ns_dict) if ns else sheet_tree.findall('.//c'):
        s_val = cell.get('s')
        if s_val is not None:
            original_s = int(s_val)
            # Look up the new cellXfs index using the source workbook index and original s value
            new_s = cellXfs_map.get((source_wb_idx, original_s), original_s)
            cell.set('s', str(new_s))

def update_table_references_in_formulas(sheet_tree, source_wb_idx, name_map, ns_dict):
    """
    Updates table references in formulas, handling both quoted and unquoted table names.
    Uses string operations instead of complex regex to avoid escaping issues.
    """
    for cell in sheet_tree.findall('.//ns:c', ns_dict):
        f_elem = cell.find('./ns:f', ns_dict)
        if f_elem is not None and f_elem.text:
            formula = f_elem.text
            new_formula = formula

            # Process table references by finding square brackets and looking backwards
            i = 0
            while i < len(new_formula):
                if new_formula[i] == '[':  # Found a potential table reference
                    # Look backwards to find the table name
                    if i > 0:
                        # Case 1: 'Table Name'[
                        if new_formula[i-1] == "'":  # Quoted name
                            # Find the matching opening quote
                            j = i - 2
                            found_start = False
                            while j >= 0:
                                if new_formula[j] == "'" and (j == 0 or new_formula[j-1] != "'"):
                                    # Found the start of the quoted name
                                    quoted_text = new_formula[j+1:i-1]
                                    # Replace doubled quotes with single quotes
                                    table_name = ""
                                    k = 0
                                    while k < len(quoted_text):
                                        if k < len(quoted_text) - 1 and quoted_text[k] == "'" and quoted_text[k+1] == "'":
                                            table_name += "'"
                                            k += 2
                                        else:
                                            table_name += quoted_text[k]
                                            k += 1

                                    key = (source_wb_idx, table_name)
                                    if key in name_map:
                                        new_name = name_map[key]
                                        # Escape single quotes in the new name
                                        escaped_name = ""
                                        for char in new_name:
                                            if char == "'":
                                                escaped_name += "''"
                                            else:
                                                escaped_name += char

                                        # Replace the old name with the new name
                                        prefix = new_formula[:j]
                                        suffix = new_formula[i:]
                                        new_formula = prefix + "'" + escaped_name + "'" + suffix
                                        i = j + len(escaped_name) + 3  # +3 for the quotes and '['
                                    else:
                                        i += 1  # Move past this bracket if no replacement
                                    found_start = True
                                    break
                                j -= 1

                            if not found_start:
                                i += 1  # Ensure we move forward if we don't find a matching quote

                        # Case 2: TableName[
                        else:
                            # Find the start of the unquoted name
                            j = i - 1
                            while j >= 0 and (new_formula[j].isalnum() or new_formula[j] == '_'):
                                j -= 1
                            j += 1  # Adjust to include the start character

                            if j < i:  # Found a valid name
                                table_name = new_formula[j:i]
                                key = (source_wb_idx, table_name)
                                if key in name_map:
                                    new_name = name_map[key]
                                    # Check if new name needs quotes
                                    if ' ' in new_name or '-' in new_name or "'" in new_name:
                                        # Escape single quotes for Excel
                                        escaped_name = ""
                                        for char in new_name:
                                            if char == "'":
                                                escaped_name += "''"
                                            else:
                                                escaped_name += char
                                        replacement = f"'{escaped_name}'"
                                    else:
                                        replacement = new_name

                                    # Replace the old name with the new name
                                    prefix = new_formula[:j]
                                    suffix = new_formula[i:]
                                    new_formula = prefix + replacement + suffix
                                    i = j + len(replacement) + 1  # +1 for the '['
                                else:
                                    i += 1  # Move past this bracket if no replacement
                            else:
                                i += 1  # Move forward if we couldn't find a valid name
                    else:
                        i += 1  # Move forward when [ is at the beginning of the string
                else:
                    i += 1  # Move forward when not a bracket

            if new_formula != formula:
                f_elem.text = new_formula

def process_sheet(sheet_info, source_zip, output_zip, sheet_index, next_sheet_rel_id, workbook_rels_root, sheets_element, temp_dir, shared_strings_map=None, pivot_cache_map=None, pivot_table_map=None,
                  table_map=None, final_pivot_names=None, final_table_names=None, pivot_table_cacheids=None, dxf_map=None, chart_map=None, drawing_map=None, cellXfs_map=None, name_map=None):
    source_wb = sheet_info['workbook']
    source_sheet_name = sheet_info['sheet_name']
    target_sheet_name = sheet_info['new_sheet_name']
    source_wb_idx = sheet_info['source_idx']

    source_sheet_idx = source_wb.sheetnames.index(source_sheet_name) + 1

    sheet_xml = f'xl/worksheets/sheet{source_sheet_idx}.xml'
    if sheet_xml in source_zip.namelist():
        sheet_content = source_zip.read(sheet_xml)
        sheet_tree = ET.fromstring(sheet_content)

        if shared_strings_map:
            update_sheet_shared_strings(sheet_tree, source_wb_idx, shared_strings_map)
        
        # Update table references in formulas
        if name_map:
            ns = get_namespace(sheet_tree)
            ns_dict = {'ns': ns} if ns else {}
            update_table_references_in_formulas(sheet_tree, source_wb_idx, name_map, ns_dict)

        update_sheet_dxf_references(sheet_tree, source_wb_idx, dxf_map)

        # Update cell style indices to reference merged styles.xml
        update_cell_style_references(sheet_tree, source_wb_idx, cellXfs_map)

        ns2_uri = "http://schemas.microsoft.com/office/spreadsheetml/2014/revision"
        uid_attr = f"{{{ns2_uri}}}uid"
        worksheet_elem = sheet_tree
        if uid_attr in worksheet_elem.attrib:
            del worksheet_elem.attrib[uid_attr]
        new_uid = "{" + str(uuid4()).upper() + "}"
        worksheet_elem.set(uid_attr, new_uid)

        worksheet_elem = sheet_tree

        ns_declarations = {
            'xmlns:x14ac': "http://schemas.microsoft.com/office/spreadsheetml/2009/9/ac",
            'xmlns:xr2': "http://schemas.microsoft.com/office/spreadsheetml/2015/revision2",
            'xmlns:xr3': "http://schemas.microsoft.com/office/spreadsheetml/2016/revision3"
        }

        worksheet_elem.attrib.update(ns_declarations)

        mc_uri = "http://schemas.openxmlformats.org/markup-compatibility/2006"
        worksheet_elem.set(f'{{{mc_uri}}}Ignorable', 'x14ac xr xr2 xr3')

        xr_uri = "http://schemas.microsoft.com/office/spreadsheetml/2014/revision"
        worksheet_elem.set(f'{{{xr_uri}}}uid', '{{{}}}'.format(str(uuid4()).upper()))

        output_sheet_xml = f'xl/worksheets/sheet{sheet_index}.xml'

        sheet_rel_id = f"rId{next_sheet_rel_id}"

        rel_element = ET.SubElement(workbook_rels_root, '{http://schemas.openxmlformats.org/package/2006/relationships}Relationship')
        rel_element.set('Id', sheet_rel_id)
        rel_element.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet')
        rel_element.set('Target', f'worksheets/sheet{sheet_index}.xml')

        sheet_element = ET.SubElement(sheets_element, '{http://schemas.openxmlformats.org/spreadsheetml/2006/main}sheet')
        sheet_element.set('name', target_sheet_name)
        sheet_element.set('sheetId', str(sheet_index))
        sheet_element.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id', sheet_rel_id)

        sheet_rels_path = f'xl/worksheets/_rels/sheet{source_sheet_idx}.xml.rels'
        if sheet_rels_path in source_zip.namelist():
            sheet_rels_content = source_zip.read(sheet_rels_path)
            sheet_rels_tree = ET.fromstring(sheet_rels_content)

            ns = 'http://schemas.openxmlformats.org/package/2006/relationships'
            new_rels_root = ET.Element('{'+ns+'}Relationships')
            new_rels_root.set('xmlns', ns)

            copied_files = set()
            rel_remaps = {}

            for rel in sheet_rels_tree.findall('.//{'+ns+'}Relationship'):
                rel_id = rel.get('Id')
                rel_type = rel.get('Type')
                rel_target = rel.get('Target')

                new_rel_id = rel_id
                new_rel = ET.SubElement(new_rels_root, '{'+ns+'}Relationship')

                if 'pivottable' in rel_type.lower():
                    pivot_id = extract_number(rel_target)
                    if pivot_id is not None and (source_wb_idx, pivot_id) in pivot_table_map:
                        new_pivot_id = pivot_table_map[(source_wb_idx, pivot_id)]
                        new_target = rel_target.replace(f'pivotTable{pivot_id}', f'pivotTable{new_pivot_id}')
                        new_rel.set('Target', new_target)

                        new_rel_id = f"rId{10000 + new_pivot_id}"
                        rel_remaps[rel_id] = new_rel_id
                        new_rel.set('Id', new_rel_id)
                        new_rel.set('Type', rel_type)

                        pivot_path = f"xl{rel_target}" if rel_target.startswith('/') else f"xl/{rel_target}"
                        if pivot_path in source_zip.namelist() and pivot_path not in copied_files:
                            pivot_content = source_zip.read(pivot_path)
                            pivot_tree = ET.fromstring(pivot_content)

                            cache_id_attr = '{http://schemas.openxmlformats.org/spreadsheetml/2006/main}cacheId'
                            if cache_id_attr in pivot_tree.attrib:
                                old_cache_id = int(pivot_tree.attrib[cache_id_attr])
                                if (source_wb_idx, old_cache_id) in pivot_cache_map:
                                    pivot_tree.attrib[cache_id_attr] = str(pivot_cache_map[(source_wb_idx, old_cache_id)])

                            main_ns = 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'

                            for attr_name in list(pivot_tree.attrib.keys()):
                                if attr_name.endswith('name') and attr_name != f'{{{main_ns}}}name':
                                    del pivot_tree.attrib[attr_name]

                            pivot_name = final_pivot_names.get((source_wb_idx, pivot_id), f'PivotTable{new_pivot_id}')
                            pivot_tree.set(f'{{{main_ns}}}name', pivot_name)

                            new_pivot_path = pivot_path.replace(f'pivotTable{pivot_id}', f'pivotTable{new_pivot_id}')
                            output_zip.writestr(new_pivot_path, ET.tostring(pivot_tree, encoding='UTF-8'))
                            copied_files.add(pivot_path)

                            pivot_rels_path = f"{os.path.dirname(pivot_path)}/_rels/{os.path.basename(pivot_path)}.rels"
                            if pivot_rels_path in source_zip.namelist() and pivot_rels_path not in copied_files:
                                pivot_rels_content = source_zip.read(pivot_rels_path)
                                pivot_rels_tree = ET.fromstring(pivot_rels_content)

                                for pivot_rel in pivot_rels_tree.findall('.//{'+ns+'}Relationship'):
                                    pivot_target = pivot_rel.get('Target')
                                    if 'pivotCacheDefinition' in pivot_target:
                                        old_cache_id = extract_number(pivot_target)
                                        if (source_wb_idx, old_cache_id) in pivot_cache_map:
                                            new_cache_id = pivot_cache_map[(source_wb_idx, old_cache_id)]
                                            pivot_rel.set('Target', pivot_target.replace(f'pivotCacheDefinition{old_cache_id}', f'pivotCacheDefinition{new_cache_id}'))

                                new_pivot_rels_path = pivot_rels_path.replace(f'pivotTable{pivot_id}', f'pivotTable{new_pivot_id}')
                                output_zip.writestr(new_pivot_rels_path, ET.tostring(pivot_rels_tree, encoding='UTF-8'))
                                copied_files.add(pivot_rels_path)
                    else:
                        new_rel.set('Id', rel_id)
                        new_rel.set('Type', rel_type)
                        new_rel.set('Target', rel_target)

                elif 'table' in rel_type.lower():
                    table_id = extract_number(rel_target)
                    if table_id is not None and (source_wb_idx, table_id) in table_map:
                        new_table_id = table_map[(source_wb_idx, table_id)]
                        new_target = rel_target.replace(f'table{table_id}', f'table{new_table_id}')
                        new_rel.set('Target', new_target)

                        new_rel_id = f"rId{20000 + new_table_id}"
                        rel_remaps[rel_id] = new_rel_id
                        new_rel.set('Id', new_rel_id)
                        new_rel.set('Type', rel_type)

                        table_path = f"xl{rel_target}" if rel_target.startswith('/') else f"xl/{rel_target}"
                        if table_path in source_zip.namelist() and table_path not in copied_files:
                            table_content = source_zip.read(table_path)
                            table_tree = ET.fromstring(table_content)

                            table_tree.attrib['id'] = str(new_table_id)

                            final_name = final_table_names.get((source_wb_idx, table_id), f'Table{new_table_id}')
                            table_tree.set('name', final_name)
                            table_tree.set('displayName', final_name)

                            xr_ns = "http://schemas.microsoft.com/office/spreadsheetml/2014/revision"
                            uid_attr = f"{{{xr_ns}}}uid"
                            if uid_attr in table_tree.attrib:
                                del table_tree.attrib[uid_attr]
                            new_uid = "{" + str(uuid4()).upper() + "}"
                            table_tree.set(uid_attr, new_uid)

                            new_table_path = table_path.replace(f'table{table_id}', f'table{new_table_id}')
                            output_zip.writestr(new_table_path, ET.tostring(table_tree, encoding='UTF-8'))
                            copied_files.add(table_path)
                    else:
                        new_rel.set('Id', rel_id)
                        new_rel.set('Type', rel_type)
                        new_rel.set('Target', rel_target)
                ###
                elif 'drawing' in rel_type.lower():
                    drawing_path = f"xl{rel_target}" if rel_target.startswith('/') else f"xl/{rel_target}"
                    drawing_path = drawing_path.replace('../', '')
                    if drawing_path in source_zip.namelist() and drawing_path not in copied_files:
                        new_rel.set('Id', rel_id)
                        new_rel.set('Type', rel_type)

                        drawing_id = extract_number(drawing_path)
                        new_drawing_id = None
                        new_drawing_path = drawing_path

                        if drawing_id is not None and (source_wb_idx, drawing_id) in drawing_map:
                            new_drawing_id = drawing_map[(source_wb_idx, drawing_id)]
                            new_drawing_path = drawing_path.replace(f'drawing{drawing_id}', f'drawing{new_drawing_id}')
                            new_rel.set('Target', rel_target.replace(f'drawing{drawing_id}', f'drawing{new_drawing_id}'))

                            # Handle the drawing relationships
                            dir_path, file_name = os.path.split(drawing_path)
                            drawing_rels_path = f"{dir_path}/_rels/{file_name}.rels"

                            if drawing_rels_path in source_zip.namelist() and drawing_rels_path not in copied_files:
                                drawing_rels_content = source_zip.read(drawing_rels_path)
                                drawing_rels_tree = ET.fromstring(drawing_rels_content)

                                # Update chart references in relationships
                                for draw_rel in drawing_rels_tree.findall('.//{'+ns+'}Relationship'):
                                    draw_rel_target = draw_rel.get('Target')

                                    # Handle chart references
                                    if '../charts/chart' in draw_rel_target:
                                        chart_id = extract_number(draw_rel_target)
                                        if chart_id is not None and (source_wb_idx, chart_id) in chart_map:
                                            new_chart_id = chart_map[(source_wb_idx, chart_id)]
                                            draw_rel.set('Target', draw_rel_target.replace(f'chart{chart_id}', f'chart{new_chart_id}'))

                                # Process all referenced files
                                for draw_rel in drawing_rels_tree.findall('.//{'+ns+'}Relationship'):
                                    draw_rel_target = draw_rel.get('Target')

                                    if draw_rel_target.startswith('../'):
                                        draw_file_path = os.path.normpath(
                                            os.path.join(os.path.dirname(drawing_path), draw_rel_target))

                                        if draw_file_path in source_zip.namelist() and draw_file_path not in copied_files:
                                            # Handle chart files specially
                                            if 'charts/chart' in draw_file_path:
                                                chart_id = extract_number(draw_file_path)
                                                if chart_id is not None and (source_wb_idx, chart_id) in chart_map:
                                                    new_chart_id = chart_map[(source_wb_idx, chart_id)]
                                                    new_chart_path = draw_file_path.replace(f'chart{chart_id}', f'chart{new_chart_id}')

                                                    # Copy the chart file
                                                    chart_content = source_zip.read(draw_file_path)
                                                    output_zip.writestr(new_chart_path, chart_content)

                                                    # Process chart relationships
                                                    chart_rels_path = f"{os.path.dirname(draw_file_path)}/_rels/chart{chart_id}.xml.rels"
                                                    if chart_rels_path in source_zip.namelist():
                                                        chart_rels_content = source_zip.read(chart_rels_path)
                                                        chart_rels_tree = ET.fromstring(chart_rels_content)

                                                        # Copy any referenced style or colors files
                                                        for chart_rel in chart_rels_tree.findall('.//{'+ns+'}Relationship'):
                                                            chart_rel_target = chart_rel.get('Target')

                                                            if 'style' in chart_rel_target or 'colors' in chart_rel_target:
                                                                file_id = extract_number(chart_rel_target)
                                                                if file_id is not None:
                                                                    new_file_path = f"xl/charts/{os.path.basename(chart_rel_target).replace(str(file_id), str(new_chart_id))}"
                                                                    chart_rel.set('Target', os.path.basename(new_file_path))

                                                                    # Copy the style/colors file
                                                                    file_path = f"xl/charts/{os.path.basename(chart_rel_target)}"
                                                                    if file_path in source_zip.namelist():
                                                                        file_content = source_zip.read(file_path)
                                                                        output_zip.writestr(new_file_path, file_content)
                                                                        copied_files.add(file_path)

                                                        # Write the updated chart relationships
                                                        new_chart_rels_path = f"{os.path.dirname(draw_file_path)}/_rels/chart{new_chart_id}.xml.rels"
                                                        output_zip.writestr(new_chart_rels_path, ET.tostring(chart_rels_tree, encoding='UTF-8'))
                                                        copied_files.add(chart_rels_path)
                                                else:
                                                    # If not being remapped, just copy it as is
                                                    draw_content = source_zip.read(draw_file_path)
                                                    output_zip.writestr(draw_file_path, draw_content)
                                            else:
                                                # For non-chart files, just copy them
                                                draw_content = source_zip.read(draw_file_path)
                                                output_zip.writestr(draw_file_path, draw_content)

                                            copied_files.add(draw_file_path)
                    else:
                        new_rel.set('Id', rel_id)
                        new_rel.set('Type', rel_type)
                        new_rel.set('Target', rel_target)
                ###
                else:
                    new_rel.set('Id', rel_id)
                    new_rel.set('Type', rel_type)
                    new_rel.set('Target', rel_target)

            if rel_remaps:
                sheet_ns = get_namespace(sheet_tree)
                if sheet_ns:
                    ns_dict = {'ns': sheet_ns}

                    for elem in sheet_tree.findall('.//*[@r:id]', {'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'}):
                        old_rel_id = elem.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                        if old_rel_id in rel_remaps:
                            elem.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id', rel_remaps[old_rel_id])

                    updated_sheet_xml = ET.tostring(sheet_tree, encoding='UTF-8')
                    output_zip.writestr(output_sheet_xml, updated_sheet_xml)

                    new_rels_xml = ET.tostring(new_rels_root, encoding='UTF-8')
                    output_zip.writestr(f'xl/worksheets/_rels/sheet{sheet_index}.xml.rels', new_rels_xml)
        else:
            output_zip.writestr(output_sheet_xml, ET.tostring(sheet_tree, encoding='UTF-8'))

        return sheet_rel_id, next_sheet_rel_id + 1

    return None, next_sheet_rel_id

def extract_number(text):
    matches = re.findall(r'\d+', text)
    if matches:
        return int(matches[0])
    return None

def get_namespace(element):
    if element.tag.startswith('{'):
        return element.tag.split('}')[0][1:]
    return None






@app.route(route="get_excel_sheet_data")
def get_excel_sheet_data(req: func.HttpRequest) -> func.HttpResponse:
    try:
        # Parse request parameters
        file_base64 = req.get_json()['file']['$content']
        file_type = req.get_json()['file']['$content-type']
        sheet_name = req.get_json()['sheet_name']
        skip_top_rows = int(req.get_json().get('skip_top_rows', 0))
        skip_bottom_rows = int(req.get_json().get('skip_bottom_rows', 0))
        skip_left_columns = int(req.get_json().get('skip_left_columns', 0))
        skip_right_columns = int(req.get_json().get('skip_right_columns', 0))
        select_columns = req.get_json().get('select_columns', None)
        if select_columns:
            select_columns = req.get_json().get('select_columns').replace(', ', ',').split(',')
        remove_columns = req.get_json().get('remove_columns', None)
        if remove_columns:
            remove_columns = req.get_json().get('remove_columns').replace(', ', ',').split(',')
        filter_query = req.get_json().get('filter_query', None)
        skip_records_count = int(req.get_json().get('skip_records_count', 0))
        top_records_count = int(req.get_json().get('top_records_count', 0))

        # Decode the base64 file
        file_bytes = base64.b64decode(file_base64)
        excel_file = BytesIO(file_bytes)

        def convert_xls_to_xlsx(xls_file, sheet_name):
            book_xls = open_workbook(file_contents=xls_file.read())
            book_xlsx = Workbook()
            sheet_xlsx = book_xlsx.create_sheet(title=sheet_name)

            sheet_xls = book_xls.sheet_by_name(sheet_name)
            for row in range(sheet_xls.nrows):
                for col in range(sheet_xls.ncols):
                    sheet_xlsx.cell(row=row+1, column=col+1).value = sheet_xls.cell_value(row, col)

            xlsx_io = BytesIO()
            book_xlsx.save(xlsx_io)
            xlsx_io.seek(0)
            return xlsx_io

        # Determine file type and convert if necessary
        if file_type == "application/vnd.ms-excel":
            # Convert xls to xlsx
            excel_file = convert_xls_to_xlsx(excel_file, sheet_name)

        # Process xlsx file
        workbook = load_workbook(excel_file)
        sheet = workbook[sheet_name]

        # Find the first row and column with data
        def find_first_data_cell(sheet):
            min_row, min_col = None, None
            # Find the first row with data
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value is not None:
                        if min_row is None or cell.row < min_row:
                            min_row = cell.row
            # Find the first column with data
            for col in sheet.iter_cols():
                for cell in col:
                    if cell.value is not None:
                        if min_col is None or cell.column < min_col:
                            min_col = cell.column

            return min_row, min_col

        first_row, first_col = find_first_data_cell(sheet)

        # Extract data starting from the first data cell
        data = sheet.iter_rows(min_row=first_row + skip_top_rows, max_row=sheet.max_row - skip_bottom_rows,
                           min_col=first_col + skip_left_columns, max_col=sheet.max_column - skip_right_columns,
                           values_only=True)

        # Extract column headers
        columns = next(data)

        # Handle duplicate column names
        def get_unique_columns(columns):
            seen = {}
            result = []
            for col in columns:
                if col in seen:
                    seen[col] += 1
                    result.append(f"{col} {seen[col]}")
                else:
                    seen[col] = 0
                    result.append(col)
            return result

        columns = get_unique_columns(columns)

        # Create DataFrame
        df = DataFrame(data, columns=columns)

        # Apply filter query
        if filter_query:
            df = df.query(filter_query)

        # Apply select or remove columns
        if select_columns:
            df = df[select_columns]
        elif remove_columns:
            df = df.drop(columns=remove_columns)

        # Apply skip and top records count
        if skip_records_count > 0:
            df = df.iloc[skip_records_count:]
        if top_records_count > 0:
            df = df.iloc[:top_records_count]

        # Convert DataFrame to JSON
        result = df.to_json(orient='records')

        return func.HttpResponse(result, mimetype="application/json")

    except Exception as e:
        if 'DataFrame columns must be unique' in str(e) or "dtype 'None    object" in str(e):
            return func.HttpResponse(f"Error: {str(e)} The program uses the 1st non-skipped row for column headers/JSON keys. If your sheet has rows with data before your column headers either remove those rows or use the skip_top_rows parameter to skip them. Alternatively, if it is already pointing to the correct header row, then the initial range may not have column headers for some columns that contain data. You will either need to add column headers in the header row or use the skip_left_columns & skip_right_columns parameters to skip those columns of unlabelled data.", status_code=400)
        else:
            return func.HttpResponse(f"Error: {str(e)}", status_code=400)






@app.route(route="extract_pdf_elements")
def extract_pdf_elements(req: func.HttpRequest) -> func.HttpResponse:
    try:
        req_json = req.get_json()
        pdf_bytes = base64_to_pdf(req_json['file_content'].get('$content'))

        text_exists, all_pages_text_exists, elements_layer_content = extract_pdf_elements(pdf_bytes)

        response_data = {
            "text_exists": text_exists,
            "all_pages_text_exists": all_pages_text_exists,
            **elements_layer_content  # Spread the text layer content directly into the response
        }

        return func.HttpResponse(
            body=json.dumps(response_data, cls=BytesEncoder),
            mimetype="application/json",
            status_code=200
        )
    except Exception as e:
        error_traceback = traceback.format_exc()
        error_response = {
            "error": str(e),
            "traceback": error_traceback
        }
        return func.HttpResponse(
            body=json.dumps(error_response, cls=BytesEncoder),
            mimetype="application/json",
            status_code=400
        )

class BytesEncoder(json.JSONEncoder):
    def default(self, obj):
        if isinstance(obj, bytes):
            return base64.b64encode(obj).decode('utf-8')
        # Handle PyMuPDF Rect objects
        elif hasattr(obj, '__class__') and obj.__class__.__name__ == 'Rect':
            return {
                "x0": obj.x0,
                "y0": obj.y0,
                "x1": obj.x1,
                "y1": obj.y1,
                "width": obj.width,
                "height": obj.height
            }
        return super().default(obj)

def base64_to_pdf(base64_string):
    file_bytes = base64.b64decode(base64_string)
    if file_bytes[0:4] != b"%PDF":
        raise ValueError("Missing the PDF file signature")
    return file_bytes

def extract_pdf_elements(pdf_bytes: bytes) -> tuple:
    """
    Extract text layer information from the PDF along with images and other elements.
    Returns:
        tuple: (text_exists, all_pages_text_exists, text_layer_content)
            - text_exists: True if at least one page has text
            - all_pages_text_exists: True if all pages have text
            - text_layer_content: Dictionary with text content, images, and other elements
    """
    pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
    text_exists = False
    all_pages_text_exists = True

    # Create response structure
    text_layer_content = {
        "document_metadata": {
            "page_count": len(pdf_document),
            "title": pdf_document.metadata.get("title", ""),
            "author": pdf_document.metadata.get("author", ""),
            "subject": pdf_document.metadata.get("subject", ""),
            "keywords": pdf_document.metadata.get("keywords", ""),
            "creation_date": pdf_document.metadata.get("creationDate", ""),
            "modification_date": pdf_document.metadata.get("modDate", "")
        },
        "elements_with_positions": [],  # Will contain all elements with position info (renamed)
        "images": [],                  # Will contain all images with metadata
        "links": [],                   # Will contain all links
        "form_fields": [],             # Will contain all form fields
        "annotations": [],             # Will contain all annotations
        "pages": []                    # Will contain simplified page-by-page information
    }

    # Process each page
    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)

        # Get basic text content
        page_text = page.get_text()
        has_text = bool(page_text.strip())

        # Update text existence flags
        text_exists = text_exists or has_text
        all_pages_text_exists = all_pages_text_exists and has_text

        # Create page-specific entry (simplified)
        page_entry = {
            "page_number": page_num + 1,
            "has_text": has_text,
            "text": page_text,
            "width": page.rect.width,
            "height": page.rect.height,
            "rotation": page.rotation
        }

        try:
            # Extract text in different formats
            # Dict format (PyMuPDF's native structured format with position data)
            dict_data = page.get_text("dict")

            # Flatten the text blocks structure
            flattened_elements = flatten_text_blocks(dict_data.get("blocks", []), page_num)
            text_layer_content["elements_with_positions"].extend(flattened_elements)

            # Extract images
            page_images = extract_images_from_page(page, page_num)

            # Add images to full document collection
            for img in page_images:
                img_with_page = img.copy()
                img_with_page["page_number"] = page_num + 1
                text_layer_content["images"].append(img_with_page)

            # Extract links
            links = page.get_links()

            # Add links to full document collection
            for link in links:
                link_with_page = link.copy()
                link_with_page["page_number"] = page_num + 1
                text_layer_content["links"].append(link_with_page)

            # Extract form fields (if any)
            form_fields = extract_form_fields(page, page_num)

            # Add form fields to full document collection
            for field in form_fields:
                text_layer_content["form_fields"].append(field)

            # Extract annotations
            annotations = extract_annotations(page, page_num)

            # Add annotations to full document collection
            for annotation in annotations:
                text_layer_content["annotations"].append(annotation)

        except Exception as e:
            # If extraction fails for this page, record the error
            error_info = {
                "message": str(e),
                "traceback": traceback.format_exc()
            }
            page_entry["extraction_error"] = error_info

        text_layer_content["pages"].append(page_entry)

    # Generate spatially-aware text representation
    text_layer_content["text"] = generate_spatially_aware_text(text_layer_content["elements_with_positions"])

    # Move elements from full_document up one level
    text_layer_content = restructure_response(text_layer_content)

    return text_exists, all_pages_text_exists, text_layer_content

def extract_images_from_page(page, page_num):
    """
    Extract images from a PDF page with metadata.
    Args:
        page: PyMuPDF page object
        page_num: Page number (0-based)
    Returns:
        list: List of dictionaries containing image data and metadata
    """
    images = []

    try:
        # Get list of image objects in the page
        image_list = page.get_images(full=True)

        for img_index, img_info in enumerate(image_list):
            try:
                # Unpack image information
                xref, smask, width, height, bpc, colorspace, altcolor, name, filter, referencer = img_info

                # Get the image metadata
                image_meta = {
                    "index": img_index,
                    "reference": xref,
                    "width": width,
                    "height": height,
                    "bits_per_component": bpc,
                    "colorspace": colorspace,
                    "name": name,
                    "filter": filter,
                    "position": None  # Will be filled if we can find it
                }

                # Try to find position information for this image
                image_rect = find_image_rectangle(page, xref)
                if image_rect:
                    image_meta["position"] = {
                        "x0": image_rect.x0,
                        "y0": image_rect.y0,
                        "x1": image_rect.x1,
                        "y1": image_rect.y1,
                        "width": image_rect.width,
                        "height": image_rect.height
                    }

                # Get image base64 data (optional - can be commented out if response size is a concern)
                try:
                    # Get the pixel map
                    pix = fitz.Pixmap(page.parent, xref)

                    # Convert to appropriate format
                    if pix.n - pix.alpha > 3:  # CMYK
                        pix = fitz.Pixmap(fitz.csRGB, pix)

                    # Convert to base64 string
                    img_bytes = pix.tobytes("png")
                    img_base64 = base64.b64encode(img_bytes).decode('utf-8')

                    # Limit image data size (optional)
                    if len(img_base64) < 1024 * 1024:  # Only include if less than ~1MB
                        image_meta["data"] = "data:image/png;base64," + img_base64
                    else:
                        image_meta["data"] = "Image data too large to include"

                    # Clean up
                    pix = None
                except Exception as e:
                    image_meta["data_error"] = str(e)

                images.append(image_meta)
            except Exception as e:
                images.append({
                    "index": img_index,
                    "error": str(e)
                })
    except Exception as e:
        images.append({
            "extraction_error": str(e),
            "traceback": traceback.format_exc()
        })

    return images

def find_image_rectangle(page, xref):
    """
    Try to find the rectangle for an image with the given xref.
    Args:
        page: PyMuPDF page object
        xref: Image reference number
    Returns:
        fitz.Rect or None: The rectangle where the image is displayed
    """
    try:
        # Iterate through the page's display list
        dl = page.get_displaylist()

        for item in dl:
            if item[0] == "i":  # Image item type
                if item[1] == xref:  # This is our image
                    return fitz.Rect(item[2])  # Return the rectangle
    except:
        pass

    return None

def extract_form_fields(page, page_num):
    """
    Extract form fields from a PDF page.
    Args:
        page: PyMuPDF page object
        page_num: Page number (0-based)
    Returns:
        list: List of dictionaries containing form field data
    """
    fields = []

    try:
        # Try to get form fields if this is a form
        widgets = page.widgets()

        for i, widget in enumerate(widgets):
            try:
                field = {
                    "index": i,
                    "page_number": page_num + 1,
                    "type": widget.field_type,
                    "type_name": widget.field_type_string,
                    "name": widget.field_name,
                    "value": widget.field_value,
                    "rect": {
                        "x0": widget.rect.x0,
                        "y0": widget.rect.y0,
                        "x1": widget.rect.x1,
                        "y1": widget.rect.y1
                    }
                }

                # Get field flags
                if hasattr(widget, "field_flags"):
                    field["flags"] = widget.field_flags

                # For choice fields, get available options
                if widget.field_type == 4:  # Choice field
                    field["options"] = widget.choice_values

                fields.append(field)
            except Exception as e:
                fields.append({
                    "index": i,
                    "page_number": page_num + 1,
                    "error": str(e)
                })
    except Exception as e:
        pass  # No form fields or error accessing them

    return fields

def extract_annotations(page, page_num):
    """
    Extract annotations from a PDF page.
    Args:
        page: PyMuPDF page object
        page_num: Page number (0-based)
    Returns:
        list: List of dictionaries containing annotation data
    """
    annotations = []

    try:
        for i, annot in enumerate(page.annots()):
            try:
                annotation = {
                    "index": i,
                    "page_number": page_num + 1,
                    "type": annot.type[0],
                    "type_name": annot.type[1],
                    "rect": {
                        "x0": annot.rect.x0,
                        "y0": annot.rect.y0,
                        "x1": annot.rect.x1,
                        "y1": annot.rect.y1
                    }
                }

                # Get contents if present
                if annot.info.get("content"):
                    annotation["content"] = annot.info["content"]

                # Get additional info depending on annotation type
                if annot.type[0] == 8:  # Highlight
                    quads = annot.vertices
                    if quads:
                        annotation["highlight_areas"] = quads
                elif annot.type[0] == 0:  # Text annotation
                    for key, value in annot.info.items():
                        if key not in ["content", "subject"] and value:
                            annotation[key] = value

                annotations.append(annotation)
            except Exception as e:
                annotations.append({
                    "index": i,
                    "page_number": page_num + 1,
                    "error": str(e)
                })
    except Exception as e:
        pass  # No annotations or error accessing them

    return annotations

def flatten_text_blocks(blocks, page_num):
    """
    Flatten the text blocks structure to have each span as a top-level object.
    Args:
        blocks: List of text blocks from PyMuPDF
        page_num: Page number (0-based)
    Returns:
        list: Flattened list of text elements
    """
    flattened_elements = []

    for block in blocks:
        # Process text blocks with lines and spans
        if block.get("type") == 0 and "lines" in block:
            for line in block["lines"]:
                if "spans" in line:
                    for span in line["spans"]:
                        # Create a new element for each span
                        element = {
                            "type": 0,  # Text
                            "page_number": page_num + 1,
                            "bbox": span.get("bbox"),
                            "text": span.get("text", ""),
                            "font": span.get("font", ""),
                            "color": span.get("color", 0),
                            "size": span.get("size", 0),
                            "flags": span.get("flags", 0),
                            "origin": {
                                "block_type": block.get("type"),
                                "block_bbox": block.get("bbox"),
                                "line_bbox": line.get("bbox"),
                                "span_number": span.get("span_number", 0)
                            }
                        }
                        flattened_elements.append(element)

        # Handle all other block types (images, paths, shadings, form XObjects, etc.)
        else:
            block["page_number"] = page_num + 1
            flattened_elements.append(block)

    return flattened_elements

def generate_spatially_aware_text(elements):
    """
    Generate text that preserves spatial relationships between text elements
    while properly handling page boundaries.
    Scales down coordinates by dividing by 2 to achieve 1 space per 2 x units
    and 1 newline per 2 y units.
    Args:
        elements: List of text elements with positions
    Returns:
        str: Spatially formatted text
    """
    # Filter only text elements
    text_elements = [e for e in elements if e.get("type") == 0 and e.get("text", "").strip()]

    if not text_elements:
        return ""

    # Calculate midpoints for text elements using scaled-down coordinates (divided by 2.5)
    for e in text_elements:
        bbox = e.get("bbox", [0, 0, 0, 0])
        # Scale down coordinates by dividing by 2
        e["mid_x"] = round((bbox[0] + bbox[2]) / 5)  # Divide by 2 for midpoint, then by 2.5 again for scaling
        e["mid_y"] = round((bbox[1] + bbox[3]) / 5)  # Divide by 2 for midpoint, then by 2.5 again for scaling
        e["char_len"] = len(e.get("text", ""))
        e["height"] = (bbox[3] - bbox[1]) / 2.5  # Scale down height by dividing by 2.5

    # Find the leftmost text element (with its midpoint and character length)
    leftmost_element = min(text_elements, key=lambda e: e["mid_x"] - (e["char_len"] * .8))
    min_x_start = leftmost_element["mid_x"] - (leftmost_element["char_len"] * .8)

    # Group elements by page
    pages = {}
    for e in text_elements:
        page_num = e.get("page_number", 1)
        if page_num not in pages:
            pages[page_num] = []
        pages[page_num].append(e)

    # Process each page separately and join with page separators
    all_pages_text = []

    for page_num in sorted(pages.keys()):
        page_elements = pages[page_num]

        # Sort by y-coordinate (rounded to nearest integer) then by x-coordinate
        page_elements.sort(key=lambda e: (e["mid_y"], e["mid_x"]))

        # Group by y-coordinate (lines)
        lines = {}
        for e in page_elements:
            y_key = e["mid_y"]
            if y_key not in lines:
                lines[y_key] = []
            lines[y_key].append(e)

        # Sort lines by y-coordinate
        sorted_y_keys = sorted(lines.keys())
        '''
        # Find the leftmost text element (with its midpoint and character length)
        leftmost_element = min(page_elements, key=lambda e: e["mid_x"] - e["char_len"] * .9)
        min_x_start = leftmost_element["mid_x"] - leftmost_element["char_len"] * .9
        '''
        # Prepare formatted text for this page
        page_lines = []

        # Process each line
        for idx, y_key in enumerate(sorted_y_keys):
            line_elements = lines[y_key]

            # Calculate vertical spacing from previous line (except for first line)
            if idx > 0:
                prev_y_key = sorted_y_keys[idx-1]

                # Get average height of elements in previous line
                prev_line = lines[prev_y_key]
                avg_prev_height = sum(e["height"] for e in prev_line) / len(prev_line)

                # Calculate normalized vertical space in number of newlines
                # Now with scaled coordinates, 1 newline represents ~2.5 original y units
                y_diff = y_key - prev_y_key
                newline_count = max(1, min(8, round(y_diff / avg_prev_height)))

                # Add appropriate number of newlines (capped at 8)
                page_lines.append("\n" * newline_count)

            # Sort elements in the line by x-coordinate
            line_elements.sort(key=lambda e: e["mid_x"])

            line_text = ""

            for i, e in enumerate(line_elements):
                text = e.get("text", "")

                # First element in the line
                if i == 0:
                    # Calculate spaces relative to the left margin
                    # With scaled coordinates, 1 space represents ~2.5 original x units
                    element_start = e["mid_x"] - (e["char_len"] * .8)
                    spaces = max(1, int(element_start - min_x_start))
                    line_text += " " * spaces + text
                else:
                    # Calculate spaces between elements
                    prev_e = line_elements[i-1]
                    prev_x_end = prev_e["mid_x"] + (prev_e["char_len"] * .8)
                    curr_x_start = e["mid_x"] - (e["char_len"] * .8)
                    spaces = max(1, int(curr_x_start - prev_x_end))
                    line_text += " " * spaces + text

            page_lines.append(line_text)

        # Join all lines for this page
        page_text = "".join(page_lines)

        # Add page number at the end of each page
        page_text += f"\n~Page|{page_num}"

        all_pages_text.append(page_text)

    # Join all pages with a standardized page separator (3 newlines)
    return "\n\n\n".join(all_pages_text)

def restructure_response(response_data):
    """
    Restructure the response by moving full_document elements up one level.
    Args:
        response_data: The original response data dictionary  
    Returns:
        dict: Restructured response
    """
    # Create a new dictionary with the desired structure
    restructured = {
        "document_metadata": response_data["document_metadata"],
        "elements_with_positions": response_data["elements_with_positions"],
        "images": response_data["images"],
        "links": response_data["links"],
        "form_fields": response_data["form_fields"],
        "annotations": response_data["annotations"],
        "text": response_data.get("text", ""),
        "pages": response_data["pages"]
    }

    return restructured







@app.route(route="document_intel_text_replica")
def document_intel_text_replica(req: func.HttpRequest) -> func.HttpResponse:
    """
    Azure Function that processes Logic Apps Analyze Document JSON and generates
    a spatially aware text replica of the PDF/image text layer.
    """
    logging.info('Processing spatial text replica request.')

    try:
        # Get JSON data from request
        req_body = req.get_json()
        if not req_body or 'pages' not in req_body:
            return func.HttpResponse(
                json.dumps({"error": "Invalid input. Expected JSON with 'pages' field."}),
                status_code=400,
                mimetype="application/json"
            )

        pages_data = req_body['pages']
        
        # Process each page and generate spatial text
        processed_pages = []
        for page in pages_data:
            processed_page = process_page(page)
            processed_pages.append(processed_page)
        
        # Generate the final spatial text replica
        spatial_text = generate_spatial_text_replica(processed_pages)
        
        return func.HttpResponse(
            spatial_text,
            status_code=200,
            mimetype="text/plain"
        )

    except Exception as e:
        logging.error(f"Error processing request: {str(e)}")
        return func.HttpResponse(
            json.dumps({"error": f"Processing error: {str(e)}"}),
            status_code=500,
            mimetype="application/json"
        )


def process_page(page):
    """
    Process a single page: detect rotation and transform coordinates to upright position.
    
    Args:
        page: Page data from Logic Apps Analyze Document
        
    Returns:
        dict: Processed page with upright-transformed words
    """
    words = page.get('words', [])
    page_number = page.get('pageNumber', 1)
    
    if not words:
        return {
            'pageNumber': page_number,
            'words': [],
            'rotation_applied': 0
        }
    
    # Transform polygon coordinates to 0.0-1.0 scale by dividing by 10
    # (assuming original coordinates are in 0-10 range based on the example)
    normalized_words = []
    for word in words:
        normalized_word = word.copy()
        if 'polygon' in word and len(word['polygon']) >= 8:
            # Normalize polygon coordinates
            polygon = word['polygon']
            normalized_polygon = [coord / 10.0 for coord in polygon]
            normalized_word['polygon'] = normalized_polygon
        normalized_words.append(normalized_word)
    
    # Detect page orientation by analyzing polygon coordinates
    rotation_needed = detect_page_rotation(normalized_words)
    
    # Transform coordinates based on detected rotation
    upright_words = transform_coordinates(normalized_words, rotation_needed)
    
    return {
        'pageNumber': page_number,
        'words': upright_words,
        'rotation_applied': rotation_needed
    }


def detect_page_rotation(words):
    """
    Detect the rotation needed to make the page upright by analyzing text bounding boxes.
    
    Args:
        words: List of word objects with normalized polygon coordinates
        
    Returns:
        int: Rotation angle needed (0, 90, 180, 270)
    """
    if not words:
        return 0
    
    rotation_counts = {0: 0, 90: 0, 180: 0, 270: 0}
    
    for word in words:
        polygon = word.get('polygon', [])
        if len(polygon) < 8:
            continue
            
        # Extract coordinates (assuming polygon format: [x1,y1,x2,y2,x3,y3,x4,y4])
        x1, y1, x2, y2, x3, y3, x4, y4 = polygon[:8]
        
        # Check different rotation scenarios based on coordinate relationships
        # Upright (0 degrees): top-left should be less than bottom-right
        if x1 < x3 and y1 < y3:
            rotation_counts[0] += 1
        
        # 180 degrees: coordinates are flipped
        elif x1 > x3 and y1 > y3:
            rotation_counts[180] += 1
        
        # 90 degrees: width and height are swapped, specific pattern
        elif x1 > x3 and y1 < y3:
            rotation_counts[90] += 1
        
        # 270 degrees: width and height are swapped, opposite pattern
        elif x1 < x3 and y1 > y3:
            rotation_counts[270] += 1
    
    # Return the rotation with the highest count
    return max(rotation_counts.items(), key=lambda x: x[1])[0]


def transform_coordinates(words, rotation):
    """
    Transform word coordinates based on the required rotation to make them upright.
    
    Args:
        words: List of word objects with normalized coordinates
        rotation: Rotation angle (0, 90, 180, 270)
        
    Returns:
        list: Words with transformed coordinates
    """
    if rotation == 0:
        return words
    
    transformed_words = []
    
    for word in words:
        transformed_word = word.copy()
        polygon = word.get('polygon', [])
        
        if len(polygon) >= 8:
            x1, y1, x2, y2, x3, y3, x4, y4 = polygon[:8]
            
            if rotation == 180:
                # 180-degree rotation: flip both x and y coordinates
                new_polygon = [
                    1 - x1, 1 - y1,
                    1 - x2, 1 - y2, 
                    1 - x3, 1 - y3,
                    1 - x4, 1 - y4
                ]
            
            elif rotation == 90:
                # 90-degree rotation: swap and transform coordinates
                new_polygon = [
                    y1, 1 - x1,
                    y2, 1 - x2,
                    y3, 1 - x3,
                    y4, 1 - x4
                ]
            
            elif rotation == 270:
                # 270-degree rotation: swap and transform coordinates differently
                new_polygon = [
                    1 - y1, x1,
                    1 - y2, x2,
                    1 - y3, x3,
                    1 - y4, x4
                ]
            
            transformed_word['polygon'] = new_polygon
        
        transformed_words.append(transformed_word)
    
    return transformed_words


def generate_spatial_text_replica(processed_pages):
    """
    Generate spatially aware text replica from processed pages.
    
    Args:
        processed_pages: List of processed page objects with upright words
        
    Returns:
        str: Spatially formatted text replica
    """
    if not processed_pages:
        return ""
    
    all_pages_text = []
    
    for page_data in processed_pages:
        page_number = page_data['pageNumber']
        words = page_data['words']
        
        if not words:
            all_pages_text.append(f"~Page|{page_number}")
            continue
        
        # Calculate midpoints and character lengths for spatial positioning
        text_elements = []
        for word in words:
            content = word.get('content', '').strip()
            if not content:
                continue
                
            polygon = word.get('polygon', [])
            if len(polygon) >= 8:
                # Calculate bounding box from polygon
                x_coords = [polygon[i] for i in range(0, 8, 2)]
                y_coords = [polygon[i] for i in range(1, 8, 2)]
                
                min_x, max_x = min(x_coords), max(x_coords)
                min_y, max_y = min(y_coords), max(y_coords)
                
                # Scale coordinates for text positioning
                mid_x = round((min_x + (max_x - min_x)/2) * 200)
                mid_y = round((min_y + (max_y - min_y)/2) * 110)
                
                text_elements.append({
                    'content': content,
                    'mid_x': mid_x,
                    'mid_y': mid_y,
                    'char_len': len(content),
                    'height': (max_y - min_y) * 110
                })
        
        if not text_elements:
            all_pages_text.append(f"~Page|{page_number}")
            continue
        
        # Generate spatial text for this page
        page_text = generate_page_spatial_text(text_elements)
        page_text += f"\n~Page|{page_number}"
        all_pages_text.append(page_text)
    
    # Join all pages with standardized separators
    return "\n\n\n".join(all_pages_text)


def generate_page_spatial_text(text_elements):
    """
    Generate spatially aware text for a single page.
    Args:
        text_elements: List of text elements with spatial coordinates
    Returns:
        str: Formatted text for the page
    """
    if not text_elements:
        return ""
    
    # Sort by y-coordinate first, then by x-coordinate
    text_elements.sort(key=lambda e: (e['mid_y'], e['mid_x']))
    
    # Find the leftmost starting position with better character width estimation
    leftmost_element = min(text_elements, key=lambda e: e['mid_x'] - (e['char_len'] * .75)) #char_len/2 * 2
    min_x_start = leftmost_element['mid_x'] - (leftmost_element['char_len'] * .75) #char_len/2 * 2
    
    # Group elements by y-coordinate (lines)
    lines = {}
    for element in text_elements:
        y_key = element['mid_y']
        if y_key not in lines:
            lines[y_key] = []
        lines[y_key].append(element)
    
    # Sort lines by y-coordinate
    sorted_y_keys = sorted(lines.keys())
    page_lines = []
    
    # Process each line
    for idx, y_key in enumerate(sorted_y_keys):
        line_elements = lines[y_key]
        
        # Calculate vertical spacing from previous line
        if idx > 0:
            prev_y_key = sorted_y_keys[idx-1]
            prev_line = lines[prev_y_key]
            avg_prev_height = sum(e['height'] for e in prev_line) / len(prev_line)
            
            # Calculate number of newlines needed
            y_diff = y_key - prev_y_key
            newline_count = max(1, min(6, round(y_diff / max(avg_prev_height, 1))))
            page_lines.append("\n" * newline_count)
        
        # Sort elements in the line by x-coordinate
        line_elements.sort(key=lambda e: e['mid_x'] - (element['char_len'] * .75)) #char_len/2 * 2
        
        line_text = ""
        for i, element in enumerate(line_elements):
            content = element['content']
            
            if i == 0:
                # First element: calculate absolute position from page left edge
                element_start = element['mid_x'] - (element['char_len'] * .75) #char_len/2 * 2
                # Use absolute positioning instead of relative to leftmost element
                spaces = max(0, int(element_start - min_x_start))  # Increase divisor to compress spacing
                line_text += " " * spaces + content
            else:
                # Calculate spaces between elements with improved character width estimation
                prev_element = line_elements[i-1]
                prev_x_end = prev_element['mid_x'] + (prev_element['char_len'] * .75) #char_len/2 * 2
                curr_x_start = element['mid_x'] - (element['char_len'] * .75) #char_len/2 * 2
                spaces = max(1, int(curr_x_start - prev_x_end)) # Increase divisor to compress spacing
                line_text += " " * spaces + content
        
        page_lines.append(line_text)
    
    return "".join(page_lines)












@app.route(route="merge_pdf_fitz")
def merge_pdf_fitz(req: func.HttpRequest) -> func.HttpResponse:
    try:
        # Form an array of base64 strings of the pdfs to merge from the $content parameters & call the merge_pdfs function on that array
        pdf_base64_strings_list = [item.get('$content') for item in req.get_json()['file_content']]
        merged_pdf_base64_string = merge_pdfs(pdf_base64_strings_list)

        ###HTTP response for MERGE operation###
        # Return the merged pdf base64 string in a Power Automate content object in an HTTP response
        return func.HttpResponse(
        body=json.dumps({
            "$content-type": "application/pdf",
            "$content": merged_pdf_base64_string
            }),
        mimetype="application/json",
        status_code=200
        )              

    except Exception as e:
        # If there is an error, log the error & then return the error message in an HTTP response
        debug_error = logging.exception(f"An error occurred: {str(e)}\n\nTraceback:\n{traceback.format_exc()}")
        return func.HttpResponse(
            f"Error: {str(e)}\n\nTraceback:\n{traceback.format_exc()}",
            status_code=500
        )

# Function used in each MERGE & SPLIT operation to get the PDF file bytes
def base64_to_pdf(base64_string):
    file_bytes = base64.b64decode(base64_string, validate=True)
    if file_bytes[0:4] != b"%PDF":
        raise ValueError("Missing the PDF file signature")
    return file_bytes

def merge_pdfs(pdf_base64_list):
    result = fitz.open()

    # First step: collect all form field values from all PDFs
    all_form_values = {}

    for pdf_base64 in pdf_base64_list:
        pdf_bytes = base64_to_pdf(pdf_base64)
        with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
            # Get form field values using the correct method name
            field_data = {}
            for page in doc:
                for widget in page.widgets():
                    if widget.field_name:
                        # Only store if this has a value
                        if widget.field_type == fitz.PDF_WIDGET_TYPE_RADIOBUTTON:
                            # For radio buttons, we need to check if it's selected
                            if hasattr(widget, 'field_flags') and (widget.field_flags & 2**15):
                                field_data[widget.field_name] = widget.field_value
                        else:
                            field_data[widget.field_name] = widget.field_value

            # Add to our collective field values
            all_form_values.update(field_data)

    # Second step: merge the PDFs with annotations preserved
    for pdf_base64 in pdf_base64_list:
        pdf_bytes = base64.b64decode(pdf_base64)
        with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
            # Use the same technique that works in your split function
            result.insert_pdf(doc, annots=True)

    # Final step: apply collected values to ensure consistent field values
    # Wait a moment before updating fields (sometimes helps with stability)
    processed_fields = set()

    for page in result:
        for widget in page.widgets():
            field_name = widget.field_name
            if field_name in all_form_values and field_name not in processed_fields:
                try:
                    if widget.field_type == fitz.PDF_WIDGET_TYPE_RADIOBUTTON:
                        # For radio buttons, handle the entire group together
                        radio_group_name = field_name

                        # Skip if we've already processed this group
                        if radio_group_name in processed_fields:
                            continue

                        # Get the target value to select
                        target_value = all_form_values.get(radio_group_name)
                        if not target_value:
                            continue  # Skip if no value to set

                        try:
                            # Try using the document-level API to set field values
                            # This works with radio buttons as groups rather than individual widgets
                            result.set_field_value(radio_group_name, target_value)
                        except AttributeError:
                            continue

                        processed_fields.add(radio_group_name)

                    else:
                        # For other field types
                        widget.field_value = all_form_values[field_name]
                        widget.update()

                    processed_fields.add(field_name)

                except Exception as e:
                    pass  # Skip if there's an issue

    # Save with settings that are compatible with your PyMuPDF version
    buffer = BytesIO()
    result.save(
        buffer, 
        garbage=0,  # No garbage collection to avoid removing form elements
        deflate=True, 
        clean=False  # Don't clean/remove any elements
    )
    buffer.seek(0)
    merged_pdf_bytes = buffer.read()

    return base64.b64encode(merged_pdf_bytes).decode("utf-8")






@app.route(route="split_pdf_fitz")
def split_pdf_fitz(req: func.HttpRequest) -> func.HttpResponse:
    try:
        req_json = req.get_json()
        pdf_bytes = base64_to_pdf(req_json['file_content'].get('$content'))
        page_numbers = req_json.get('pages')
        split_text = req_json.get('split_text')
        split_regex = req_json.get('split_regex')

        if page_numbers:
            split_base64_strings = split_pdf_by_page_numbers(pdf_bytes, page_numbers)
        elif split_text or split_regex:
            # Determine if PDF has a text layer
            if pdf_has_text_layer(pdf_bytes):
                split_base64_strings = split_pdf_by_text(pdf_bytes, split_text, split_regex)
            else:
                return func.HttpResponse("Text & regex methods do not work on PDFs without text layers. Use a different method or only use on PDFs with text layers.", status_code=400)
        else:
            return func.HttpResponse("Invalid. Must provide a 'pages' array to split by page, a 'split_text' to split by exact text, or a 'split_regex to split by text matching a regex expression.", status_code=400)

        response_data = [{"$content-type": "application/pdf", "$content": pdf} for pdf in split_base64_strings]

        return func.HttpResponse(
            body=json.dumps(response_data),
            mimetype="application/json",
            status_code=200
        )

    except Exception as e:
        logging.exception(f"An error occurred: {str(e)}\n\nTraceback:\n{traceback.format_exc()}")
        return func.HttpResponse(
            f"Error: {str(e)}\n\nTraceback:\n{traceback.format_exc()}",
            status_code=500
        )

def base64_to_pdf(base64_string):
    file_bytes = base64.b64decode(base64_string, validate=True)
    if file_bytes[0:4] != b"%PDF":
        raise ValueError("Missing the PDF file signature")
    return file_bytes

def pdf_has_text_layer(pdf_bytes: bytes) -> bool:
    """
    Check if the PDF contains a text layer.
    Returns True if text is found, False otherwise.
    """
    pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
    has_text = False

    for page_num in range(min(len(pdf_document), 3)):  # Check first 3 pages for efficiency
        page = pdf_document[page_num]
        text = page.get_text()
        if text.strip():
            has_text = True
            break

    pdf_document.close()
    return has_text

def pdf_has_form_fields(pdf_bytes: bytes) -> bool:
    """
    Check if the PDF contains any form fields throughout the entire document.
    Returns True if any form fields are found, False otherwise.
    """
    pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
    has_fields = False

    # Check all pages for form fields
    for page_num in range(len(pdf_document)):
        page = pdf_document[page_num]
        try:
            widgets = page.widgets()
            if len(widgets) > 0:
                has_fields = True
                break
        except Exception as e:
            logging.warning(f"Error checking for widgets on page {page_num}: {str(e)}")

    # Also check for AcroForm in the PDF catalog
    try:
        if "AcroForm" in pdf_document.get_pdf_catalog():
            has_fields = True
    except Exception as e:
        logging.warning(f"Error checking for AcroForm: {str(e)}")

    pdf_document.close()
    return has_fields

def get_form_fields_info(pdf_bytes):
    """
    Extract comprehensive form field information including all metadata.
    This function creates a more detailed mapping of fields to ensure proper preservation.
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    # Map fields to pages
    page_to_fields = defaultdict(set)
    field_to_pages = defaultdict(set)

    # Extract all field values and properties
    field_data = {}

    # Process each page
    for page_num in range(len(doc)):
        page = doc[page_num]

        # Get widgets from the page
        try:
            for widget in page.widgets():
                try:
                    field_name = widget.field_name
                    if field_name:
                        # Map field to page
                        page_to_fields[page_num].add(field_name)
                        field_to_pages[field_name].add(page_num)

                        # Store complete field data
                        if field_name not in field_data:
                            field_data[field_name] = {
                                'value': widget.field_value if hasattr(widget, 'field_value') else None,
                                'type': widget.field_type if hasattr(widget, 'field_type') else None,
                                'flags': widget.field_flags if hasattr(widget, 'field_flags') else None,
                                'rect': widget.rect if hasattr(widget, 'rect') else None,
                                'appearance': None,  # Will be populated if needed
                            }
                except Exception as e:
                    logging.warning(f"Error processing widget on page {page_num}: {str(e)}")
        except Exception as e:
            logging.warning(f"Error accessing widgets on page {page_num}: {str(e)}")

    doc.close()
    return page_to_fields, field_to_pages, field_data

def process_split_document(pdf_bytes, start_page, end_page):
    """
    Creates a new document from the specified page range and ensures form fields are preserved.
    Using a different approach to ensure consistent form field preservation across all splits.
    """
    # Open the source document
    source_doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    # Create a new document for the page range
    new_doc = fitz.open()

    # Insert pages with complete annotations (crucial for form fields)
    new_doc.insert_pdf(source_doc, from_page=start_page, to_page=end_page, annots=True)

    # Explicitly handle XFA forms if present
    try:
        if hasattr(source_doc, "xref_xml_metadata") and source_doc.xref_xml_metadata > 0:
            # Copy XFA form data if available
            if hasattr(new_doc, "set_xml_metadata") and hasattr(source_doc, "xml_metadata"):
                new_doc.set_xml_metadata(source_doc.xml_metadata)
    except Exception as e:
        logging.warning(f"Error handling XFA forms: {str(e)}")

    # Ensure AcroForm is preserved
    try:
        if "AcroForm" in source_doc.get_pdf_catalog():
            # If there's an AcroForm in the source, make sure we're preserving all form-related elements
            logging.info("AcroForm found in source document, ensuring preservation")
    except Exception as e:
        logging.warning(f"Error checking for AcroForm: {str(e)}")

    # Save the document with careful settings to preserve all form functionality
    try:
        buffer = BytesIO()
        # Use specific PDF settings that maximize form preservation
        new_doc.save(
            buffer, 
            garbage=0,  # No garbage collection to avoid removing form elements
            deflate=True, 
            clean=False,  # Don't clean/remove any elements
            encryption=False,
            permissions=int(
                fitz.PDF_PERM_ACCESSIBILITY |
                fitz.PDF_PERM_PRINT |
                fitz.PDF_PERM_COPY |
                fitz.PDF_PERM_ANNOTATE
            ),
            preserve_annots=True,  # IMPORTANT: Ensure annotations are preserved
            embedded_files=True    # Keep embedded files if any
        )
        buffer.seek(0)
        pdf_bytes = buffer.read()
    except Exception as e:
        logging.warning(f"Error saving with options: {str(e)}. Using fallback method.")
        try:
            # Fallback with different options
            buffer = BytesIO()
            new_doc.save(buffer, garbage=0, clean=False)
            buffer.seek(0)
            pdf_bytes = buffer.read()
        except Exception as e2:
            logging.warning(f"Fallback method also failed: {str(e2)}. Using tobytes.")
            pdf_bytes = new_doc.tobytes()

    # Close both documents to free resources
    new_doc.close()
    source_doc.close()

    return base64.b64encode(pdf_bytes).decode("utf-8")

def split_pdf_by_page_numbers(pdf_bytes, page_numbers):
    """
    Split PDF by page numbers, preserving form fields and their values.
    Returns a list of base64-encoded PDF documents.
    """
    # Check if the PDF has form fields (checking ALL pages)
    has_form_fields = pdf_has_form_fields(pdf_bytes)
    if has_form_fields:
        logging.info("PDF contains form fields - using form-preserving splitting")

    # Open the document to get total page count
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    total_pages = len(doc)
    doc.close()

    result_base64_strings = []

    # Ensure the page_numbers list starts with 1
    if not page_numbers or page_numbers[0] != 1:
        page_numbers = [1] + (page_numbers if page_numbers else [])

    # Add the total number of pages + 1 if the last split point is not the end
    # This ensures we include the last page in our calculations
    if page_numbers[-1] <= total_pages:
        page_numbers.append(total_pages + 1)

    # Process each page range
    for i in range(len(page_numbers) - 1):
        start_page = page_numbers[i] - 1  # Convert to 0-based index
        end_page = page_numbers[i+1] - 2  # Convert to 0-based index

        # For the last range, ensure we include the final page
        if i == len(page_numbers) - 2:
            end_page = total_pages - 1  # Make sure to include the last page

        # Skip invalid ranges
        if start_page > end_page or start_page < 0 or end_page >= total_pages:
            continue

        # Process the document for this page range
        base64_pdf = process_split_document(pdf_bytes, start_page, end_page)
        result_base64_strings.append(base64_pdf)

    return result_base64_strings

def split_pdf_by_text(pdf_bytes, split_text=None, split_regex=None):
    """
    Split PDF by exact text occurrence or regex match, preserving form fields and their values.
    Returns a list of base64-encoded PDF documents.
    """
    has_form_fields = pdf_has_form_fields(pdf_bytes)
    if has_form_fields:
        logging.info("PDF contains form fields - using form-preserving splitting")

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    total_pages = len(doc)
    split_pages = [0]  # Start with page 0

    for page_num in range(total_pages):
        page = doc[page_num]
        text = page.get_text()

        if split_regex:
            if re.search(split_regex, text) and page_num > 0:
                split_pages.append(page_num)
        elif split_text:
            if split_text in text and page_num > 0:
                split_pages.append(page_num)

    if split_pages[-1] != total_pages - 1:
        split_pages.append(total_pages)
    else:
        split_pages.append(total_pages)

    doc.close()

    result_base64_strings = []
    for i in range(len(split_pages) - 1):
        start_page = split_pages[i]
        end_page = split_pages[i + 1] - 1

        if end_page < start_page:
            continue

        base64_pdf = process_split_document(pdf_bytes, start_page, end_page)
        result_base64_strings.append(base64_pdf)

    return result_base64_strings






@app.route(route="rotate_pdf_pages")
def rotate_pdf_pages(req: func.HttpRequest) -> func.HttpResponse:
    try:
        req_body = req.get_json()

        file_content = req_body.get("file_content", {})
        pdf_base64 = file_content.get("$content")
        if not pdf_base64:
            return func.HttpResponse(json.dumps({"error": "No PDF content provided"}), status_code=400)

        rotations = req_body.get("rotations", [])
        # Check if it's a JSON array with objects having "page" and "rotate"
        if not((isinstance(rotations, list) and 
            all(isinstance(item, dict) for item in rotations) and 
            all("page" in item and "rotate" in item for item in rotations))):
            return func.HttpResponse(json.dumps({"error": "JSON array of rotation instructions was not provided or was incorrect. Correct example... 'rotations':[{'page':1, 'rotate':180},{'page':2, 'rotate':270}]"}), status_code=400)

        pdf_bytes = base64.b64decode(pdf_base64)

        corrected_pdf_bytes = rotate_pdf_pages(pdf_bytes, rotations)
        corrected_pdf_base64 = base64.b64encode(corrected_pdf_bytes).decode("utf-8")

        response_content = {
            "$content-type": "application/pdf",
            "$content": corrected_pdf_base64
        }

        return func.HttpResponse(json.dumps(response_content), mimetype="application/json")

    except Exception as e:
        return func.HttpResponse(
            f"Error: {str(e)}\n\nTraceback:\n{traceback.format_exc()}",
            status_code=500
        )

def rotate_pdf_pages(pdf_bytes, rotation_instructions):
    if not rotation_instructions:
        return pdf_bytes  # No rotation needed, return original

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    for page_info in rotation_instructions:
        if isinstance(page_info["rotate"], str):
            page_num = int(page_info["page"])
        else:
            page_num = page_info["page"]

        if isinstance(page_info["rotate"], str):
            rotate_amount = int(page_info["rotate"])
        else:
            rotate_amount = page_info["rotate"]

        page = doc[page_num - 1]
        current_rotation = page.rotation  # Get existing rotation
        new_rotation = (current_rotation + rotate_amount) % 360  # Apply rotation cumulatively

        page.set_rotation(new_rotation)

    pdf_bytes_out = BytesIO()
    doc.save(pdf_bytes_out)
    return pdf_bytes_out.getvalue()






@app.route(route="replace_text_pdf")
def replace_text_pdf(req: func.HttpRequest) -> func.HttpResponse:
    """
    Azure Function HTTP trigger to replace text in a PDF using PyMuPDF.
    Returns a file content object with the modified PDF.
    """
    try:
        # Get and validate the request body
        req_body = req.get_json()

        # Extract parameters with proper validation
        search_text = req_body.get("search_text")
        search_regex = req_body.get("search_regex")
        replace_text = req_body.get("replace_text", "")

        # Extract file content and validate
        file_content_obj = req_body.get("file_content", {})
        if isinstance(file_content_obj, dict):
            file_content = file_content_obj.get("$content")
        else:
            file_content = file_content_obj

        if not file_content:
            return func.HttpResponse("Missing or invalid file_content parameter", status_code=400)

        if not ((search_text and replace_text is not None) or (search_regex and replace_text is not None)):
            return func.HttpResponse("Must provide either search_text or search_regex with replace_text", status_code=400)

        # Decode the base64 content
        try:
            pdf_bytes = base64.b64decode(file_content)
        except Exception as e:
            return func.HttpResponse(f"Invalid base64 encoding: {str(e)}", status_code=400)

        # Check if the PDF contains a text layer
        if not pdf_has_text_layer(pdf_bytes):
            raise ValueError("The provided PDF does not contain a text layer and appears to be scanned images. This script only replaces text on PDFs with text layers.")

        # Perform text replacement
        try:
            modified_pdf = replace_text_in_pdf_improved(pdf_bytes, search_text, search_regex, replace_text)
        except Exception as e:
            logging.error("Error replacing text in PDF: %s", str(e))
            return func.HttpResponse(f"Error processing PDF content: {str(e)}", status_code=500)

        # Encode to base64
        output_pdf_base64 = base64.b64encode(modified_pdf).decode("utf-8")

        # Create the file content object
        response_body = json.dumps({
            "$content-type": "application/pdf",
            "$content": output_pdf_base64
        })

        # Return the modified PDF
        return func.HttpResponse(
            body=response_body,
            mimetype="application/json",
            status_code=200
        )

    except ValueError as e:
        return func.HttpResponse(f"Invalid JSON in request body: {str(e)}\n\nTraceback:\n{traceback.format_exc()}", status_code=400)
    except Exception as e:
        logging.error("Unexpected error: %s", str(e))
        return func.HttpResponse(
            f"Error: {str(e)}\n\nTraceback:\n{traceback.format_exc()}",
            status_code=500
        )

def pdf_has_text_layer(pdf_bytes: bytes) -> bool:
    """
    Check if the PDF contains a text layer.
    Returns True if text is found, False otherwise.
    """
    pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        text = page.get_text()
        if text.strip():
            return True
    return False

def replace_text_in_pdf_improved(pdf_bytes, search_text=None, search_regex=None, replace_text=""):
    """
    Improved method to replace text in a PDF using PyMuPDF's search_for and 
    redaction methods for more accurate text replacement.
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    try:
        for page_num in range(len(doc)):
            page = doc[page_num]

            if search_regex:
                # For regex searches, we need to get all text and process it differently
                page_text = page.get_text("dict")

                for block in page_text["blocks"]:
                    if "lines" not in block:
                        continue

                    for line in block["lines"]:
                        if "spans" not in line:
                            continue

                        for span in line["spans"]:
                            text = span["text"]

                            # Check if regex matches
                            if re.search(search_regex, text):
                                # For each match, create a list of matches
                                matches = list(re.finditer(search_regex, text))

                                for match in matches:
                                    # Get match position in the span text
                                    start, end = match.span()

                                    # Calculate character positions within the span
                                    char_width = (span["bbox"][2] - span["bbox"][0]) / len(text) if len(text) > 0 else 0

                                    # Create rectangle for just the matched part
                                    match_rect = fitz.Rect(
                                        span["bbox"][0] + start * char_width,
                                        span["bbox"][1],
                                        span["bbox"][0] + end * char_width,
                                        span["bbox"][3]
                                    )

                                    # Create replacement text (for regex, apply substitution)
                                    replacement = re.sub(search_regex, replace_text, match.group(0))

                                    # Add redaction annotation for this match
                                    redact_annot = page.add_redact_annot(match_rect, replacement)

                                    # Apply the redaction
                                    page.apply_redactions()

            # Determine what to search for
            elif search_text:
                # Direct text search
                text_instances = page.search_for(search_text)

                for inst in text_instances:
                    # Create a redaction annotation for the found text
                    redact_annot = page.add_redact_annot(inst, replace_text)

                    # Apply the redaction which removes the text
                    page.apply_redactions()

        # Save to memory buffer
        output_stream = BytesIO()
        doc.save(output_stream)
        return output_stream.getvalue()

    finally:
        doc.close()






@app.route(route="images_to_pdf")
def images_to_pdf(req: func.HttpRequest) -> func.HttpResponse: 
    try:
        req_body = req.get_json()
    except Exception as e:
        logging.error("Failed to get JSON body: %s", e)
        return func.HttpResponse("Invalid JSON", status_code=400)

    # Support both a raw list or a dict with an "images" key.
    if isinstance(req_body, dict) and "file_content" in req_body:
        images = req_body.get("file_content", [])
        stretch = req_body.get("fit_to_page", "NO").upper() == "YES"
        orientation_value = req_body.get("orientation") #LANDSCAPE or PORTRAIT
        pageOrientation = orientation_value.upper() if orientation_value else None
    elif isinstance(req_body, list):
        images = req_body
        stretch = False
        pageOrientation = None
    else:
        return func.HttpResponse("Invalid input format", status_code=400)

    if not images or not isinstance(images, list):
        return func.HttpResponse("No images provided", status_code=400)

    # Prepare an in-memory buffer for the PDF.
    pdf_buffer = BytesIO()
    temp_files = []

    try:
        c = canvas.Canvas(pdf_buffer)

        # Define A4 portrait and landscape sizes.
        A4_portrait = A4             # (595, 842) in points
        A4_landscape = (A4[1], A4[0])  # (842, 595)

        # Process each image and add a PDF page for it.
        for image_obj in images:
            try:
                # Each image object should have "$content-type" and "$content".
                content_type = image_obj.get("$content-type")
                image_data_base64 = image_obj.get("$content")
                if not image_data_base64:
                    logging.warning("Image object is missing '$content'")
                    continue

                # Decode the base64 string.
                image_data = base64.b64decode(image_data_base64)
                image_stream = BytesIO(image_data)

                # Open image with Pillow.
                with Image.open(image_stream) as img:
                    img.load()  # Ensure image is loaded
                    image_width, image_height = img.size

                    # Determine the page size based on orientation.
                    if pageOrientation is None:
                        # Auto-detect: if image is wider than tall, use landscape.
                        page_size = A4_landscape if image_width > image_height else A4_portrait
                    else:
                        # Use provided orientation.
                        if pageOrientation == "LANDSCAPE":
                            page_size = A4_landscape
                        else:
                            page_size = A4_portrait

                    # Set the page size for this page.
                    c.setPageSize(page_size)
                    page_width, page_height = page_size

                    # Create a temporary file for the image
                    fd, temp_path = tempfile.mkstemp(suffix='.png')
                    temp_files.append(temp_path)  # Track for cleanup
                    os.close(fd)

                    # Save the image to the temporary file
                    img.save(temp_path, format="PNG")

                    if stretch:
                        # Stretch image to fill entire page (may distort aspect ratio).
                        draw_width = page_width
                        draw_height = page_height
                        x = 0
                        y = 0
                    else:
                        # Calculate scale factor (only scale down if image is too large).
                        scale_factor = min(1, min(page_width / image_width, page_height / image_height))
                        draw_width = image_width * scale_factor
                        draw_height = image_height * scale_factor
                        # Center the image on the page.
                        x = (page_width - draw_width) / 2
                        y = (page_height - draw_height) / 2

                    # Draw the image onto the PDF page using the file path
                    c.drawImage(temp_path, x, y, width=draw_width, height=draw_height)
                    c.showPage()  # Finalize the current page

            except Exception as e:
                logging.error("Error processing an image: %s", e)
                return func.HttpResponse(f"Error generating PDF: {str(e)}\n\nTraceback:\n{traceback.format_exc()}", status_code=500)

        # Finalize the PDF.
        c.save()
        pdf_bytes = pdf_buffer.getvalue()

        # Encode the PDF as a base64 string.
        pdf_base64 = base64.b64encode(pdf_bytes).decode('utf-8')
        result = {
            "$content-type": "application/pdf",
            "$content": pdf_base64
        }
        return func.HttpResponse(json.dumps(result), mimetype="application/json")

    except Exception as e:
        logging.error("Error generating PDF: %s", e)
        return func.HttpResponse(f"Error generating PDF: {str(e)}", status_code=500)
    finally:
        pdf_buffer.close()
        # Clean up temporary files
        for temp_file in temp_files:
            try:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
            except Exception as e:
                logging.warning(f"Failed to remove temporary file {temp_file}: {e}")







@app.route(route="pdf_to_images")
def pdf_to_images(req: func.HttpRequest) -> func.HttpResponse:
    if req.method == 'POST':
        try:
            # Get base64 string from HTTP call
            pdf_base64_string = req.get_json()['file_content']['$content']
            # Run the function with the base64 string
            image_base64_strings = pdf_to_image_base64(pdf_base64_string)
            response = [{"$content-type": "image/png", "$content": image_base64} for image_base64 in image_base64_strings]

            return func.HttpResponse(
            body=json.dumps(response),
            mimetype="application/json",
            status_code=200
    )

        except Exception as e:
            # If there is an error, log the error & then return the error message in an HTTP response
            debug_error = logging.exception(f"An error occurred: {str(e)}\n\nTraceback:\n{traceback.format_exc()}")
            return func.HttpResponse(
            f"Error: {str(e)}\n\nTraceback:\n{traceback.format_exc()}",
            status_code=500
            )

def pdf_to_image_base64(pdf_base64_str):
                # Decode the base64 PDF string
                pdf_bytes = base64.b64decode(pdf_base64_str)

                # Read the PDF using PyMuPDF
                doc = fitz.open(stream=pdf_bytes)

                # Initialize an empty list to store image base64 strings
                image_base64_list = []

                # Iterate through each page in the PDF
                for page_num in range(doc.page_count):
                    page = doc.load_page(page_num)
                    pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))  # Adjust zoom as needed

                    # Convert the image data to a numpy array
                    image_np = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.h, pix.w, pix.n)
                    image_np = np.ascontiguousarray(image_np[..., [2, 1, 0]])  # Convert BGR to RGB

                    # Encode the image data as base64
                    image_bytes = BytesIO()
                    Image.fromarray(image_np).save(image_bytes, format="PNG")  # You can choose other formats
                    image_base64 = base64.b64encode(image_bytes.getvalue()).decode("ascii")
                    image_base64_list.append(image_base64)

                return image_base64_list






@app.route(route="pdf_to_html", auth_level=func.AuthLevel.FUNCTION)
def pdf_to_html(req: func.HttpRequest) -> func.HttpResponse:
    try:
        req_body = req.get_json()
        file_content = req_body.get("file_content", {})

        if "$content" not in file_content:
            return func.HttpResponse("Missing '$content' in file_content.", status_code=400)

        pdf_base64 = file_content["$content"]
        pdf_bytes = base64.b64decode(pdf_base64)

        # Check if the PDF contains a text layer
        if not pdf_has_text_layer(pdf_bytes):
            raise ValueError("The provided PDF does not contain a text layer and appears to be scanned images. This script only converts PDFs with text layers.")

        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=True) as temp_file:
            temp_file.write(pdf_bytes)
            temp_file.flush()

            try:
                doc = fitz.open(temp_file.name)
                html_output = []

                # CSS optimized for PDF conversion tools
                css = """
                <style>
                    @page {
                        margin: 0;
                        size: auto;
                    }
                    body {
                        margin: 0;
                        padding: 0;
                    }
                    .pdf-container {
                        position: relative;
                        width: 100%;
                    }
                    .pdf-page {
                        position: relative;
                        width: 100%;
                        page-break-after: always;
                        transform-origin: top left;
                        overflow: visible;
                    }
                    .pdf-element {
                        position: absolute;
                        margin: 0;
                        padding: 0;
                        line-height: 1;
                        white-space: pre;
                        transform-origin: top left;
                    }
                    .pdf-image {
                        position: absolute;
                        margin: 0;
                        padding: 0;
                        transform-origin: top left;
                    }
                </style>
                """

                html_output.append(f"""<!DOCTYPE html>
                <html>
                <head>
                    <meta charset="utf-8">
                    <meta name="viewport" content="width=device-width, initial-scale=1.0">
                    <title>PDF Converted to HTML</title>
                    {css}
                </head>
                <body>
                <div class="pdf-container">
                """)

                for page_num, page in enumerate(doc):
                    # Get physical page dimensions
                    media_box = page.mediabox
                    page_width = media_box.width
                    page_height = media_box.height

                    # Create a page container that preserves PDF dimensions exactly
                    html_output.append(f"""
                    <div class="pdf-page" data-page-number="{page_num + 1}" 
                         style="width: {page_width}pt; height: {page_height}pt;">
                    """)

                    # Extract text blocks with their positions
                    text_blocks = page.get_text("dict")["blocks"]

                    for block in text_blocks:
                        if "lines" in block:  # Text block
                            for line in block["lines"]:
                                for span in line["spans"]:
                                    text = span["text"]
                                    if text.strip():
                                        # Get position and styling
                                        x0, y0, x1, y1 = span["bbox"]
                                        font_size = span["size"]
                                        font_name = span["font"]
                                        color = f"#{span['color']:06x}" if "color" in span else "#000000"

                                        # Use pt units to match PDF measurements
                                        style = (
                                            f"left: {x0}pt; top: {y0}pt; "
                                            f"font-size: {font_size}pt; "
                                            f"color: {color}; "
                                            f"width: {x1 - x0}pt; height: {y1 - y0}pt;"
                                        )

                                        html_output.append(
                                            f'<div class="pdf-element" style="{style}">{text}</div>'
                                        )

                    # Extract images with their positions
                    for img_index, img in enumerate(page.get_images(full=True)):
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        image_ext = base_image["ext"]
                        image_b64 = base64.b64encode(image_bytes).decode('utf-8')

                        # Try to find image position on the page
                        img_rect = None
                        for block in text_blocks:
                            if block.get("type") == 1:  # Image block
                                img_rect = block["bbox"]
                                break

                        if not img_rect:
                            # Fall back to image extraction method
                            pix = fitz.Pixmap(doc, xref)
                            # Place near top left if we can't determine position
                            img_rect = [20, 20, 20 + pix.width, 20 + pix.height]

                        x0, y0, x1, y1 = img_rect
                        width = x1 - x0
                        height = y1 - y0

                        # Use pt units for image positioning
                        style = f"left: {x0}pt; top: {y0}pt; width: {width}pt; height: {height}pt;"

                        html_output.append(
                            f'<img class="pdf-image" src="data:image/{image_ext};base64,{image_b64}" '
                            f'alt="Image from page {page_num + 1}" style="{style}">'
                        )

                    html_output.append("</div>")

                html_output.append("</div></body></html>")

                # Add a script to help ensure accurate scaling
                html_output.append("""
                <script>
                document.addEventListener('DOMContentLoaded', function() {
                    // This helps some HTML-to-PDF converters properly scale the content
                    const pages = document.querySelectorAll('.pdf-page');
                    pages.forEach(page => {
                        const width = parseFloat(getComputedStyle(page).width);
                        const height = parseFloat(getComputedStyle(page).height);
                        page.style.minWidth = width + 'px';
                        page.style.minHeight = height + 'px';
                    });
                });
                </script>
                """)

            finally:
                if 'doc' in locals():
                    doc.close()

        return func.HttpResponse("".join(html_output), mimetype="text/html")

    except Exception as e:
        error_trace = traceback.format_exc()
        logging.error(f"Error processing PDF: {error_trace}")
        return func.HttpResponse(f"Error processing PDF:\n{error_trace}", status_code=500)

def pdf_has_text_layer(pdf_bytes: bytes) -> bool:
    """
    Check if the PDF contains a text layer.
    Returns True if text is found, False otherwise.
    """
    pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        text = page.get_text()
        if text.strip():
            return True
    return False           











@app.route(route="msg_to_html")
def msg_to_html(req: func.HttpRequest) -> func.HttpResponse:
    try:
        # Parse the request JSON
        req_body = req.get_json()
        
        if not req_body or 'file_content' not in req_body:
            return func.HttpResponse(
                "Error: Missing 'file_content' parameter",
                status_code=400
            )
        
        file_content = req_body['file_content']
        
        # Validate the content structure
        if not isinstance(file_content, dict) or '$content' not in file_content:
            return func.HttpResponse(
                "Error: Invalid file_content format. Expected {'$content-type': '...', '$content': '...'}",
                status_code=400
            )
        
        # Extract and decode the base64 content
        base64_content = file_content['$content']
        
        # Decode the base64 content
        try:
            msg_bytes = base64.b64decode(base64_content)
        except Exception as decode_error:
            return func.HttpResponse(
                f"Error: Failed to decode base64 content: {str(decode_error)}",
                status_code=400
            )
        
        # Convert MSG to structured data
        result = convert_msg_to_data(msg_bytes)

        # Return the result as JSON
        return func.HttpResponse(
            json.dumps(result, indent=2),
            status_code=200,
            headers={'Content-Type': 'application/json'}
        )
        
    except Exception as e:
        # If there is an error, log the error & then return the error message in an HTTP response
        debug_error = logging.exception(f"An error occurred: {str(e)}\n\nTraceback:\n{traceback.format_exc()}")
        return func.HttpResponse(
            f"Error: {str(e)}\n\nTraceback:\n{traceback.format_exc()}",
            status_code=500
        )

def convert_msg_to_data(msg_bytes):
    """
    Convert MSG file bytes to structured data with both HTML and plain text
    """
    try:
        # Parse MSG file using extract-msg
        msg_data = parse_msg_with_extract_msg(msg_bytes)
        
        # Validate that we extracted meaningful content
        if not any([
            msg_data.get('subject', '').strip(),
            msg_data.get('sender_name', '').strip(),
            msg_data.get('sender_email', '').strip(),
            msg_data.get('body', '').strip(),
            msg_data.get('body_html', '').strip()
        ]):
            raise Exception("MSG file parsing failed - no meaningful content could be extracted")
        
        # Generate both HTML and plain text outputs
        html_content = generate_simple_html(msg_data).replace('\n','')
        plain_text = generate_plain_text_with_csv_tables(msg_data)
        
        return {
            "html": html_content,
            "plain_text": plain_text,
            "metadata": {
                "subject": msg_data.get('subject', ''),
                "sender_name": msg_data.get('sender_name', ''),
                "sender_email": msg_data.get('sender_email', ''),
                "recipients": msg_data.get('recipients', []),
                "cc_recipients": msg_data.get('cc_recipients', []),
                "bcc_recipients": msg_data.get('bcc_recipients', []),
                "sent_time": msg_data.get('sent_time', ''),
                "received_time": msg_data.get('received_time', ''),
                "attachments": msg_data.get('attachments', [])
            }
        }
        
    except Exception as e:
        # Re-raise the exception to fail the function
        raise Exception(f"MSG file could not be parsed: {str(e)}")

def convert_to_string(value):
    """
    Convert bytes or other types to string safely
    """
    if isinstance(value, bytes):
        try:
            return value.decode('utf-8', errors='ignore')
        except:
            return str(value, errors='ignore')
    elif value is None:
        return ''
    else:
        return str(value)

def parse_msg_with_extract_msg(msg_bytes):
    """
    Parse MSG file using the extract-msg library
    """
    msg_data = {
        'subject': '',
        'sender_name': '',
        'sender_email': '',
        'recipients': [],
        'cc_recipients': [],
        'bcc_recipients': [],
        'sent_time': '',
        'received_time': '',
        'body': '',
        'body_html': '',
        'attachments': []
    }
    
    try:
        # Create a BytesIO object from the msg_bytes
        msg_stream = BytesIO(msg_bytes)
        
        # Parse the MSG file
        with ExtractMessage(msg_stream) as msg:
            # Extract basic properties - handle bytes to string conversion
            subject_raw = getattr(msg, 'subject', '') or ''
            sender_name_raw = getattr(msg, 'sender', '') or ''
            sender_email_raw = getattr(msg, 'senderEmail', '') or ''
            
            # Convert bytes to string if necessary
            msg_data['subject'] = convert_to_string(subject_raw)
            msg_data['sender_name'] = convert_to_string(sender_name_raw)
            msg_data['sender_email'] = convert_to_string(sender_email_raw)
            
            # Handle recipients - ensure strings not bytes
            to_recipients = []
            cc_recipients = []
            bcc_recipients = []
            
            if hasattr(msg, 'recipients') and msg.recipients:
                for recipient in msg.recipients:
                    recipient_type = getattr(recipient, 'type', None)
                    
                    # Convert bytes to string if necessary
                    recipient_email = convert_to_string(getattr(recipient, 'email', '') or getattr(recipient, 'name', '') or '')
                    
                    if recipient_email:
                        if recipient_type == 'CC':
                            cc_recipients.append(recipient_email)
                        elif recipient_type == 'BCC':
                            bcc_recipients.append(recipient_email)
                        else:  # TO or default
                            to_recipients.append(recipient_email)
            
            msg_data['recipients'] = to_recipients
            msg_data['cc_recipients'] = cc_recipients
            msg_data['bcc_recipients'] = bcc_recipients
            
            # Extract dates
            if hasattr(msg, 'date') and msg.date:
                msg_data['sent_time'] = str(msg.date)
            
            # Try multiple ways to get the received time
            if hasattr(msg, 'receivedTime') and msg.receivedTime:
                msg_data['received_time'] = str(msg.receivedTime)
            elif hasattr(msg, 'creationTime') and msg.creationTime:
                msg_data['received_time'] = str(msg.creationTime)
            
            # Extract body content - handle both string and bytes
            body_raw = getattr(msg, 'body', '') or ''
            body_html_raw = getattr(msg, 'htmlBody', '') or ''
            
            # Convert bytes to string if necessary
            msg_data['body'] = convert_to_string(body_raw)
            msg_data['body_html'] = convert_to_string(body_html_raw)
            
            # If no plain text body but we have HTML, try to extract text
            if not msg_data['body'] and msg_data['body_html']:
                try:
                    # Simple HTML to text conversion
                    text = re.sub(r'<[^>]+>', '', msg_data['body_html'])
                    text = re.sub(r'\s+', ' ', text)  # Normalize whitespace
                    msg_data['body'] = text.strip()
                except:
                    pass
            
            # Extract attachment information - handle bytes
            if hasattr(msg, 'attachments') and msg.attachments:
                for attachment in msg.attachments:
                    filename_raw = getattr(attachment, 'longFilename', '') or getattr(attachment, 'shortFilename', '') or 'Unknown'
                    attachment_info = {
                        'filename': convert_to_string(filename_raw),
                        'size': getattr(attachment, 'size', 0),
                        'type': convert_to_string(getattr(attachment, 'type', 'Unknown'))
                    }
                    msg_data['attachments'].append(attachment_info)
            
            logging.info(f"Successfully parsed MSG file using extract-msg library")
            
    except Exception as e:
        logging.error(f"Error parsing MSG with extract-msg: {str(e)}")
        logging.error(f"Traceback: {traceback.format_exc()}")
        raise Exception(f"extract-msg library failed to parse MSG file: {str(e)}")
    
    return msg_data

def generate_simple_html(msg_data):
    """
    Generate simple HTML output focused on metadata and content
    """
    html_template = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Email Message</title>
    <style>
        body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 20px auto; padding: 20px; }}
        .email-header {{ background-color: #f8f9fa; padding: 15px; border-left: 4px solid #007bff; margin-bottom: 20px; }}
        .email-header h2 {{ margin: 0 0 10px 0; color: #333; }}
        .meta-item {{ margin: 5px 0; }}
        .meta-label {{ font-weight: bold; color: #555; }}
        .email-content {{ border: 1px solid #e0e0e0; padding: 15px; background-color: #fefefe; }}
        .recipients {{ color: #007bff; }}
    </style>
</head>
<body>
    <div class="email-header">
        <h2>{subject}</h2>
        <div class="meta-item"><span class="meta-label">From:</span> {sender}</div>
        <div class="meta-item"><span class="meta-label">To:</span> <span class="recipients">{to_recipients}</span></div>
        {cc_section}
        {bcc_section}
        <div class="meta-item"><span class="meta-label">Sent:</span> {sent_time}</div>
        <div class="meta-item"><span class="meta-label">Received:</span> {received_time}</div>
    </div>
    
    <div class="email-content">
        {content}
    </div>
</body>
</html>
"""
    
    # Format sender
    sender = ""
    if msg_data.get('sender_name', '').strip():
        sender = html.escape(msg_data['sender_name'])
        if msg_data.get('sender_email', '').strip():
            sender += f" &lt;{html.escape(msg_data['sender_email'])}&gt;"
    elif msg_data.get('sender_email', '').strip():
        sender = html.escape(msg_data['sender_email'])
    else:
        sender = "Unknown"
    
    # Format recipients
    to_recipients = ", ".join(msg_data.get('recipients', [])) or "Unknown"
    
    # CC section
    cc_section = ""
    if msg_data.get('cc_recipients'):
        cc_list = ", ".join(msg_data['cc_recipients'])
        cc_section = f'<div class="meta-item"><span class="meta-label">CC:</span> <span class="recipients">{html.escape(cc_list)}</span></div>'
    
    # BCC section
    bcc_section = ""
    if msg_data.get('bcc_recipients'):
        bcc_list = ", ".join(msg_data['bcc_recipients'])
        bcc_section = f'<div class="meta-item"><span class="meta-label">BCC:</span> <span class="recipients">{html.escape(bcc_list)}</span></div>'
    
    # Content - prefer HTML if available, otherwise plain text
    content = ""
    if msg_data.get('body_html', '').strip():
        # Use HTML content directly (already formatted) but remove literal \n characters
        content = msg_data['body_html']
    elif msg_data.get('body', '').strip():
        # Convert plain text to HTML (preserve line breaks)
        plain_text = html.escape(msg_data['body'])
        content = plain_text.replace('\n', '<br>\n')
    else:
        content = "<em>No message content available</em>"
    
    return html_template.format(
        subject=html.escape(msg_data.get('subject', 'No Subject')),
        sender=sender,
        to_recipients=html.escape(to_recipients),
        cc_section=cc_section,
        bcc_section=bcc_section,
        sent_time=html.escape(msg_data.get('sent_time', 'Unknown')),
        received_time=html.escape(msg_data.get('received_time', 'Unknown')),
        content=content
    )

def html_decode_and_format(text):
    """
    Decode HTML entities and format text for readability
    """
    if not text:
        return ""
    
    # First decode HTML entities
    decoded = html.unescape(text)
    
    # Replace literal \n characters with actual newlines
    decoded = decoded.replace('\\n', '\n')
    
    # Handle common HTML entities that might be missed
    decoded = decoded.replace('&nbsp;', ' ')
    decoded = decoded.replace('&lt;', '<')
    decoded = decoded.replace('&gt;', '>')
    decoded = decoded.replace('&amp;', '&')
    decoded = decoded.replace('&quot;', '"')
    decoded = decoded.replace('&#39;', "'")
    
    return decoded

def generate_plain_text_with_csv_tables(msg_data):
    """
    Generate plain text with HTML tables converted to CSV format
    """
    # Start with metadata
    lines = []
    lines.append(f"Subject: {msg_data.get('subject', 'No Subject')}")
    
    # Sender
    sender = ""
    if msg_data.get('sender_name', '').strip():
        sender = msg_data['sender_name']
        if msg_data.get('sender_email', '').strip():
            sender += f" <{msg_data['sender_email']}>"
    elif msg_data.get('sender_email', '').strip():
        sender = msg_data['sender_email']
    else:
        sender = "Unknown"
    lines.append(f"From: {sender}")
    
    # Recipients
    if msg_data.get('recipients'):
        lines.append(f"To: {', '.join(msg_data['recipients'])}")
    if msg_data.get('cc_recipients'):
        lines.append(f"CC: {', '.join(msg_data['cc_recipients'])}")
    if msg_data.get('bcc_recipients'):
        lines.append(f"BCC: {', '.join(msg_data['bcc_recipients'])}")
    
    # Dates
    if msg_data.get('sent_time'):
        lines.append(f"Sent: {msg_data['sent_time']}")
    if msg_data.get('received_time'):
        lines.append(f"Received: {msg_data['received_time']}")
    
    lines.append("")  # Empty line before content
    lines.append("--- MESSAGE CONTENT ---")
    lines.append("")
    
    # Process content - convert HTML tables to CSV
    content = ""
    if msg_data.get('body_html', '').strip():
        content = convert_html_tables_to_csv(msg_data['body_html'])
    elif msg_data.get('body', '').strip():
        content = html_decode_and_format(msg_data['body'])
    else:
        content = "No message content available"
    
    lines.append(content)
    
    # Add attachments info if any
    if msg_data.get('attachments'):
        lines.append("")
        lines.append("--- ATTACHMENTS ---")
        for attachment in msg_data['attachments']:
            size_kb = attachment['size'] / 1024 if attachment['size'] > 0 else 0
            size_str = f" ({size_kb:.1f} KB)" if size_kb > 0 else ""
            lines.append(f"- {attachment['filename']}{size_str}")
    
    return '\n'.join(lines)

def convert_html_tables_to_csv(html_content):
    """
    Convert HTML tables in content to CSV format while preserving other content
    """
    if not html_content:
        return ""
    
    # Find all tables in the HTML
    table_pattern = r'<table[^>]*>(.*?)</table>'
    tables = re.findall(table_pattern, html_content, re.DOTALL | re.IGNORECASE)
    
    if not tables:
        # No tables found, just clean up HTML and return formatted text
        return clean_html_to_text(html_content)
    
    # Process the HTML content piece by piece
    result = html_content
    
    for table_html in tables:
        csv_content = convert_single_table_to_csv(table_html)
        # Replace the table with CSV content
        table_pattern_replace = r'<table[^>]*>.*?</table>'
        result = re.sub(table_pattern_replace, f'\n\n[TABLE]\n{csv_content}\n[/TABLE]\n\n', result, count=1, flags=re.DOTALL | re.IGNORECASE)
    
    # Clean up remaining HTML
    result = clean_html_to_text(result)
    
    return result

def clean_html_to_text(html_content):
    """
    Convert HTML to clean, readable plain text
    """
    if not html_content:
        return ""
    
    text = html_content
    
    # Replace common HTML elements with appropriate text formatting
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'</?p[^>]*>', '\n\n', text, flags=re.IGNORECASE)
    text = re.sub(r'</?div[^>]*>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'<h[1-6][^>]*>(.*?)</h[1-6]>', r'\n\n\1\n\n', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'<li[^>]*>(.*?)</li>', r'\n• \1', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'</?ul[^>]*>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'</?ol[^>]*>', '\n', text, flags=re.IGNORECASE)
    
    # Remove all remaining HTML tags
    text = re.sub(r'<[^>]+>', '', text)
    
    # Decode HTML entities and clean up formatting
    text = html_decode_and_format(text)
    
    # Clean up whitespace
    text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs to single space
    text = re.sub(r'\n[ \t]+', '\n', text)  # Remove leading whitespace on lines
    text = re.sub(r'[ \t]+\n', '\n', text)  # Remove trailing whitespace on lines
    text = re.sub(r'\n{3,}', '\n\n', text)  # Multiple newlines to double newlines max
    
    return text.strip()

def convert_single_table_to_csv(table_html):
    """
    Convert a single HTML table to CSV format
    """
    if not table_html:
        return ""
    
    # Extract rows
    row_pattern = r'<tr[^>]*>(.*?)</tr>'
    rows = re.findall(row_pattern, table_html, re.DOTALL | re.IGNORECASE)
    
    if not rows:
        return ""
    
    csv_rows = []
    
    for row_html in rows:
        # Extract cells (both th and td)
        cell_pattern = r'<(?:th|td)[^>]*>(.*?)</(?:th|td)>'
        cells = re.findall(cell_pattern, row_html, re.DOTALL | re.IGNORECASE)
        
        if cells:
            # Clean up cell content
            clean_cells = []
            for cell in cells:
                # Remove HTML tags from cell content
                clean_cell = re.sub(r'<[^>]+>', '', cell)
                # Decode HTML entities
                clean_cell = html_decode_and_format(clean_cell)
                # Normalize whitespace but preserve some structure
                clean_cell = re.sub(r'\s+', ' ', clean_cell).strip()
                # Escape commas and quotes for CSV
                if ',' in clean_cell or '"' in clean_cell or '\n' in clean_cell:
                    clean_cell = '"' + clean_cell.replace('"', '""') + '"'
                clean_cells.append(clean_cell)
            
            csv_rows.append(','.join(clean_cells))
    
    if csv_rows:
        return '\n'.join(csv_rows)
    else:
        return ""









@app.route(route="pdf_to_word")
def pdf_to_word(req: func.HttpRequest) -> func.HttpResponse:
    try:
        # Parse JSON body
        req_body = req.get_json()
        file_content = req_body.get('file_content')
        if not file_content:
            raise ValueError("No 'file_content' provided in the request.")

        # Decode the base64-encoded PDF content
        pdf_bytes = base64.b64decode(file_content['$content'])

        # Check if the PDF contains a text layer
        if not pdf_has_text_layer(pdf_bytes):
            raise ValueError("The provided PDF does not contain a text layer and appears to be scanned images. This script only converts PDFs with text layers.")

        # Convert PDF to DOCX
        docx_bytes = convert_pdf_to_docx(pdf_bytes)

        # Encode the DOCX bytes to base64
        docx_base64 = base64.b64encode(docx_bytes).decode('utf-8')

        # Create the response content
        response_content = {
            "$content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "$content": docx_base64
        }

        # Return the response as JSON
        return func.HttpResponse(
            body=json.dumps(response_content),
            status_code=200,
            mimetype="application/json"
        )

    except ValueError as ve:
        logging.error(f"ValueError: {str(ve)}\n\nTraceback:\n{traceback.format_exc()}")
        return func.HttpResponse(
            str(ve),
            status_code=400
        )
    except Exception as e:
        logging.error(f"Exception: {str(e)}\n\nTraceback:\n{traceback.format_exc()}")
        return func.HttpResponse(
            f"Error: {str(e)}\n\nTraceback:\n{traceback.format_exc()}",
            status_code=500
        )

def pdf_has_text_layer(pdf_bytes: bytes) -> bool:
    """
    Check if the PDF contains a text layer.
    Returns True if text is found, False otherwise.
    """
    pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        text = page.get_text()
        if text.strip():
            return True
    return False

def convert_pdf_to_docx(pdf_bytes: bytes) -> bytes:
    """
    Convert PDF bytes to DOCX bytes using temporary files.
    """
    # Create a temporary file for the PDF
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as pdf_temp_file:
        pdf_temp_file.write(pdf_bytes)
        pdf_temp_file_path = pdf_temp_file.name

    # Create a temporary file for the DOCX
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as docx_temp_file:
        docx_temp_file_path = docx_temp_file.name

    try:
        # Convert PDF to DOCX
        converter = Converter(pdf_temp_file_path)
        converter.convert(docx_temp_file_path)
        converter.close()

        # Read the DOCX file into bytes
        with open(docx_temp_file_path, 'rb') as docx_file:
            docx_bytes = docx_file.read()

    finally:
        # Clean up temporary files
        os.remove(pdf_temp_file_path)
        os.remove(docx_temp_file_path)

    return docx_bytes






@app.route(route="extract_word_elements")
def extract_word_elements(req: func.HttpRequest) -> func.HttpResponse:
    try:
        # Get request body
        req_body = req.get_json()['file_content']

        # Check if the required fields are in the request
        if not req_body or '$content-type' not in req_body or '$content' not in req_body:
            return func.HttpResponse(
                json.dumps({"error": "Invalid request format. Expected 'file_content': {'$content-type': '...', '$content': '...'}"}),
                mimetype="application/json",
                status_code=400
            )

        # Verify content type
        content_type = req_body['$content-type']
        if content_type != "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return func.HttpResponse(
                json.dumps({"error": f"Unsupported content type: {content_type}. Only Word documents are supported."}),
                mimetype="application/json",
                status_code=400
            )

        # Decode base64 content
        base64_content = req_body['$content']
        try:
            binary_content = base64.b64decode(base64_content)
        except Exception as e:
            return func.HttpResponse(
                json.dumps({"error": f"Failed to decode base64 content: {str(e)}"}),
                mimetype="application/json",
                status_code=400
            )

        # Process Word document
        result = process_word_document(binary_content)

        return func.HttpResponse(
            json.dumps(result, indent=2),
            mimetype="application/json"
        )

    except Exception as e:
        error_trace = traceback.format_exc()
        logging.error(f"Error processing request: {str(e)}\n{error_trace}")
        return func.HttpResponse(
            json.dumps({
                "error": str(e),
                "traceback": error_trace
            }),
            mimetype="application/json",
            status_code=500
        )

def process_word_document(binary_content):
    """Process Word document and extract various information"""
    result = {}

    # Create file-like object from binary content
    docx_file = BytesIO(binary_content)

    # Check if document is password protected
    is_protected = check_password_protection(docx_file)
    result["is_password_protected"] = is_protected

    # Extract XML content
    xml_content = extract_xml_content(docx_file)
    result["xml_content"] = xml_content

    # Reset file pointer for python-docx processing
    docx_file.seek(0)

    try:
        # Use python-docx to extract additional information
        doc = Document(docx_file)

        # Get metadata
        result["metadata"] = extract_metadata(doc, xml_content)

        # Get document statistics
        result["statistics"] = extract_statistics(doc, xml_content)

        # Get page orientation
        result["section_orientation"] = extract_section_orientation(xml_content)

        # Get document text
        result["document_text"] = extract_document_text(doc)

        # Get tables
        result["tables"] = extract_tables(doc)

        # Get images
        result["images"] = extract_images(docx_file)

        # Get hyperlinks
        result["hyperlinks"] = extract_hyperlinks(doc, xml_content)

        # Get table of contents
        result["table_of_contents"] = extract_table_of_contents(doc, xml_content)

        # Get headings
        result["headings"] = extract_headings(doc)

        # Get comments
        result["comments"] = extract_comments(xml_content)

        # Get digital signatures
        result["digital_signatures"] = extract_digital_signatures(xml_content)

    except Exception as e:
        error_trace = traceback.format_exc()
        logging.error(f"Error processing document: {str(e)}\n{error_trace}")
        result["error"] = str(e)
        result["traceback"] = error_trace

    return result

def check_password_protection(docx_file):
    """Check if the document is password protected"""
    try:
        with ZipFile(docx_file) as zip_ref:
            # If we can read the contents, it's not password protected at the ZIP level
            pass
        return False
    except BadZipFile:
        # If we can't open the ZIP, it might be password protected
        return True
    finally:
        # Reset file pointer
        docx_file.seek(0)

def extract_xml_content(docx_file):
    """Extract all XML content from the Word document"""
    xml_files = {}

    with ZipFile(docx_file) as zip_ref:
        for file_info in zip_ref.infolist():
            if file_info.filename.endswith('.xml') or file_info.filename.endswith('.rels'):
                try:
                    content = zip_ref.read(file_info.filename).decode('utf-8')
                    xml_files[file_info.filename] = content
                except UnicodeDecodeError:
                    # Some XML files might have different encoding
                    xml_files[file_info.filename] = "Unable to decode XML content"

    # Reset file pointer
    docx_file.seek(0)
    return xml_files

def extract_metadata(doc, xml_content):
    """Extract document metadata"""
    metadata = {}

    # Get core properties if available
    core_props_file = 'docProps/core.xml'
    app_props_file = 'docProps/app.xml'

    if core_props_file in xml_content:
        try:
            root = ET.fromstring(xml_content[core_props_file])
            # Define namespaces
            namespaces = {
                'cp': 'http://schemas.openxmlformats.org/package/2006/metadata/core-properties',
                'dc': 'http://purl.org/dc/elements/1.1/',
                'dcterms': 'http://purl.org/dc/terms/',
                'dcmitype': 'http://purl.org/dc/dcmitype/',
                'xsi': 'http://www.w3.org/2001/XMLSchema-instance'
            }

            # Extract common metadata fields
            for ns, prefix in namespaces.items():
                for element in root.findall(f'.//{{{prefix}}}*'):
                    tag = element.tag.split('}')[-1]
                    if element.text:
                        metadata[tag] = element.text

        except Exception as e:
            metadata['error_core_props'] = str(e)

    if app_props_file in xml_content:
        try:
            root = ET.fromstring(xml_content[app_props_file])
            # Extended properties namespace
            ep_ns = 'http://schemas.openxmlformats.org/officeDocument/2006/extended-properties'

            for element in root.findall(f'.//{{{ep_ns}}}*'):
                tag = element.tag.split('}')[-1]
                if element.text:
                    metadata[tag] = element.text

        except Exception as e:
            metadata['error_app_props'] = str(e)

    # Add docx document properties if available
    try:
        core_properties = doc.core_properties
        metadata['author'] = core_properties.author
        metadata['created'] = str(core_properties.created) if core_properties.created else None
        metadata['last_modified_by'] = core_properties.last_modified_by
        metadata['modified'] = str(core_properties.modified) if core_properties.modified else None
        metadata['title'] = core_properties.title
        metadata['subject'] = core_properties.subject
        metadata['keywords'] = core_properties.keywords
        metadata['category'] = core_properties.category
        metadata['comments'] = core_properties.comments
        metadata['revision'] = core_properties.revision
    except Exception as e:
        metadata['error_docx_props'] = str(e)

    return metadata

def extract_statistics(doc, xml_content):
    """Extract document statistics"""
    statistics = {}

    # Count paragraphs, tables, sections
    statistics['paragraph_count'] = len(doc.paragraphs)
    statistics['table_count'] = len(doc.tables)
    statistics['section_count'] = len(doc.sections)

    # Word and character count
    word_count = 0
    char_count = 0
    for para in doc.paragraphs:
        text = para.text
        if text:
            words = text.split()
            word_count += len(words)
            char_count += len(text)

    statistics['word_count'] = word_count
    statistics['character_count'] = char_count

    # Count XML elements in document.xml
    if 'word/document.xml' in xml_content:
        try:
            root = ET.fromstring(xml_content['word/document.xml'])
            all_elements = root.findall('.//*')
            statistics['document_xml_element_count'] = len(all_elements)
        except Exception as e:
            statistics['document_xml_element_count_error'] = str(e)

    # Try to detect document language
    statistics['document_language'] = detect_document_language(doc, xml_content)

    return statistics

def detect_document_language(doc, xml_content):
    """Attempt to detect the document language"""
    # Check settings.xml for language information
    if 'word/settings.xml' in xml_content:
        try:
            root = ET.fromstring(xml_content['word/settings.xml'])
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

            # Look for language settings
            lang_elements = root.findall('.//w:lang', ns)
            if lang_elements:
                lang_attrs = lang_elements[0].attrib
                # Look for common language attributes
                for key in ['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 
                           '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bidi',
                           '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia']:
                    if key in lang_attrs:
                        return lang_attrs[key]
        except Exception:
            pass

    # Default to unknown if not found
    return "unknown"

def extract_section_orientation(xml_content):
    """Extract page orientation information"""
    orientations = []

    # Check section properties in document.xml
    if 'word/document.xml' in xml_content:
        try:
            root = ET.fromstring(xml_content['word/document.xml'])
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

            # Find all section properties
            section_props = root.findall('.//w:sectPr', ns)

            for i, sect_prop in enumerate(section_props):
                orientation = "portrait"  # Default orientation

                # Find page size element
                page_size = sect_prop.find('.//w:pgSz', ns)
                if page_size is not None:
                    # Check orientation attribute
                    orient_attr = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}orient'
                    if orient_attr in page_size.attrib and page_size.attrib[orient_attr] == 'landscape':
                        orientation = "landscape"

                    # Or check width/height
                    width_attr = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w'
                    height_attr = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}h'

                    if width_attr in page_size.attrib and height_attr in page_size.attrib:
                        width = int(page_size.attrib[width_attr])
                        height = int(page_size.attrib[height_attr])

                        if width > height:
                            orientation = "landscape"
                        else:
                            orientation = "portrait"

                orientations.append({
                    "section": i + 1,
                    "orientation": orientation
                })

        except Exception as e:
            return [{"error": str(e)}]

    # If no sections found, return default
    if not orientations:
        orientations.append({
            "section": 1,
            "orientation": "portrait"  # Default assumption
        })

    return orientations

def extract_document_text(doc):
    """Extract all text from the document including tables, footnotes, etc."""
    full_text = []

    # Helper function to process a block of content (could be paragraph or table)
    def process_content_block(block):
        if isinstance(block, Paragraph):
            full_text.append(block.text)
        elif isinstance(block, WordTable):
            for row in block.rows:
                for cell in row.cells:
                    # Process paragraphs and tables within cells recursively
                    for item in cell.paragraphs:
                        full_text.append(item.text)

    # Get all content in document order
    all_blocks = []
    for item in doc.element.body:
        if isinstance(item, CT_P):
            all_blocks.append(Paragraph(item, doc))
        elif isinstance(item, CT_Tbl):
            all_blocks.append(WordTable(item, doc))

    # Process each block
    for block in all_blocks:
        process_content_block(block)

    # Add footnotes/endnotes text (would need document.xml parsing for correct order)
    # This is a simplified approach

    return full_text

def extract_tables(doc):
    """Extract tables from the document"""
    tables_data = []

    for i, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text)
            table_data.append(row_data)

        # Generate CSV representation
        csv_data = ""
        for row in table_data:
            csv_row = []
            for cell in row:
                # Escape double quotes by doubling them and wrap in quotes
                escaped_cell = '"' + cell.replace('"', '""') + '"'
                csv_row.append(escaped_cell)
            csv_data += ",".join(csv_row) + "\r\n"

        tables_data.append({
            "table_index": i,
            "rows": len(table.rows),
            "columns": len(table.rows[0].cells) if table.rows else 0,
            "data": table_data,
            "csv": csv_data
        })

    return tables_data

def extract_images(docx_file):
    """Extract images from the document"""
    images = []

    with ZipFile(docx_file) as zip_ref:
        # Get list of all files in the zip
        file_list = zip_ref.namelist()

        # Find image files (typically in word/media/)
        image_files = [f for f in file_list if f.startswith('word/media/')]

        for i, img_path in enumerate(image_files):
            try:
                # Get the binary content of the image
                img_binary = zip_ref.read(img_path)

                # Determine content type based on file extension
                extension = img_path.split('.')[-1].lower()
                content_type = {
                    'png': 'image/png',
                    'jpeg': 'image/jpeg',
                    'jpg': 'image/jpeg',
                    'gif': 'image/gif',
                    'bmp': 'image/bmp',
                    'tiff': 'image/tiff',
                    'tif': 'image/tiff',
                    'emf': 'image/emf',
                    'wmf': 'image/wmf',
                }.get(extension, 'application/octet-stream')

                # Base64 encode the image
                img_base64 = base64.b64encode(img_binary).decode('utf-8')

                # Add to results
                images.append({
                    "image_index": i,
                    "filename": img_path.split('/')[-1],
                    "content_type": content_type,
                    "content": {"$content-type": content_type, "$content": img_base64}
                })

            except Exception as e:
                images.append({
                    "image_index": i,
                    "filename": img_path.split('/')[-1],
                    "error": str(e)
                })

    # Reset file pointer
    docx_file.seek(0)
    return images

def extract_hyperlinks(doc, xml_content):
    """Extract hyperlinks from the document"""
    hyperlinks = []

    # We need to parse document.xml to get hyperlinks with their text
    if 'word/document.xml' in xml_content:
        try:
            root = ET.fromstring(xml_content['word/document.xml'])
            ns = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
            }

            # Find all hyperlink elements
            hyperlink_elements = root.findall('.//w:hyperlink', ns)

            for i, elem in enumerate(hyperlink_elements):
                # Get relationship ID
                r_id = elem.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')

                # Get the link text
                text_elements = elem.findall('.//w:t', ns)
                link_text = ' '.join([e.text for e in text_elements if e.text])

                hyperlinks.append({
                    "hyperlink_index": i,
                    "text": link_text,
                    "relationship_id": r_id,
                    "url": get_hyperlink_target(r_id, xml_content)
                })

        except Exception as e:
            hyperlinks.append({
                "error": str(e)
            })

    return hyperlinks

def get_hyperlink_target(r_id, xml_content):
    """Get the target URL for a hyperlink relationship ID"""
    if not r_id or 'word/_rels/document.xml.rels' not in xml_content:
        return None

    try:
        # Parse relationships file
        root = ET.fromstring(xml_content['word/_rels/document.xml.rels'])

        # Find relationship with matching ID
        for rel in root.findall('.//{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
            if rel.get('Id') == r_id:
                return rel.get('Target')

    except Exception:
        pass

    return None

def extract_table_of_contents(doc, xml_content):
    """Extract table of contents if it exists"""
    toc_entries = []

    # Check for TOC fields in document.xml
    if 'word/document.xml' in xml_content:
        try:
            root = ET.fromstring(xml_content['word/document.xml'])
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

            # Find TOC field begin markers
            toc_begin = root.findall('.//w:fldChar[@w:fldCharType="begin"]/../../../..', ns)

            for toc in toc_begin:
                # Look for TOC instruction text
                instr_text = toc.findall('.//w:instrText', ns)
                for instr in instr_text:
                    if instr.text and 'TOC' in instr.text:
                        # This confirms it's a TOC
                        # Now find all paragraphs between TOC begin and end
                        # This is a simplified approach and may not work for all documents
                        for para in doc.paragraphs:
                            if "TOC" in para.text or "Table of Contents" in para.text:
                                continue  # Skip TOC title

                            if hasattr(para, 'style') and para.style and para.style.name:
                                # Check if it's a TOC style (typically TOC1, TOC2, etc.)
                                if para.style.name.startswith('TOC'):
                                    level = int(para.style.name.replace('TOC', '')) if para.style.name[3:].isdigit() else 0
                                    toc_entries.append({
                                        "level": level,
                                        "text": para.text,
                                    })

        except Exception as e:
            toc_entries.append({
                "error": str(e)
            })

    return toc_entries

def extract_headings(doc):
    """Extract headings from the document"""
    headings = []

    for i, para in enumerate(doc.paragraphs):
        if hasattr(para, 'style') and para.style and para.style.name:
            # Check if paragraph style is a heading style
            if para.style.name.startswith('Heading'):
                try:
                    level = int(para.style.name.replace('Heading', '')) if para.style.name[7:].isdigit() else 0
                    headings.append({
                        "index": i,
                        "level": level,
                        "text": para.text
                    })
                except ValueError:
                    # If heading level can't be determined
                    headings.append({
                        "index": i,
                        "level": 0,  # Default level
                        "text": para.text,
                        "style": para.style.name
                    })
            # Also check for custom heading styles
            elif 'head' in para.style.name.lower():
                headings.append({
                    "index": i,
                    "text": para.text,
                    "style": para.style.name
                })

    return headings

def extract_comments(xml_content):
    """Extract comments from the document"""
    comments = []

    # Check for comments in comments.xml
    if 'word/comments.xml' in xml_content:
        try:
            root = ET.fromstring(xml_content['word/comments.xml'])
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

            # Find all comment elements
            comment_elements = root.findall('.//w:comment', ns)

            for i, comment in enumerate(comment_elements):
                # Get comment ID
                comment_id = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')

                # Get author
                author = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author', 'Unknown')

                # Get date
                date_str = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date', None)

                # Get comment text
                text_elements = comment.findall('.//w:t', ns)
                comment_text = ' '.join([e.text for e in text_elements if e.text])

                comments.append({
                    "comment_index": i,
                    "id": comment_id,
                    "author": author,
                    "date": date_str,
                    "text": comment_text
                })

        except Exception as e:
            comments.append({
                "error": str(e)
            })

    return comments

def extract_digital_signatures(xml_content):
    """Extract digital signatures from the document"""
    signatures = []

    # Check for digital signature info
    signature_files = [f for f in xml_content.keys() if 'digitalSignature' in f or '_signatures' in f]

    if not signature_files:
        return []  # No signatures found

    for sig_file in signature_files:
        try:
            root = ET.fromstring(xml_content[sig_file])

            # This is a simplified approach as signature XML can vary
            # Try to extract common signature elements
            for i, sig_element in enumerate(root.findall('.//*')):
                if 'signature' in sig_element.tag.lower():
                    sig_info = {
                        "signature_index": i,
                        "type": sig_element.tag.split('}')[-1]
                    }

                    # Extract available properties
                    for child in sig_element:
                        tag = child.tag.split('}')[-1]
                        if child.text:
                            sig_info[tag] = child.text

                    signatures.append(sig_info)

        except Exception as e:
            signatures.append({
                "error": str(e),
                "file": sig_file
            })

    return signatures






@app.route(route="extract_main_word_xml")
def extract_main_word_xml(req: func.HttpRequest) -> func.HttpResponse:
    try:
        # Parse request JSON
        req_body = req.get_json()
        file_content_base64 = req_body.get("file_content")["$content"]

        if not file_content_base64:
            return func.HttpResponse("Missing 'file_content' in request.", status_code=400)

        # Decode base64 content
        file_bytes = base64.b64decode(file_content_base64)

        # Load Word file as ZIP
        zip_buffer = BytesIO(file_bytes)
        extracted_data = {}

        with ZipFile(zip_buffer, 'r') as docx_zip:
            for file_name in docx_zip.namelist():
                # Filter to relevant XML files
                if (file_name.startswith("word/document") or file_name=="word/comments.xml" or file_name.startswith("word/_rels/document") or file_name.startswith("word/header") or file_name.startswith("word/footer") or file_name.startswith("word/endnotes") or file_name.startswith("word/footnotes") or file_name.startswith("word/numbering") or file_name.startswith("word/theme") or file_name.startswith("word/theme/theme")) and (file_name.endswith(".xml") or file_name.endswith(".rels")):
                    with docx_zip.open(file_name) as file:
                        extracted_data[file_name] = file.read().decode("utf-8")

        return func.HttpResponse(json.dumps(extracted_data, indent=2), mimetype="application/json")

    except Exception as e:
        logging.error(f"Error processing file: {str(e)}\n\nTraceback:\n{traceback.format_exc()}")
        return func.HttpResponse(
            f"Error: {str(e)}\n\nTraceback:\n{traceback.format_exc()}",
            status_code=500
        )






@app.route(route="update_word_xml")
def update_word_xml(req: func.HttpRequest) -> func.HttpResponse:
    try:
        # Parse request JSON
        req_body = req.get_json()
        file_content_base64 = req_body.get("file_content", {}).get("$content")
        updated_xml_data = req_body.get("updated_xml")

        if not file_content_base64 or not updated_xml_data:
            return func.HttpResponse("Missing required inputs.", status_code=400)

        # Decode base64 original Word file
        file_bytes = base64.b64decode(file_content_base64)
        zip_buffer = BytesIO(file_bytes)
        output_buffer = BytesIO()

        with ZipFile(zip_buffer, 'r') as docx_zip:
            with ZipFile(output_buffer, 'w', ZIP_DEFLATED) as new_docx_zip:
                for file_name in docx_zip.namelist():
                    with docx_zip.open(file_name) as file:
                        # Replace XML content if it's updated
                        if file_name in updated_xml_data:
                            new_content = updated_xml_data[file_name].encode("utf-8")
                        else:
                            new_content = file.read()
                        new_docx_zip.writestr(file_name, new_content)

                # Add new XML files that are in updated_xml_data but not in the original document
                for new_file_name, new_content in updated_xml_data.items():
                    if new_file_name not in docx_zip.namelist():
                        new_docx_zip.writestr(new_file_name, new_content.encode("utf-8"))

        # Get base64-encoded updated Word file
        output_buffer.seek(0)
        base64_encoded_docx = base64.b64encode(output_buffer.getvalue()).decode("utf-8")

        response_data = {
            "$content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "$content": base64_encoded_docx
        }

        return func.HttpResponse(json.dumps(response_data), mimetype="application/json")

    except Exception as e:
        logging.error(f"Error updating file: {str(e)}\n\nTraceback:\n{traceback.format_exc()}")
        return func.HttpResponse(
            f"Error: {str(e)}\n\nTraceback:\n{traceback.format_exc()}",
            status_code=500
        )






@app.route(route="merge_word")
def merge_word(req: func.HttpRequest) -> func.HttpResponse:
    try:
        req_body = req.get_json()
        file_contents = req_body.get("file_content", [])
        # Updated parameter for header/footer behavior with default "NONE"
        headers_footers = req_body.get("headers_footers", "NONE").upper()

        # Validate headers_footers option
        if headers_footers not in ["NONE", "FIRST", "LAST"]:
            return func.HttpResponse(f"Invalid headers_footers value. Must be one of: NONE, FIRST, LAST", status_code=400)

        if not file_contents:
            return func.HttpResponse("Missing or empty 'file_content' array", status_code=400)

        word_files = []
        for file_obj in file_contents:
            if file_obj.get("$content-type") != "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return func.HttpResponse("Invalid file type detected.", status_code=400)
            word_files.append(base64.b64decode(file_obj["$content"]))

        merged_content = merge_documents(word_files, headers_footers)

        response_data = {
            "$content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "$content": base64.b64encode(merged_content).decode('utf-8')
        }

        return func.HttpResponse(
            body=json.dumps(response_data),
            status_code=200,
            mimetype="application/json"
        )
    except Exception as e:
        logging.error(f"Error merge_word: {str(e)}\n\nTraceback:\n{traceback.format_exc()}")
        return func.HttpResponse(f"merge_word: {str(e)}\n\nTraceback:\n{traceback.format_exc()}", status_code=500)

def merge_documents(doc_contents, headers_footers="NONE"):
    try:
        base_temp_dir = tempfile.mkdtemp(prefix="docx_merge_")

        try:
            if not doc_contents:
                raise ValueError("doc_contents is empty, cannot merge documents.")

            # Extract all documents and analyze them
            doc_files = []
            all_extract_dirs = []
            base_file_index = 0  # Default to first file as base

            # Step 1: Extract all files and determine which file to use as base
            for i, content in enumerate(doc_contents):
                doc_file = os.path.join(base_temp_dir, f"doc_{i}.docx")
                with open(doc_file, 'wb') as f:
                    f.write(content)
                doc_files.append(doc_file)

                # Extract the docx for analysis
                extract_dir = os.path.join(base_temp_dir, f"doc_{i}_extract")
                os.makedirs(extract_dir, exist_ok=True)
                with ZipFile(doc_file, 'r') as zip_ref:
                    zip_ref.extractall(extract_dir)
                all_extract_dirs.append(extract_dir)

            # Determine which file to use as the base for headers/footers
            if headers_footers in ["FIRST", "LAST"]:
                # Find files with headers/footers
                files_with_headers_footers = []
                for i, extract_dir in enumerate(all_extract_dirs):
                    has_headers = len(glob.glob(os.path.join(extract_dir, "word", "header*.xml"))) > 0
                    has_footers = len(glob.glob(os.path.join(extract_dir, "word", "footer*.xml"))) > 0

                    if has_headers or has_footers:
                        files_with_headers_footers.append(i)

                if files_with_headers_footers:
                    if headers_footers == "FIRST":
                        base_file_index = files_with_headers_footers[0]
                    else:  # LAST
                        base_file_index = files_with_headers_footers[-1]

            # Step 2: Create result directory using the base file
            # Extract base file to result directory to preserve headers/footers structure
            result_extract_dir = os.path.join(base_temp_dir, "result_extract")
            shutil.copytree(all_extract_dirs[base_file_index], result_extract_dir)

            # Step 3: Process all doc content to add to the base file
            master_doc = Document(doc_files[base_file_index])

            # Get the section properties before clearing the content
            # This is the key fix - we need to preserve the sectPr element that contains header/footer references
            last_section = master_doc.sections[-1]
            section_properties = None

            # Extract the original sectPr XML element
            for element in master_doc.element.body.xpath('.//w:sectPr'):
                section_properties = deepcopy(element)
                break

            # Clear the content of master_doc to start fresh but preserve structure
            while len(master_doc.element.body) > 0:
                master_doc.element.body.remove(master_doc.element.body[0])

            # Add content from all files including the base file (in proper order)
            for i, doc_file in enumerate(doc_files):
                doc_to_process = Document(doc_file)

                # Add page break between documents, except for the first one
                if i > 0:
                    master_doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

                for element in doc_to_process.element.body:
                    if element.tag.endswith('p') or element.tag.endswith('tbl'):
                        # Skip the sectPr element if it's embedded in a paragraph
                        if not (element.tag.endswith('p') and element.xpath('.//w:sectPr')):
                            master_doc.element.body.append(deepcopy(element))

            # Re-add the section properties at the end to preserve headers/footers
            if section_properties is not None and headers_footers != "NONE":
                # Append the section properties to the end of the document
                master_doc.element.body.append(section_properties)

            # Save the initial merged document
            temp_result = os.path.join(base_temp_dir, "temp_result.docx")
            master_doc.save(temp_result)

            # Update document.xml in result directory with new content
            with ZipFile(temp_result, 'r') as zip_ref:
                document_content = zip_ref.read('word/document.xml')

            with open(os.path.join(result_extract_dir, "word", "document.xml"), 'wb') as f:
                f.write(document_content)

            # Process all files to merge additional elements
            for i, doc_file in enumerate(doc_files):
                if i == base_file_index and headers_footers != "NONE":
                    # Skip the header/footer processing for the base file
                    # as its headers/footers are already in place
                    continue

                doc_extract_dir = all_extract_dirs[i]
                process_document_components(doc_extract_dir, result_extract_dir, i, headers_footers == "NONE")

            # Merge content types from all documents
            merge_content_types(all_extract_dirs, result_extract_dir)

            # Create final document
            final_docx_path = os.path.join(base_temp_dir, "final_merged.docx")
            create_final_docx(result_extract_dir, final_docx_path)

            with open(final_docx_path, 'rb') as f:
                final_content = f.read()

            return final_content

        finally:
            try:
                shutil.rmtree(base_temp_dir)
            except:
                pass

    except Exception as e:
        error_msg = f"Error merge_documents: {str(e)}\n{traceback.format_exc()}"
        logging.error(error_msg)
        raise RuntimeError(f"merge_documents: {error_msg}")

def process_document_components(source_dir, target_dir, doc_index, ignore_headers_footers=True):
    """Process and merge document components (styles, media, comments, hyperlinks)"""
    try:
        # Handle comments
        doc_comments_path = os.path.join(source_dir, "word", "comments.xml")
        if os.path.exists(doc_comments_path):
            merge_comments(source_dir, target_dir, doc_index)

        # Handle styles
        merge_styles(source_dir, target_dir)

        # Handle headers/footers if needed (when ignore_headers_footers is False and not the base file)
        if not ignore_headers_footers:
            # Headers/footers are handled by using the base file directly
            pass

        # Create target media directory first
        result_media_dir = os.path.join(target_dir, "word", "media")
        os.makedirs(result_media_dir, exist_ok=True)

        # Handle media files and their relationships
        # First check for media in the root directory
        root_media_dir = os.path.join(source_dir, "media")
        if os.path.exists(root_media_dir):
            # Move images from root media to word/media
            for media_file in os.listdir(root_media_dir):
                source_media = os.path.join(root_media_dir, media_file)
                new_media_name = f"doc_{doc_index}_{media_file}"
                target_media = os.path.join(result_media_dir, new_media_name)
                shutil.copy2(source_media, target_media)

        # Then check for media in the word directory
        word_media_dir = os.path.join(source_dir, "word", "media")
        if os.path.exists(word_media_dir):
            # Copy media files with new names to avoid conflicts
            for media_file in os.listdir(word_media_dir):
                source_media = os.path.join(word_media_dir, media_file)
                new_media_name = f"doc_{doc_index}_{media_file}"
                target_media = os.path.join(result_media_dir, new_media_name)
                shutil.copy2(source_media, target_media)

        # Create a mapping of old image paths to new ones
        image_path_map = {}

        # Map both root and word media paths
        if os.path.exists(root_media_dir):
            for media_file in os.listdir(root_media_dir):
                image_path_map[f"media/{media_file}"] = f"media/doc_{doc_index}_{media_file}"
                image_path_map[f"/media/{media_file}"] = f"media/doc_{doc_index}_{media_file}"

        if os.path.exists(word_media_dir):
            for media_file in os.listdir(word_media_dir):
                image_path_map[f"media/{media_file}"] = f"media/doc_{doc_index}_{media_file}"

        # Now handle the relationships
        source_rels_path = os.path.join(source_dir, "word", "_rels", "document.xml.rels")
        target_rels_path = os.path.join(target_dir, "word", "_rels", "document.xml.rels")

        if os.path.exists(source_rels_path) and os.path.exists(target_rels_path):
            merge_media_relationships(source_rels_path, target_rels_path, image_path_map, doc_index)

        # Handle hyperlinks
        fix_hyperlinks(source_dir, target_dir, doc_index)

    except Exception as e:
        logging.error(f"Error process_document_components: {str(e)}")
        raise RuntimeError(f"process_document_components: {str(e)}")

def merge_media_relationships(source_rels_path, target_rels_path, image_path_map, doc_index):
    """Merge media relationships from source to target document."""
    try:
        # Parse source and target relationship files
        source_tree = etree.parse(source_rels_path)
        source_root = source_tree.getroot()

        target_tree = etree.parse(target_rels_path)
        target_root = target_tree.getroot()

        # Find the highest existing rId in target
        next_id = 1
        id_map = {}  # Map old rIds to new ones

        for rel in target_root.findall(".//{*}Relationship"):
            rel_id = rel.get("Id", "")
            if rel_id and rel_id.startswith("rId"):
                try:
                    id_num = int(rel_id[3:])
                    next_id = max(next_id, id_num + 1)
                except ValueError:
                    pass

        # Process media relationships from source
        for rel in source_root.findall(".//{*}Relationship"):
            rel_type = rel.get("Type", "")
            rel_target = rel.get("Target", "")
            rel_id = rel.get("Id", "")

            # Check if this is an image relationship
            if "image" in rel_type.lower():
                # Fix Target path format - make sure it starts with "media/" without leading slash
                if rel_target.startswith("/media/"):
                    rel_target = rel_target[1:]  # Remove leading slash
                elif not rel_target.startswith("media/"):
                    rel_target = "media/" + os.path.basename(rel_target)

                # Map to new target if in our path mapping
                if rel_target in image_path_map:
                    new_target = image_path_map[rel_target]
                else:
                    # Create new target for images not in mapping
                    new_target = f"media/doc_{doc_index}_{os.path.basename(rel_target)}"
                    # Copy the file if it exists
                    source_file = os.path.join(os.path.dirname(os.path.dirname(source_rels_path)), rel_target)
                    target_file = os.path.join(os.path.dirname(os.path.dirname(target_rels_path)), new_target)
                    os.makedirs(os.path.dirname(target_file), exist_ok=True)
                    if os.path.exists(source_file):
                        shutil.copy2(source_file, target_file)

                new_id = f"rId{next_id}"
                next_id += 1

                # Create new relationship
                new_rel = etree.SubElement(target_root, "{http://schemas.openxmlformats.org/package/2006/relationships}Relationship")
                new_rel.set("Id", new_id)
                new_rel.set("Type", rel.get("Type"))
                new_rel.set("Target", new_target)

                # Remember the mapping from old to new rId
                id_map[rel_id] = new_id

        # Save the updated relationships file
        target_tree.write(target_rels_path, encoding='UTF-8', xml_declaration=True)

        # Now update the document.xml to reference the new relationship IDs
        document_path = os.path.join(os.path.dirname(os.path.dirname(target_rels_path)), "document.xml")
        if os.path.exists(document_path) and id_map:
            update_document_image_refs(document_path, id_map, doc_index)

    except Exception as e:
        logging.error(f"Error merge_media_relationships: {str(e)}")
        raise RuntimeError(f"merge_media_relationships: {str(e)}")

def update_document_image_refs(document_path, id_map, doc_index):
    """Update image references in document.xml with new relationship IDs."""
    try:
        # Parse document
        parser = etree.XMLParser(remove_blank_text=True)
        doc_tree = etree.parse(document_path, parser)
        doc_root = doc_tree.getroot()

        # Define namespaces
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
            'wp14': 'http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing',
            'a14': 'http://schemas.microsoft.com/office/drawing/2010/main'
        }

        # Find paragraphs with drawings after the document break
        paragraphs = doc_root.xpath("//w:p", namespaces=namespaces)

        # Find page breaks to determine document boundaries
        page_breaks = []
        for i, p in enumerate(paragraphs):
            if p.xpath(".//w:br[@w:type='page']", namespaces=namespaces):
                page_breaks.append(i)

        # Start processing from the correct section
        start_index = 0
        if doc_index > 0 and page_breaks:
            # Find the relevant page break for this document
            if doc_index <= len(page_breaks):
                start_index = page_breaks[doc_index - 1] + 1

        # Process paragraphs from the start_index
        for i in range(start_index, len(paragraphs)):
            # Update blip references (images)
            for blip in paragraphs[i].xpath(".//a:blip", namespaces=namespaces):
                old_id = blip.get("{" + namespaces['r'] + "}embed")
                if old_id in id_map:
                    blip.set("{" + namespaces['r'] + "}embed", id_map[old_id])

                    # Ensure all necessary namespaces are defined in parent elements
                    # This fixes the namespace references in the XML
                    drawing_element = blip.xpath("./ancestor::w:drawing", namespaces=namespaces)
                    if drawing_element:
                        ensure_namespaces(drawing_element[0], namespaces)

        # Save the updated document
        doc_tree.write(document_path, encoding='UTF-8', xml_declaration=True)

    except Exception as e:
        logging.error(f"Error update_document_image_refs: {str(e)}")
        raise RuntimeError(f"update_document_image_refs: {str(e)}")

def update_media_references(target_dir, old_media_name, new_media_name, doc_index):
    """Update references to media files in the document and relationships."""
    try:
        # Update document.xml.rels
        rels_path = os.path.join(target_dir, "word", "_rels", "document.xml.rels")
        if os.path.exists(rels_path):
            # Parse the relationships file
            rels_tree = etree.parse(rels_path)
            rels_root = rels_tree.getroot()

            # Find existing media relationships
            for rel in rels_root.findall(".//{*}Relationship"):
                target = rel.get("Target")
                if target and "media/" in target and old_media_name in target:
                    # Update the target to point to the new media name
                    new_target = target.replace(old_media_name, new_media_name)
                    rel.set("Target", new_target)

            # Save the updated relationships file
            rels_tree.write(rels_path, encoding='UTF-8', xml_declaration=True)

        # Update document.xml
        doc_path = os.path.join(target_dir, "word", "document.xml")
        if os.path.exists(doc_path):
            # Parse document
            doc_tree = etree.parse(doc_path)
            doc_root = doc_tree.getroot()

            # Define namespaces for XPath queries
            namespaces = {
                'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
                'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
            }

            # Find image elements with blip references
            for blip in doc_root.xpath("//a:blip", namespaces=namespaces):
                embed = blip.get("{" + namespaces['r'] + "}embed", "")
                if embed:
                    # The embed attribute refers to a relationship ID in document.xml.rels
                    # We don't change this directly, as we updated the relationship target above
                    pass

            # Save the updated document
            doc_tree.write(doc_path, encoding='UTF-8', xml_declaration=True)

    except Exception as e:
        logging.error(f"Error update_media_references: {e}")
        raise RuntimeError(f"update_media_references: {str(e)}")

def ensure_namespaces(element, namespaces):
    """Ensure all necessary namespace declarations are present in the element and its children."""
    # Add missing namespace declarations to elements where needed
    for prefix, uri in namespaces.items():
        nsmap = element.nsmap if hasattr(element, 'nsmap') else {}
        if prefix not in nsmap:
            # We can't directly modify nsmap in lxml, but we can add attributes
            if prefix == 'xmlns':
                element.set('xmlns', uri)
            else:
                element.set(f'xmlns:{prefix}', uri)

    # Recursively ensure namespaces for all children with specific tags
    for tag in ['wp:inline', 'a:graphic', 'a:graphicData', 'pic:pic', 'pic:nvPicPr', 'pic:blipFill', 'pic:spPr']:
        for child in element.xpath(f'.//{tag}', namespaces=namespaces):
            ensure_namespaces(child, namespaces)

def fix_media_paths(extract_dir):
    """Fix media paths in the relationship files to ensure they use the correct format."""
    try:
        # Fix document.xml.rels
        rels_path = os.path.join(extract_dir, "word", "_rels", "document.xml.rels")
        if os.path.exists(rels_path):
            # Parse the relationships file
            tree = etree.parse(rels_path)
            root = tree.getroot()

            # Check all relationships
            relationships_modified = False
            for rel in root.findall(".//{*}Relationship"):
                rel_type = rel.get("Type", "")
                target = rel.get("Target", "")

                # Fix image paths only
                if "image" in rel_type.lower() and target:
                    # Ensure all image paths are in word/media folder without leading slash
                    if target.startswith("/media/"):
                        # Update to correct path format
                        rel.set("Target", target[1:])
                        relationships_modified = True

            # Save only if modified
            if relationships_modified:
                tree.write(rels_path, encoding='UTF-8', xml_declaration=True)

        # Ensure all images are in word/media folder
        root_media_dir = os.path.join(extract_dir, "media")
        word_media_dir = os.path.join(extract_dir, "word", "media")

        # Create word/media if it doesn't exist
        os.makedirs(word_media_dir, exist_ok=True)

        # Move any images from root media to word/media
        if os.path.exists(root_media_dir):
            for media_file in os.listdir(root_media_dir):
                source_media = os.path.join(root_media_dir, media_file)
                target_media = os.path.join(word_media_dir, media_file)
                if not os.path.exists(target_media):
                    shutil.copy2(source_media, target_media)

    except Exception as e:
        logging.error(f"Error fix_media_paths: {e}")
        raise RuntimeError(f"fix_media_paths: {str(e)}")

def merge_comments(source_dir, target_dir, doc_index):
    """Merge comments from a source document into the target document."""
    try:
        source_comments_path = os.path.join(source_dir, "word", "comments.xml")
        target_comments_path = os.path.join(target_dir, "word", "comments.xml")
        target_comments_dir = os.path.dirname(target_comments_path)

        os.makedirs(target_comments_dir, exist_ok=True)

        # Define namespaces
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        }

        # If source comments file doesn't exist, nothing to do
        if not os.path.exists(source_comments_path):
            return

        # Parse source comments file
        source_tree = etree.parse(source_comments_path)
        source_root = source_tree.getroot()

        # If target comments file doesn't exist yet, create it
        if not os.path.exists(target_comments_path):
            # Create a new comments file
            comments_root = etree.Element("{" + namespaces['w'] + "}comments")
            comments_tree = etree.ElementTree(comments_root)
            comments_tree.write(target_comments_path, encoding='UTF-8', xml_declaration=True)

        # Parse target comments file
        target_tree = etree.parse(target_comments_path)
        target_root = target_tree.getroot()

        # Find current highest comment ID
        existing_ids = target_root.xpath("//w:comment/@w:id", namespaces=namespaces)
        next_id = 0
        if existing_ids:
            next_id = max(int(x) for x in existing_ids) + 1

        # Map to store old ID to new ID mapping
        id_map = {}

        # Process each comment
        for comment in source_root.xpath("//w:comment", namespaces=namespaces):
            old_id = comment.get("{" + namespaces['w'] + "}id")
            new_id = str(next_id)

            # Update ID mapping
            id_map[old_id] = new_id

            # Update comment ID
            comment.set("{" + namespaces['w'] + "}id", new_id)
            next_id += 1

            # Deep copy to avoid XML namespace issues
            comment_copy = deepcopy(comment)
            target_root.append(comment_copy)

        # Save the updated comments file
        target_tree.write(target_comments_path, encoding='UTF-8', xml_declaration=True)

        # Update document.xml to reference the new comment IDs
        document_path = os.path.join(target_dir, "word", "document.xml")

        if os.path.exists(document_path):
            update_comment_references(document_path, id_map, doc_index, namespaces)

    except Exception as e:
        logging.error(f"Error merge_comments: {e}")
        raise RuntimeError(f"merge_comments: {str(e)}")

def update_comment_references(document_path, id_map, doc_index, namespaces):
    """Update comment references in the document.xml file."""
    try:
        # Parse the document
        doc_tree = etree.parse(document_path)
        doc_root = doc_tree.getroot()

        # Process all paragraphs after the doc_index page break
        paragraphs = doc_root.xpath("//w:p", namespaces=namespaces)

        # Find page breaks to determine document boundaries
        page_breaks = []
        for i, p in enumerate(paragraphs):
            if p.xpath(".//w:br[@w:type='page']", namespaces=namespaces):
                page_breaks.append(i)

        # Start processing from the correct section
        start_index = 0
        if doc_index > 0 and page_breaks:
            # Find the relevant page break for this document
            if doc_index <= len(page_breaks):
                start_index = page_breaks[doc_index - 1] + 1

        # Process paragraphs from the start_index
        for i in range(start_index, len(paragraphs)):
            # Process comment range starts
            for element in paragraphs[i].xpath(".//w:commentRangeStart", namespaces=namespaces):
                old_id = element.get("{" + namespaces['w'] + "}id")
                if old_id in id_map:
                    element.set("{" + namespaces['w'] + "}id", id_map[old_id])

            # Process comment range ends
            for element in paragraphs[i].xpath(".//w:commentRangeEnd", namespaces=namespaces):
                old_id = element.get("{" + namespaces['w'] + "}id")
                if old_id in id_map:
                    element.set("{" + namespaces['w'] + "}id", id_map[old_id])

            # Process comment references
            for element in paragraphs[i].xpath(".//w:commentReference", namespaces=namespaces):
                old_id = element.get("{" + namespaces['w'] + "}id")
                if old_id in id_map:
                    element.set("{" + namespaces['w'] + "}id", id_map[old_id])

        # Save the updated document
        doc_tree.write(document_path, encoding='UTF-8', xml_declaration=True)

    except Exception as e:
        logging.error(f"Error update_comment_references: {e}")
        raise RuntimeError(f"update_comment_references: {str(e)}")

def merge_styles(source_dir, target_dir):
    """Merge styles from source document into target document."""
    try:
        source_styles_path = os.path.join(source_dir, "word", "styles.xml")
        target_styles_path = os.path.join(target_dir, "word", "styles.xml")

        # If source doesn't have styles, nothing to do
        if not os.path.exists(source_styles_path):
            return

        # Define namespaces
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        }

        # If target doesn't have styles yet, copy the source file
        if not os.path.exists(target_styles_path):
            os.makedirs(os.path.dirname(target_styles_path), exist_ok=True)
            source_tree = etree.parse(source_styles_path)
            source_tree.write(target_styles_path, encoding='UTF-8', xml_declaration=True)
            return

        # Parse both files
        source_tree = etree.parse(source_styles_path)
        source_root = source_tree.getroot()

        target_tree = etree.parse(target_styles_path)
        target_root = target_tree.getroot()

        # Get all style IDs from target
        target_style_ids = set()
        for style in target_root.xpath("//w:style", namespaces=namespaces):
            style_id = style.get("{" + namespaces['w'] + "}styleId")
            if style_id:
                target_style_ids.add(style_id)

        # Add unique styles from source
        for style in source_root.xpath("//w:style", namespaces=namespaces):
            style_id = style.get("{" + namespaces['w'] + "}styleId")
            if style_id and style_id not in target_style_ids:
                # Copy this style to the target
                target_root.append(etree.fromstring(etree.tostring(style)))
                target_style_ids.add(style_id)

        # Save the merged styles
        target_tree.write(target_styles_path, encoding='UTF-8', xml_declaration=True)

    except Exception as e:
        logging.error(f"Error merge_styles: {e}")
        raise RuntimeError(f"merge_styles: {str(e)}")

def fix_hyperlinks(source_dir, target_dir, doc_index):
    """Fix hyperlinks in the merged document."""
    try:
        source_rels_path = os.path.join(source_dir, "word", "_rels", "document.xml.rels")
        target_rels_path = os.path.join(target_dir, "word", "_rels", "document.xml.rels")

        if not os.path.exists(source_rels_path):
            return

        # Ensure target directory exists
        os.makedirs(os.path.dirname(target_rels_path), exist_ok=True)

        # Parse source relationships
        source_tree = etree.parse(source_rels_path)
        source_root = source_tree.getroot()

        # Parse or create target relationships
        if os.path.exists(target_rels_path):
            target_tree = etree.parse(target_rels_path)
            target_root = target_tree.getroot()
        else:
            # Create a new relationships file
            nsmap = {"xmlns": "http://schemas.openxmlformats.org/package/2006/relationships"}
            target_root = etree.Element("{http://schemas.openxmlformats.org/package/2006/relationships}Relationships")
            target_tree = etree.ElementTree(target_root)

        # Get all hyperlink relationships from source
        hyperlink_rels = []
        for rel in source_root.findall(".//{*}Relationship"):
            if rel.get("Type") == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink":
                hyperlink_rels.append(rel)

        if not hyperlink_rels:
            return

        # Find the highest existing rId
        next_id = 1
        for rel in target_root.findall(".//{*}Relationship"):
            rel_id = rel.get("Id", "")
            if rel_id and rel_id.startswith("rId"):
                try:
                    id_num = int(rel_id[3:])
                    next_id = max(next_id, id_num + 1)
                except ValueError:
                    pass

        # Create a mapping of old IDs to new IDs
        hyperlink_rel_map = {}

        # Add hyperlink relationships to target
        for rel in hyperlink_rels:
            old_id = rel.get("Id")
            new_id = f"rId{next_id}"
            next_id += 1

            # Create new relationship
            new_rel = etree.SubElement(target_root, "{http://schemas.openxmlformats.org/package/2006/relationships}Relationship")
            new_rel.set("Id", new_id)
            new_rel.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink")
            new_rel.set("Target", rel.get("Target"))
            new_rel.set("TargetMode", "External")

            hyperlink_rel_map[old_id] = new_id

        # Save the updated relationships file
        target_tree.write(target_rels_path, encoding='UTF-8', xml_declaration=True)

        # Update hyperlinks in document.xml
        document_path = os.path.join(target_dir, "word", "document.xml")
        update_hyperlink_references(document_path, hyperlink_rel_map)

    except Exception as e:
        logging.error(f"Error fix_hyperlinks: {str(e)}")
        raise RuntimeError(f"fix_hyperlinks: {str(e)}")

def update_hyperlink_references(document_path, hyperlink_rel_map):
    """Update hyperlink references in the document."""
    try:
        if not os.path.exists(document_path) or not hyperlink_rel_map:
            return

        # Parse document
        doc_tree = etree.parse(document_path)
        doc_root = doc_tree.getroot()

        # Define namespaces
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
        }

        # Find all hyperlinks
        hyperlinks = doc_root.xpath("//w:hyperlink", namespaces=namespaces)

        # Update hyperlink relationship IDs
        for hyperlink in hyperlinks:
            old_id = hyperlink.get("{" + namespaces['r'] + "}id")
            if old_id in hyperlink_rel_map:
                hyperlink.set("{" + namespaces['r'] + "}id", hyperlink_rel_map[old_id])

        # Save updated document
        doc_tree.write(document_path, encoding='UTF-8', xml_declaration=True)

    except Exception as e:
        logging.error(f"Error update_hyperlink_references: {e}")
        raise RuntimeError(f"update_hyperlink_references: {str(e)}")

def merge_content_types(source_dirs, target_dir):
    """Merge content types from all source directories into the target directory."""
    try:
        # Path to target content types file
        target_content_types_path = os.path.join(target_dir, "[Content_Types].xml")

        # Parse target content types file
        target_tree = None
        target_root = None

        if os.path.exists(target_content_types_path):
            target_tree = etree.parse(target_content_types_path)
            target_root = target_tree.getroot()
        else:
            # Create a new content types file if it doesn't exist
            root_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
            target_root = etree.Element("{" + root_ns + "}Types")
            target_tree = etree.ElementTree(target_root)

        # Track existing content types to avoid duplicates
        existing_defaults = set()
        existing_overrides = set()

        # Get existing content types
        for default in target_root.findall(".//{*}Default"):
            extension = default.get("Extension", "")
            if extension:
                existing_defaults.add(extension.lower())

        for override in target_root.findall(".//{*}Override"):
            part_name = override.get("PartName", "")
            if part_name:
                existing_overrides.add(part_name)

        # Process each source directory
        for source_dir in source_dirs:
            source_content_types_path = os.path.join(source_dir, "[Content_Types].xml")
            if os.path.exists(source_content_types_path):
                source_tree = etree.parse(source_content_types_path)
                source_root = source_tree.getroot()

                # Add missing Default entries (file extension mappings)
                for default in source_root.findall(".//{*}Default"):
                    extension = default.get("Extension", "")
                    content_type = default.get("ContentType", "")

                    if extension and content_type and extension.lower() not in existing_defaults:
                        # Create a new Default element
                        new_default = etree.SubElement(target_root, "{" + target_root.nsmap[None] + "}Default")
                        new_default.set("Extension", extension)
                        new_default.set("ContentType", content_type)
                        existing_defaults.add(extension.lower())
                        logging.info(f"Added content type mapping for extension: {extension}")

                # Add missing Override entries (specific part mappings)
                for override in source_root.findall(".//{*}Override"):
                    part_name = override.get("PartName", "")
                    content_type = override.get("ContentType", "")

                    # Skip existing overrides, but include new ones
                    if part_name and content_type and part_name not in existing_overrides:
                        # Check if this is a part we actually have in the target
                        part_path = part_name.lstrip('/')
                        if os.path.exists(os.path.join(target_dir, part_path)):
                            new_override = etree.SubElement(target_root, "{" + target_root.nsmap[None] + "}Override")
                            new_override.set("PartName", part_name)
                            new_override.set("ContentType", content_type)
                            existing_overrides.add(part_name)

        # Ensure common image content types are included
        common_image_types = {
            "jpeg": "image/jpeg",
            "jpg": "image/jpeg",
            "png": "image/png",
            "gif": "image/gif",
            "bmp": "image/bmp",
            "tiff": "image/tiff",
            "tif": "image/tiff",
            "wmf": "image/x-wmf",
            "emf": "image/x-emf"
        }

        for ext, content_type in common_image_types.items():
            if ext not in existing_defaults:
                new_default = etree.SubElement(target_root, "{" + target_root.nsmap[None] + "}Default")
                new_default.set("Extension", ext)
                new_default.set("ContentType", content_type)
                existing_defaults.add(ext)
                logging.info(f"Added common image content type mapping for extension: {ext}")

        # Save the updated content types file
        target_tree.write(target_content_types_path, encoding='UTF-8', xml_declaration=True)

    except Exception as e:
        logging.error(f"Error merge_content_types: {str(e)}")
        raise RuntimeError(f"merge_content_types: {str(e)}")

def create_final_docx(extract_dir, output_path):
    """Create a final DOCX file from extracted directory."""
    try:
        # Fix media paths before creating the final docx
        fix_media_paths(extract_dir)

        # Check content types file for completeness
        content_types_path = os.path.join(extract_dir, "[Content_Types].xml")
        ensure_common_content_types(content_types_path)

        # Create a new zip file
        with ZipFile(output_path, 'w', ZIP_DEFLATED) as zipf:
            # First add the content types file
            if os.path.exists(content_types_path):
                zipf.write(content_types_path, "[Content_Types].xml")

            # Then add the _rels folder
            rels_dir = os.path.join(extract_dir, "_rels")
            if os.path.exists(rels_dir):
                for root, dirs, files in os.walk(rels_dir):
                    for file in files:
                        file_path = os.path.join(root, file)
                        rel_path = os.path.relpath(file_path, extract_dir)
                        zipf.write(file_path, rel_path)

            # Then add all other files
            for root, dirs, files in os.walk(extract_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    # Skip the files we've already added
                    if file == "[Content_Types].xml" or root.startswith(rels_dir):
                        continue
                    # Calculate relative path for the zip
                    rel_path = os.path.relpath(file_path, extract_dir)
                    zipf.write(file_path, rel_path)

    except Exception as e:
        logging.error(f"Error create_final_docx: {e}")
        raise RuntimeError(f"create_final_docx: {str(e)}")

def ensure_common_content_types(content_types_path):
    """Make sure common content types are defined in the file."""
    try:
        if not os.path.exists(content_types_path):
            # Create a basic content types file if it doesn't exist
            root_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
            root = etree.Element("{" + root_ns + "}Types", nsmap={None: root_ns})
            tree = etree.ElementTree(root)
        else:
            tree = etree.parse(content_types_path)
            root = tree.getroot()

        # Track existing content types
        existing_extensions = set()

        # Get existing default content types
        for default in root.findall(".//{*}Default"):
            extension = default.get("Extension", "")
            if extension:
                existing_extensions.add(extension.lower())

        # Common content types for Office documents
        common_types = {
            # Text and document formats
            "xml": "application/xml",
            "rels": "application/vnd.openxmlformats-package.relationships+xml",

            # Image formats
            "jpeg": "image/jpeg",
            "jpg": "image/jpeg",
            "png": "image/png",
            "gif": "image/gif",
            "bmp": "image/bmp",
            "tiff": "image/tiff",
            "tif": "image/tiff",
            "wmf": "image/x-wmf",
            "emf": "image/x-emf",

            # Other media formats
            "bin": "application/vnd.openxmlformats-officedocument.oleObject",
        }

        # Add any missing content types
        modified = False
        for ext, content_type in common_types.items():
            if ext.lower() not in existing_extensions:
                new_default = etree.SubElement(root, "{" + root.nsmap[None] + "}Default")
                new_default.set("Extension", ext)
                new_default.set("ContentType", content_type)
                modified = True
                logging.info(f"Added missing content type for extension: {ext}")

        # Save if modified
        if modified:
            tree.write(content_types_path, encoding='UTF-8', xml_declaration=True)

    except Exception as e:
        logging.error(f"Error ensure_common_content_types: {str(e)}")






@app.route(route="split_word")
def split_word(req: func.HttpRequest) -> func.HttpResponse:
    try:
        req_json = req.get_json()

        word_content = req_json.get('file_content', {}).get('$content')
        word_bytes = base64.b64decode(word_content) if word_content else None

        if not word_bytes:
            return func.HttpResponse("Missing Word file content.", status_code=400)

        split_text = req_json.get('split_text')
        split_regex = req_json.get('split_regex')
        split_offset = req_json.get('split_offset', 0)

        if not split_text and not split_regex:
            return func.HttpResponse("Provide either 'split_text' or 'split_regex'.", status_code=400)

        split_base64_strings = split_word_by_text(word_bytes, split_text, split_regex, split_offset)

        response_data = [{
            "$content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "$content": doc_base64
        } for doc_base64 in split_base64_strings]

        return func.HttpResponse(
            body=json.dumps(response_data),
            mimetype="application/json",
            status_code=200
        )
    except Exception as e:
        error_msg = f"Error splitting documents: {str(e)}\n{traceback.format_exc()}"
        logging.exception(error_msg)
        return func.HttpResponse(error_msg, status_code=400)

def split_word_by_text(word_bytes: bytes, split_text: str, split_regex: str, split_offset: int) -> List[str]:
    try:
        split_positions = find_text_positions_in_docx(word_bytes, split_text, split_regex)

        if not split_positions:
            return [base64.b64encode(word_bytes).decode("utf-8")]

        adjusted_positions = sorted(set(max(0, pos + split_offset) for pos in split_positions))

        result_base64_strings = []

        # Ensure elements before the first adjusted split position are not lost
        if adjusted_positions[0] > 0:
            result_base64_strings.append(base64.b64encode(create_word_split(word_bytes, 0, adjusted_positions[0])).decode("utf-8"))

        for i in range(len(adjusted_positions)):
            start_position = adjusted_positions[i]
            end_position = adjusted_positions[i + 1] if i < len(adjusted_positions) - 1 else None
            split_doc_bytes = create_word_split(word_bytes, start_position, end_position)
            result_base64_strings.append(base64.b64encode(split_doc_bytes).decode("utf-8"))

        return result_base64_strings
    except Exception as e:
        logging.error(f"split_word_by_text: {e}")
        raise RuntimeError(f"split_word_by_text: {str(e)}")

def find_text_positions_in_docx(word_bytes: bytes, split_text: str, split_regex: str) -> List[int]:
    try:
        with tempfile.TemporaryDirectory() as temp_dir:
            docx_path = os.path.join(temp_dir, "temp.docx")
            with open(docx_path, "wb") as f:
                f.write(word_bytes)

            extract_dir = os.path.join(temp_dir, "extracted")
            os.mkdir(extract_dir)
            with ZipFile(docx_path, 'r') as zip_ref:
                zip_ref.extractall(extract_dir)

            doc_xml_path = os.path.join(extract_dir, "word", "document.xml")
            parser = etree.XMLParser(recover=True)
            doc_tree = etree.parse(doc_xml_path, parser)
            root = doc_tree.getroot()

            namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            body = root.find(".//w:body", namespaces=namespaces)
            if not body:
                return []

            split_positions = []
            for i, element in enumerate(body):
                text_content = ''.join(element.xpath(".//w:t/text()", namespaces=namespaces))
                if (split_text and split_text in text_content) or (split_regex and re.search(split_regex, text_content)):
                    split_positions.append(i)

            return split_positions
    except Exception as e:
        logging.error(f"find_text_positions_in_docx: {e}")
        raise RuntimeError(f"find_text_positions_in_docx: {str(e)}")

def create_word_split(word_bytes: bytes, start_idx: int, end_idx: int) -> bytes:
    try:
        with tempfile.TemporaryDirectory() as temp_dir:
            docx_path = os.path.join(temp_dir, "original.docx")
            with open(docx_path, "wb") as f:
                f.write(word_bytes)

            extract_dir = os.path.join(temp_dir, "extracted")
            os.mkdir(extract_dir)
            with ZipFile(docx_path, 'r') as zip_ref:
                zip_ref.extractall(extract_dir)

            doc_xml_path = os.path.join(extract_dir, "word", "document.xml")
            parser = etree.XMLParser(recover=True)
            doc_tree = etree.parse(doc_xml_path, parser)
            root = doc_tree.getroot()

            namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            body = root.find(".//w:body", namespaces=namespaces)
            if not body:
                raise ValueError("Document body not found.")

            all_elements = list(body)
            new_body = all_elements[start_idx:end_idx] if end_idx is not None else all_elements[start_idx:]

            for child in list(body):
                body.remove(child)

            for element in new_body:
                body.append(element)

            with open(doc_xml_path, "wb") as f:
                f.write(etree.tostring(doc_tree, encoding="UTF-8", xml_declaration=True))

            new_docx_path = os.path.join(temp_dir, "split.docx")
            with ZipFile(new_docx_path, "w", ZIP_DEFLATED) as new_zip:
                for foldername, subfolders, filenames in os.walk(extract_dir):
                    for filename in filenames:
                        file_path = os.path.join(foldername, filename)
                        arcname = os.path.relpath(file_path, extract_dir)
                        new_zip.write(file_path, arcname)

            with open(new_docx_path, "rb") as f:
                return f.read()
    except Exception as e:
        logging.error(f"create_word_split: {e}")
        raise RuntimeError(f"create_word_split: {str(e)}")






@app.route(route="replace_text_word")
def replace_text_word(req: func.HttpRequest) -> func.HttpResponse:
    try:
        req_body = req.get_json()
        file_content = req_body.get("file_content")
        search_text = req_body.get("search_text")
        replace_text = req_body.get("replace_text")
        search_regex = req_body.get("search_regex")

        if not file_content or "$content" not in file_content:
            return func.HttpResponse("Missing 'file_content' with '$content'", status_code=400)

        if not ((search_text and replace_text is not None) or (search_regex and replace_text is not None)):
            return func.HttpResponse("Must provide either search_text or search_regex with replace_text", status_code=400)

        docx_base64 = file_content["$content"]
        docx_bytes = base64.b64decode(docx_base64)

        # Load Word document
        doc_stream = BytesIO(docx_bytes)
        doc = Document(doc_stream)

        # Process document based on search type
        if search_regex:
            process_document(doc, search_regex, replace_text, is_regex=True)
        else:
            process_document(doc, search_text, replace_text, is_regex=False)

        # Convert modified document back to bytes
        output_stream = BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)
        modified_docx_base64 = base64.b64encode(output_stream.read()).decode("utf-8")

        # Prepare response
        response_content = {
            "$content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "$content": modified_docx_base64
        }

        return func.HttpResponse(json.dumps(response_content), mimetype="application/json")

    except Exception as e:
        error_msg = f"Error replacing text: {str(e)}\n{traceback.format_exc()}"
        logging.exception(error_msg)
        return func.HttpResponse(error_msg, status_code=400)

def process_document(doc, search_pattern, replace_text, is_regex=False):
    """Process entire document for text matches."""
    # Process regular paragraphs
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph, search_pattern, replace_text, is_regex)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph, search_pattern, replace_text, is_regex)

def process_paragraph(paragraph, search_pattern, replace_text, is_regex=False):
    """Process paragraph content with preservation of formatting."""
    # First, check if the text exists in this paragraph
    full_text = paragraph.text
    needs_processing = False

    if is_regex:
        needs_processing = re.search(search_pattern, full_text) is not None
    else:
        needs_processing = search_pattern in full_text

    if not needs_processing:
        return  # Skip processing if no match

    # Handle the paragraph preserving all formatting
    runs = paragraph.runs
    if not runs:
        return

    # Create a map of the paragraph text with formatting details
    text_map = []
    for i, run in enumerate(runs):
        is_hyperlink = is_run_hyperlink(run)
        text_map.append({
            'index': i,
            'text': run.text,
            'bold': run.bold,
            'italic': run.italic, 
            'underline': run.underline,
            'hyperlink': is_hyperlink,
            'run': run
        })

    # Create a composite view of the full text with positions
    composite = []
    pos = 0
    for item in text_map:
        text = item['text']
        if not text:
            continue
        for char_pos, char in enumerate(text):
            composite.append({
                'char': char,
                'pos': pos + char_pos,
                'run_index': item['index'],
                'formatting': {
                    'bold': item['bold'],
                    'italic': item['italic'],
                    'underline': item['underline'],
                    'hyperlink': item['hyperlink']
                }
            })
        pos += len(text)

    # Find matches and make replacements
    matches = []
    if is_regex:
        for match in re.finditer(search_pattern, full_text):
            matches.append((match.start(), match.end(), match.group()))
    else:
        start = 0
        while True:
            idx = full_text.find(search_pattern, start)
            if idx == -1:
                break
            matches.append((idx, idx + len(search_pattern), search_pattern))
            start = idx + 1

    # Handle each match (in reverse to prevent position shifts)
    if matches:
        for match_start, match_end, _ in reversed(matches):
            # Find run indexes for the match
            start_run_idx = None
            end_run_idx = None
            for i, char_info in enumerate(composite):
                if char_info['pos'] == match_start:
                    start_run_idx = char_info['run_index']
                if char_info['pos'] == match_end - 1:
                    end_run_idx = char_info['run_index']
                    break

            if start_run_idx is None or end_run_idx is None:
                continue  # Skip if we can't find the run boundaries

            # Get formatting from the first matched character
            target_formatting = None
            for char_info in composite:
                if char_info['pos'] == match_start:
                    target_formatting = char_info['formatting']
                    break

            # Apply replacement with correct formatting
            replace_in_runs(paragraph, runs, start_run_idx, end_run_idx, 
                           match_start, match_end, replace_text, target_formatting)

def replace_in_runs(paragraph, runs, start_run_idx, end_run_idx, 
                   match_start, match_end, replace_text, formatting):
    """Replace text while preserving formatting and handling complex cases."""
    # Get full paragraph text for position calculations
    full_text = paragraph.text

    # Get text before and after the matched segment
    start_pos = 0
    for i in range(start_run_idx):
        start_pos += len(runs[i].text or '')

    # Calculate offset within the starting run
    start_offset = match_start - start_pos

    # Create a list of operations to perform on runs
    operations = []

    # Handle the starting run
    if start_run_idx == end_run_idx:
        # Match is contained within a single run
        run = runs[start_run_idx]
        run_text = run.text
        before = run_text[:start_offset]
        after = run_text[start_offset + (match_end - match_start):]
        operations.append((start_run_idx, before + replace_text + after))
    else:
        # Match spans multiple runs
        # Handle first run
        run = runs[start_run_idx]
        run_text = run.text
        before = run_text[:start_offset]
        operations.append((start_run_idx, before + replace_text))

        # Clear intermediate runs
        for i in range(start_run_idx + 1, end_run_idx):
            operations.append((i, ""))

        # Handle last run
        run = runs[end_run_idx]
        run_text = run.text
        end_offset = (match_end - start_pos) - sum(len(runs[i].text or '') for i in range(start_run_idx, end_run_idx))
        after = run_text[end_offset:]
        operations.append((end_run_idx, after))

    # Apply operations in reverse order to avoid index shifting
    for run_idx, new_text in sorted(operations, reverse=True):
        # Apply original formatting to the run
        run = runs[run_idx]
        run.text = new_text

        # Restore formatting if available
        if formatting:
            if formatting.get('bold') is not None:
                run.bold = formatting['bold']
            if formatting.get('italic') is not None:
                run.italic = formatting['italic']
            if formatting.get('underline') is not None:
                run.underline = formatting['underline']

def is_run_hyperlink(run):
    """Check if a run is part of a hyperlink."""
    try:
        xml_str = run._r.xml
        return 'w:hyperlink' in xml_str or hasattr(run, '_hyperlink')
    except:
        return False






@app.route(route="word_to_html")
def word_to_html(req: func.HttpRequest) -> func.HttpResponse:
    try:
        # Get the file content and content type from the request
        file_content = req.get_json().get('file_content', None)
        if file_content is None or '$content' not in file_content or '$content-type' not in file_content:
            raise ValueError("No file content or content type provided in the request")

        # Decode the base64 string
        base64_doc = file_content['$content']
        word_file_content = base64.b64decode(base64_doc)

        # Convert the .docx content to HTML using Mammoth
        with BytesIO(word_file_content) as doc_file:
            result = mammoth.convert_to_html(doc_file)
            html_content = result.value  # The generated HTML

        # Inject CSS for table gridlines (borders)
        html_content = add_table_borders(html_content)

        # Return the response JSON object
        return func.HttpResponse(
            str(html_content),
            mimetype="text/html",
            status_code=200
        )

    except Exception as e:
        logging.error(f"Error: {str(e)}\n\nTraceback:\n{traceback.format_exc()}")
        return func.HttpResponse(
            f"Error: {str(e)}\n\nTraceback:\n{traceback.format_exc()}",
            status_code=500
        )

def add_table_borders(html_content: str) -> str:
    """
    Add CSS styles for table gridlines to the generated HTML content.
    """
    # Find all <table> elements and add inline CSS to maintain gridlines
    html_content = html_content.replace(
        '<table>',
        '<table style="border-collapse: collapse; border: 1px solid black;">'
    )

    # Add border style to <td> and <th> elements inside tables
    html_content = html_content.replace(
        '<td>',
        '<td style="border: 1px solid black; padding: 5px;">'
    )
    html_content = html_content.replace(
        '<th>',
        '<th style="border: 1px solid black; padding: 5px;">'
    )

    return html_content






@app.route(route="html_to_word")
def html_to_word(req: func.HttpRequest) -> func.HttpResponse:
    try:
        # Parse request
        req_body = req.get_json()
        html_content = req_body.get('html')

        if not html_content:
            raise ValueError("Invalid input: 'html' field is required")

        # Initialize DOCX document
        doc = Document()
        soup = BeautifulSoup(html_content, 'html.parser')

        # HTML to Word Style Mapping
        header_styles = {
            'h1': 'Heading 1',
            'h2': 'Heading 2',
            'h3': 'Heading 3'
        }

        # Process elements
        for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'table', 'img']):
            if element.name in header_styles:
                doc.add_paragraph(element.get_text(strip=True), style=header_styles[element.name])

            elif element.name == 'p':
                paragraph = doc.add_paragraph()
                for content in element.contents:
                    if isinstance(content, str):
                        paragraph.add_run(content)
                    elif content.name in ['b', 'strong']:
                        paragraph.add_run(content.get_text(strip=True)).bold = True
                    elif content.name in ['i', 'em']:
                        paragraph.add_run(content.get_text(strip=True)).italic = True
                    elif content.name == 'u':
                        paragraph.add_run(content.get_text(strip=True)).underline = True
                    elif content.name == 'a':  # Handle links properly
                        hyperlink = content.get('href')
                        if hyperlink:
                            add_hyperlink(paragraph, hyperlink, content.get_text(strip=True), doc)

            elif element.name == 'table':
                rows = element.find_all('tr')
                if not rows:
                    continue

                cols = max(len(row.find_all(['td', 'th'])) for row in rows)
                table = doc.add_table(rows=0, cols=cols)

                # Apply "Table Grid" style to show borders
                table.style = 'Table Grid'

                for row in rows:
                    cells = row.find_all(['td', 'th'])
                    row_cells = table.add_row().cells
                    for i, cell in enumerate(cells):
                        row_cells[i].text = cell.get_text(strip=True)
                        if cell.name == 'th':  # Make headers bold
                            for run in row_cells[i].paragraphs[0].runs:
                                run.bold = True

            elif element.name == 'img':
                img_src = element.get('src')
                if img_src:
                    if img_src.startswith('http'):
                        try:
                            response = requests.get(img_src, stream=True)
                            response.raise_for_status()

                            # Verify the content is an image
                            content_type = response.headers.get('Content-Type', '')
                            if 'image' not in content_type:
                                raise ValueError(f"Invalid image URL: {img_src}")

                            img_stream = BytesIO(response.content)
                            doc.add_picture(img_stream, width=Inches(4))
                        except requests.RequestException as e:
                            logging.warning(f"Failed to fetch image: {img_src} - {e}")

                    elif img_src.startswith('data:image'):
                        try:
                            header, encoded = img_src.split(',', 1)
                            img_data = base64.b64decode(encoded)
                            img_stream = BytesIO(img_data)
                            doc.add_picture(img_stream, width=Inches(4))
                        except Exception as e:
                            logging.warning(f"Failed to decode base64 image - {e}")

        # Save DOCX file
        byte_io = BytesIO()
        doc.save(byte_io)
        byte_io.seek(0)

        # Base64 encode DOCX
        base64_encoded_docx = base64.b64encode(byte_io.read()).decode('utf-8')

        # Create JSON response
        response = {
            "$content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "$content": base64_encoded_docx
        }

        return func.HttpResponse(json.dumps(response), mimetype='application/json')

    except Exception as e:
        logging.error(f"Error processing request: {e}")
        return func.HttpResponse(
            f"Error: {str(e)}\n\nTraceback:\n{traceback.format_exc()}",
            status_code=500
        )

def add_hyperlink(paragraph, url, text, doc):
    """
    Adds a clickable hyperlink to a paragraph in a Word document.

    :param paragraph: The paragraph where the hyperlink should be added.
    :param url: The hyperlink URL.
    :param text: The visible text for the hyperlink.
    :param doc: The Document object, needed to add relationships.
    """

    # Add hyperlink relationship using the document part
    r_id = doc.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)  # Properly assign the relationship ID

    # Create the run for hyperlink text
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Style: Make the hyperlink blue and underlined
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')  # Blue text
    rPr.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')  # Underline
    rPr.append(underline)

    run.append(rPr)

    # Add hyperlink text
    text_element = OxmlElement('w:t')
    text_element.text = text  # The visible text for the hyperlink
    run.append(text_element)

    # Build the hyperlink
    hyperlink.append(run)
    paragraph._element.append(hyperlink)

    return hyperlink






@app.route(route="resize_image")
def resize_image(req: func.HttpRequest) -> func.HttpResponse:
    try:
        req_body = req.get_json()

        # Validate input
        file_content = req_body.get("file_content")
        if not file_content or "$content" not in file_content or "$content-type" not in file_content:
            return error_response("Invalid input format", 400)

        # Decode base64 image data
        content_type = file_content["$content-type"]
        if content_type not in ["image/png", "image/jpeg", "image/jpg"]:
            return error_response("Unsupported image format", 400)

        image_bytes = base64.b64decode(file_content["$content"])
        image = Image.open(BytesIO(image_bytes))

        # Extract and validate resizing parameters
        width = int(req_body.get("width")) if req_body.get("width") else None
        height = int(req_body.get("height")) if req_body.get("height") else None
        scale = float(req_body.get("scale")) / 100.0 if req_body.get("scale") else None
        max_width = int(req_body.get("max_width")) if req_body.get("max_width") else None
        max_height = int(req_body.get("max_height")) if req_body.get("max_height") else None
        fit_mode = req_body.get("fit_mode", "FIT").upper() # Options: "fit", "fill", "stretch"

        # Validate positive values
        for param_name, param_value in [("width", width), ("height", height), 
                                      ("max_width", max_width), ("max_height", max_height)]:
            if param_value is not None and param_value <= 0:
                return error_response(f"{param_name} must be positive", 400)

        if scale is not None and scale <= 0:
            return error_response("scale must be positive", 400)

        # Only perform resize if parameters are provided
        if any([width, height, scale, max_width, max_height]):
            image = resize_image_with_params(image, width, height, scale, max_width, max_height, fit_mode)

        # Convert image to appropriate format and save to buffer
        output_buffer = BytesIO()
        image_format = "JPEG" if "jpeg" in content_type.lower() or "jpg" in content_type.lower() else "PNG"

        # Convert RGBA to RGB for JPEG
        if image_format == "JPEG" and image.mode in ("RGBA", "P"):
            image = image.convert("RGB")

        # Save with appropriate settings
        save_params = {"format": image_format}
        if image_format == "JPEG":
            save_params["quality"] = 95
        else:  # PNG
            save_params["compress_level"] = 6

        image.save(output_buffer, **save_params)
        output_buffer.seek(0)
        processed_image_bytes = output_buffer.getvalue()

        # Return the resized image
        encoded_image = base64.b64encode(processed_image_bytes).decode("utf-8")
        return func.HttpResponse(
            json.dumps({
                "$content-type": content_type,
                "$content": encoded_image
            }), 
            mimetype="application/json"
        )

    except Exception as e:
        logging.error(f"Image processing error: {str(e)}")
        logging.error(traceback.format_exc())
        return error_response(f"Error splitting documents: {str(e)}\n{traceback.format_exc()}", 500)

def resize_image_with_params(image, width, height, scale, max_width, max_height, fit_mode):
    """Resize image based on provided parameters"""
    if width and height:
        # Width and height specified
        new_size = (width, height)

        if fit_mode == "STRETCH":
            # Simple stretch to target dimensions
            return image.resize(new_size, Image.LANCZOS)

        # Calculate target and image aspect ratios
        img_ratio = image.width / image.height
        target_ratio = width / height

        if fit_mode == "FILL":
            # Crop and resize to fill target dimensions
            if img_ratio > target_ratio:  # Image is wider
                new_height = image.height
                new_width = int(new_height * target_ratio)
            else:  # Image is taller
                new_width = image.width
                new_height = int(new_width / target_ratio)

            # Crop to center
            left = (image.width - new_width) // 2
            top = (image.height - new_height) // 2
            right = left + new_width
            bottom = top + new_height
            image = image.crop((left, top, right, bottom))
            return image.resize(new_size, Image.LANCZOS)

        else:  # fit mode - maintain aspect ratio
            if img_ratio > target_ratio:  # Width is the limiting factor
                new_width = width
                new_height = int(width / img_ratio)
            else:  # Height is the limiting factor
                new_height = height
                new_width = int(height * img_ratio)

            return image.resize((new_width, new_height), Image.LANCZOS)

    elif scale:
        # Scale by percentage
        new_size = (int(image.width * scale), int(image.height * scale))
        return image.resize(new_size, Image.LANCZOS)

    elif max_width or max_height:
        # Constrain to maximum dimensions
        target_width = max_width if max_width else image.width
        target_height = max_height if max_height else image.height

        # Calculate scaling factors
        width_factor = target_width / image.width
        height_factor = target_height / image.height

        # Use the smaller factor to ensure the image fits within constraints
        factor = min(width_factor, height_factor)

        new_size = (int(image.width * factor), int(image.height * factor))
        return image.resize(new_size, Image.LANCZOS)

    return image  # Return original if no resize needed

def error_response(message, status_code):
    """Create a standardized error response"""
    return func.HttpResponse(
        json.dumps({"error": message}),
        status_code=status_code,
        mimetype="application/json"
    )






@app.route(route="compress_image_or_pdf")
def compress_image_or_pdf(req: func.HttpRequest) -> func.HttpResponse:
    try:
        try:
            req_body = req.get_json()
        except ValueError:
            return func.HttpResponse("Invalid JSON", status_code=400)

        compression_rate = req_body.get('compression_rate', 0.8)
        text_pdf_to_jpg = req_body.get('text_pdf_to_jpg', 'NO').upper()  # Standardize input
        file_content = req_body.get('file_content')

        if file_content:
            content_type = file_content.get('$content-type')
            content = file_content.get('$content')

            if content_type and content:
                decoded_content = base64.b64decode(content)

                if 'image' in content_type:
                    compressed_content = convert_and_compress_image(decoded_content, compression_rate)
                elif 'pdf' in content_type:
                    compressed_content = process_pdf(decoded_content, compression_rate, text_pdf_to_jpg)
                else:
                    return func.HttpResponse("Unsupported content type", status_code=400)

                encoded_content = base64.b64encode(compressed_content).decode('utf-8')
                return func.HttpResponse(
                    json.dumps({'$content-type': content_type, '$content': encoded_content}),
                    status_code=200,
                    mimetype="application/json"
                )
            else:
                return func.HttpResponse("Missing content or content type", status_code=400)
        else:
            return func.HttpResponse("Missing file content", status_code=400)
    except Exception as e:
        logging.error(f"Error: {str(e)}\n\nTraceback:\n{traceback.format_exc()}")
        return func.HttpResponse(
            f"Error: {str(e)}\n\nTraceback:\n{traceback.format_exc()}",
            status_code=500
        ) 

def convert_and_compress_image(image_data, compression_rate):
    """
    Converts PNG to JPEG (if needed) and compresses images.
    """
    try:
        with BytesIO(image_data) as input_buffer:
            with Image.open(input_buffer) as img:
                output_buffer = BytesIO()

                # Convert PNG to JPEG
                if img.format == "PNG":
                    img = img.convert("RGB")  # JPEG doesn't support transparency
                    format = "JPEG"
                else:
                    format = img.format  # Keep original format

                # Ensure compression rate
                quality = max(1, min(int(compression_rate * 95), 95)) if format == "JPEG" else None

                img.save(output_buffer, format=format, quality=quality)
                return output_buffer.getvalue()
    except Exception as e:
        logging.error(f"Error processing image: {e}")
        raise

def process_pdf(pdf_data, compression_rate, text_pdf_to_jpg):
    """
    Determines if the PDF is text-based or image-based.
    If text-based and 'text_pdf_to_jpg' is 'YES', convert it to images before compressing.
    """
    try:
        with fitz.open(stream=pdf_data, filetype="pdf") as doc:
            # Check if the PDF contains selectable text
            if any(page.get_text("text") for page in doc):
                if text_pdf_to_jpg == "YES":
                    logging.info("PDF is text-based, but converting to images as requested.")
                    return compress_image_based_pdf(doc, compression_rate)
                logging.info("PDF is text-based. Returning original.")
                return pdf_data  # Return original text-based PDFs

            logging.info("PDF is image-based. Processing images.")
            return compress_image_based_pdf(doc, compression_rate)
    except Exception as e:
        logging.error(f"Error processing PDF: {e}")
        raise

def compress_image_based_pdf(doc, compression_rate):
    """
    Converts all pages of an image-based PDF to JPEG, then recombines into a new PDF.
    """
    try:
        new_doc = fitz.open()  # Create a new blank PDF

        for page in doc:
            pix = page.get_pixmap()
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            # Compress image
            output_buffer = BytesIO()
            img.save(output_buffer, format="JPEG", quality=int(compression_rate * 95))
            compressed_img = output_buffer.getvalue()

            # Create a new blank page in the output PDF
            new_page = new_doc.new_page(width=pix.width, height=pix.height)
            img_rect = fitz.Rect(0, 0, pix.width, pix.height)

            # Insert the compressed image into the new page
            new_page.insert_image(img_rect, stream=compressed_img)

        output_buffer = BytesIO()
        new_doc.save(output_buffer, garbage=4, deflate=True, clean=True, incremental=False)
        return output_buffer.getvalue()
    except Exception as e:
        logging.error(f"Error compressing image-based PDF: {e}")
        raise






@app.route(route="zip_files")
def zip_files(req: func.HttpRequest) -> func.HttpResponse:
    try:
        # Parse the JSON object from the request body
        req_body = req.get_json()
        file_array = req_body.get('file_array', [])

        # Create an in-memory bytes buffer to hold the ZIP file
        zip_buffer = BytesIO()

        # Initialize the ZIP file
        with ZipFile(zip_buffer, 'w', ZIP_DEFLATED) as zip_file:
            for file_obj in file_array:
                # Extract file details
                file_name = file_obj['name']
                file_content_base64 = file_obj['$content']

                # Decode the base64 content
                file_content = base64.b64decode(file_content_base64)

                # Write the file to the ZIP archive
                zip_file.writestr(file_name, file_content)

        # Ensure the buffer's position is at the beginning
        zip_buffer.seek(0)

        # Encode the ZIP file's content to base64
        zip_base64 = base64.b64encode(zip_buffer.read()).decode('utf-8')

        # Create the response object
        response = {
            '$content-type': 'application/zip',
            '$content': zip_base64
        }

        return func.HttpResponse(json.dumps(response), mimetype='application/json')

    except Exception as e:
        return func.HttpResponse(
            f"Error: {str(e)}\n\nTraceback:\n{traceback.format_exc()}",
            status_code=500
        )